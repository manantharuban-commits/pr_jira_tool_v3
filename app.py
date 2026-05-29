"""
PR  Jira Ticket Tool  v3
─────────────────────────────────────────────────────────────────────────────
NEW in v3
  • Word Document Preview  – live HTML-rendered preview panel (screenshots +
    comments shown inline) with Print / Open-in-Word buttons
  • Dynamic Jira Fields    – add / remove / reorder fields; each field has a
    label, type (text / dropdown / date / number), default value, and an
    "include in Jira" toggle
  • Custom fields          – "Add Field" button lets user create any field on
    the fly with label + optional dropdown choices
  • Default-value manager  – Settings → Field Defaults; saved to config.json
    so every fresh launch pre-fills values automatically
─────────────────────────────────────────────────────────────────────────────
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext, simpledialog
import json, os, re, threading, webbrowser, base64, copy, tempfile, html
from datetime import datetime

import requests
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
#  Paths & palette
# ─────────────────────────────────────────────────────────────────────────────
BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE        = os.path.join(BASE_DIR, "config.json")
FIELD_VALUES_FILE  = os.path.join(BASE_DIR, "field_values.json")
KEYWORDS_FILE      = os.path.join(BASE_DIR, "keywords.json")

APP_NAME     = "PRism"
APP_SUBTITLE = "PR  →  Jira · Doc · SQL"


def _draw_prism(d, sz):
    """Draw PRism logo onto an ImageDraw canvas of size sz×sz (RGB, dark bg)."""
    cx, cy = sz // 2 - sz // 25, sz // 2
    # prism body
    tl   = (cx - sz * 24 // 100, cy - sz * 30 // 100)
    bl   = (cx - sz * 24 // 100, cy + sz * 30 // 100)
    apex = (cx + sz * 30 // 100, cy)
    d.polygon([tl, bl, apex], fill=(30, 40, 60))
    bw = max(2, sz // 50)
    d.polygon([tl, bl, apex], outline=(88, 166, 255), width=bw)
    # incoming beam
    d.line([(sz * 4 // 100, cy), (tl[0], cy)],
           fill=(210, 225, 255), width=max(2, sz // 60))
    # refracted rays
    ex = sz - sz * 6 // 100
    for ey, col in [
        (cy - sz * 25 // 100, (63, 185, 80)),
        (cy,                  (88, 166, 255)),
        (cy + sz * 25 // 100, (210, 153, 34)),
    ]:
        lw = max(2, sz // 48)
        d.line([(apex[0], apex[1]), (ex, ey)], fill=col, width=lw)
    return tl, bl, apex


def _generate_app_assets():
    """Generate prism_icon.ico + prism_logo.png in BASE_DIR.
    Uses RGB images (no alpha) for maximum ICO compatibility.
    Returns (logo_path, icon_path) or (None, None) on failure."""
    try:
        from PIL import Image, ImageDraw, ImageFont

        icon_path = os.path.join(BASE_DIR, "prism_icon.ico")
        logo_path = os.path.join(BASE_DIR, "prism_logo.png")
        BG = (13, 17, 23)   # #0d1117

        # Build one image per ICO size to keep quality at small sizes
        ico_imgs = []
        for sz in (256, 128, 64, 48, 32, 16):
            img = Image.new("RGB", (sz, sz), BG)
            d   = ImageDraw.Draw(img)
            # rounded-square bg tint
            pad = max(1, sz // 20)
            try:
                d.rounded_rectangle([pad, pad, sz - pad - 1, sz - pad - 1],
                                    radius=sz // 7, fill=(22, 27, 34))
            except AttributeError:
                d.rectangle([pad, pad, sz - pad - 1, sz - pad - 1],
                            fill=(22, 27, 34))
            _draw_prism(d, sz)
            # "PR" text only for larger sizes
            if sz >= 48:
                fnt = None
                for fp in ["C:/Windows/Fonts/arialbd.ttf",
                           "C:/Windows/Fonts/calibrib.ttf",
                           "C:/Windows/Fonts/segoeui.ttf"]:
                    try:
                        fnt = ImageFont.truetype(fp, sz * 22 // 100)
                        break
                    except Exception:
                        pass
                if fnt:
                    tx = sz * 18 // 100
                    ty = sz * 34 // 100
                    d.text((tx, ty), "PR", font=fnt, fill=(255, 255, 255))
            ico_imgs.append(img)

        # Save ICO — first image is primary, rest are alternate sizes
        ico_imgs[0].save(
            icon_path, format="ICO",
            append_images=ico_imgs[1:],
            sizes=[(i.width, i.height) for i in ico_imgs],
        )

        # Save PNG logo at 256
        ico_imgs[0].save(logo_path, format="PNG")
        return logo_path, icon_path

    except Exception as exc:
        return None, None

THEMES = {
    "dark": dict(
        bg         = "#0d1117",
        surface    = "#161b22",
        card       = "#1c2128",
        card2      = "#21262d",
        border     = "#30363d",
        shadow     = "#090c10",
        accent     = "#58a6ff",
        accent_dim = "#132236",
        green      = "#3fb950",
        yellow     = "#d29922",
        red        = "#f85149",
        purple     = "#bc8cff",
        orange     = "#e3b341",
        text       = "#e6edf3",
        muted      = "#8b949e",
        inp        = "#0d1117",
        hover      = "#1f6feb",
        cb_sel     = "#2ea043",
        tag        = "dark",
        st_badge   = dict(added="#1f6931", modified="#1c4a6e", removed="#8a1f1f",
                          renamed="#5a2d82", copied="#1c4a6e", changed="#1c4a6e"),
        st_row     = dict(added="#0e2318", modified="#0c1c2c", removed="#200e0e",
                          renamed="#180e2a", copied="#0c1c2c", changed="#0c1c2c"),
    ),
    "light": dict(
        bg         = "#f6f8fa",
        surface    = "#ffffff",
        card       = "#ffffff",
        card2      = "#f6f8fa",
        border     = "#d0d7de",
        shadow     = "#b8c0c8",
        accent     = "#0969da",
        accent_dim = "#ddeeff",
        green      = "#1a7f37",
        yellow     = "#9a6700",
        red        = "#cf222e",
        purple     = "#8250df",
        orange     = "#bc6800",
        text       = "#1f2328",
        muted      = "#57606a",
        inp        = "#ffffff",
        hover      = "#0550ae",
        cb_sel     = "#cce5ff",
        tag        = "light",
        st_badge   = dict(added="#1a7f37", modified="#0969da", removed="#cf222e",
                          renamed="#8250df", copied="#0969da", changed="#0969da"),
        st_row     = dict(added="#eafbee", modified="#eef6ff", removed="#fff5f5",
                          renamed="#f7f0ff", copied="#eef6ff", changed="#eef6ff"),
    ),
}

C = dict(THEMES["dark"])   # mutable — updated in-place on theme switch

STATUS_EMOJI  = dict(added="A", modified="M", removed="D", renamed="R", copied="C", changed="~")
STATUS_BADGE  = dict(C["st_badge"])
STATUS_ROW    = dict(C["st_row"])

STATUS_VERB   = dict(added="Add", modified="Update", removed="Remove",
                     renamed="Rename", copied="Copy", changed="Change")

FIELD_HINTS = {
    "owner":       "Person responsible for resolving or triaging this issue.",
    "environment": "Environment where the issue was discovered (e.g. Production, Staging).",
    "issue_type":  "Jira issue category — Bug, Task, Story, Epic, etc.",
    "dependency":  "Enter 'Yes' if this ticket depends on another PR or ticket.",
    "git_link":    "GitHub Pull Request URL — auto-filled when a PR is fetched.",
    "doc_name":    "Output file name for the generated Word document (no .docx extension). Auto-synced from Jira Summary.",
    "doc_title":   "Heading shown at the top of the Word document. Auto-synced from Jira Summary; edit to override.",
    "summary":     "Jira ticket title — becomes the Issue Summary field in Jira.",
    "priority":    "Priority level shown in Jira (Highest → Lowest).",
    "component":   "Application component or module affected by this issue.",
    "fix_version": "Release version in which this issue is fixed.",
}

# ─────────────────────────────────────────────────────────────────────────────
#  Tooltip helper
# ─────────────────────────────────────────────────────────────────────────────

class Tooltip:
    """Lightweight hover tooltip attached to any widget."""
    def __init__(self, widget, text):
        self._w    = widget
        self._text = text
        self._tip  = None
        widget.bind("<Enter>",       self._show, "+")
        widget.bind("<Leave>",       self._hide, "+")
        widget.bind("<ButtonPress>", self._hide, "+")

    def _show(self, _=None):
        if self._tip or not self._text:
            return
        wx, wy = self._w.winfo_rootx(), self._w.winfo_rooty()
        wh     = self._w.winfo_height()
        self._tip = tk.Toplevel(self._w)
        self._tip.wm_overrideredirect(True)
        self._tip.wm_geometry(f"+{wx + 16}+{wy + wh + 6}")
        tk.Label(self._tip, text=self._text,
                 bg="#24292f", fg="#e6edf3",
                 font=("Segoe UI", 9), relief="flat",
                 padx=10, pady=7, wraplength=300,
                 justify="left").pack()

    def _hide(self, _=None):
        if self._tip:
            self._tip.destroy()
            self._tip = None

# ─────────────────────────────────────────────────────────────────────────────
#  Dependency-order scoring for PR files
# ─────────────────────────────────────────────────────────────────────────────

def _dep_score(filename):
    """Lower score = fewer dependencies = should appear first in the step list."""
    fname = filename.lower()
    base  = os.path.basename(fname)
    name  = os.path.splitext(base)[0]

    # Test files first — a file named test_auth_service.py is a test, not a service
    if any(x in fname for x in ['test', 'spec', '__tests__', '.test.', '.spec.']):
        return 5
    if any(x in name for x in ['config', 'const', 'constant', 'setting', 'settings',
                                 'setup', 'init', 'env', 'properties', 'defaults']):
        return 0
    if any(x in name for x in ['type', 'types', 'model', 'models', 'schema', 'schemas',
                                 'interface', 'interfaces', 'enum', 'enums', 'dto', 'entity']):
        return 1
    if any(x in name for x in ['util', 'utils', 'helper', 'helpers', 'common',
                                 'shared', 'base', 'mixin', 'lib', 'libs', 'core']):
        return 2
    if any(x in name for x in ['service', 'services', 'api', 'client', 'repo',
                                 'repository', 'store', 'stores', 'hook', 'hooks', 'dao']):
        return 3
    if any(x in name for x in ['component', 'controller', 'view', 'page',
                                 'screen', 'widget', 'handler', 'middleware', 'route']):
        return 4
    return 3   # default: mid-tier

# ─────────────────────────────────────────────────────────────────────────────
#  Built-in field definitions  (label, key, type, choices, default, jira_key)
# ─────────────────────────────────────────────────────────────────────────────
BUILTIN_FIELDS = [
    # ── Jira ticket fields (sent directly to Jira API, shown in "Jira Fields" card) ──
    {"label": "Project",      "key": "project",     "type": "text",
     "choices": [], "default": "", "jira_key": "project",    "jira_field": True,
     "required": True,  "enabled": True, "show_label_in_jira": False},
    {"label": "Jira Summary", "key": "summary",     "type": "text",
     "choices": [], "default": "", "jira_key": "summary",    "jira_field": True,
     "required": True,  "enabled": True, "show_label_in_jira": False},
    {"label": "Issue Type",   "key": "issue_type",  "type": "dropdown",
     "choices": ["Bug","Enhancement","Task","Story","Epic","Sub-task","Incident"],
     "default": "Bug", "jira_key": "issuetype",              "jira_field": True,
     "required": True,  "enabled": True, "show_label_in_jira": False},
    {"label": "Reporter",          "key": "reporter",       "type": "text",
     "choices": [], "default": "", "jira_key": "reporter",           "jira_field": True,
     "required": False, "enabled": True,  "show_label_in_jira": False},
    # ── Additional Jira API fields (disabled by default — enable via ⊞ Jira Fields button) ──
    {"label": "Assignee",          "key": "assignee",       "type": "text",
     "choices": [], "default": "", "jira_key": "assignee",           "jira_field": True,
     "required": False, "enabled": True,  "show_label_in_jira": False},
    {"label": "Priority",          "key": "jira_priority",  "type": "dropdown",
     "choices": ["Highest","High","Medium","Low","Lowest"],
     "default": "Medium",           "jira_key": "priority",          "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Labels",            "key": "labels",         "type": "text",
     "choices": [], "default": "", "jira_key": "labels",             "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Components",        "key": "jira_components","type": "text",
     "choices": [], "default": "", "jira_key": "components",         "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Fix Versions",      "key": "jira_fix_versions","type": "text",
     "choices": [], "default": "", "jira_key": "fixVersions",        "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Affects Versions",  "key": "aff_versions",   "type": "text",
     "choices": [], "default": "", "jira_key": "versions",           "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Due Date",          "key": "due_date",       "type": "date",
     "choices": [], "default": "", "jira_key": "duedate",            "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Environment",       "key": "jira_env",       "type": "text",
     "choices": [], "default": "", "jira_key": "environment",        "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Epic Link",         "key": "epic_link",      "type": "text",
     "choices": [], "default": "", "jira_key": "customfield_10014",  "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Sprint",            "key": "sprint",         "type": "text",
     "choices": [], "default": "", "jira_key": "customfield_10020",  "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Story Points",      "key": "story_points",   "type": "number",
     "choices": [], "default": "", "jira_key": "customfield_10016",  "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Original Estimate", "key": "orig_estimate",  "type": "text",
     "choices": [], "default": "", "jira_key": "timeoriginalestimate","jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Parent Issue",      "key": "parent_issue",   "type": "text",
     "choices": [], "default": "", "jira_key": "parent",             "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Security Level",    "key": "security_level", "type": "text",
     "choices": [], "default": "", "jira_key": "security",           "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    {"label": "Watchers",          "key": "watchers",       "type": "text",
     "choices": [], "default": "", "jira_key": "watches",            "jira_field": True,
     "required": False, "enabled": False, "show_label_in_jira": False},
    # ── Description-only fields (appear in Jira ticket body as a table, not as top-level API fields) ──
    {"label": "Dependency",   "key": "dependency",  "type": "text",
     "choices": [], "default": "", "jira_key": "dependency",
     "required": False, "enabled": True, "show_label_in_jira": True},
    {"label": "Git PR Link",  "key": "git_link",    "type": "text",
     "choices": [], "default": "", "jira_key": "git_link",
     "required": False, "enabled": True, "show_label_in_jira": True},
    {"label": "Doc File Name", "key": "doc_name",    "type": "text",
     "choices": [], "default": "", "jira_key": None,
     "required": False, "enabled": True, "show_label_in_jira": False},
    {"label": "Doc Title",    "key": "doc_title",   "type": "text",
     "choices": [], "default": "", "jira_key": None,
     "required": False, "enabled": True, "show_label_in_jira": False},
    {"label": "Priority",     "key": "priority",    "type": "dropdown",
     "choices": ["Highest","High","Medium","Low","Lowest"],
     "default": "Medium", "jira_key": "priority",
     "required": False, "enabled": True, "show_label_in_jira": True},
    {"label": "Component",    "key": "component",   "type": "text",
     "choices": [], "default": "", "jira_key": "components",
     "required": False, "enabled": False, "show_label_in_jira": True},
    {"label": "Fix Version",  "key": "fix_version", "type": "text",
     "choices": [], "default": "", "jira_key": "fixVersions",
     "required": False, "enabled": False, "show_label_in_jira": True},
]

DEFAULT_CONFIG = {
    "github_token_file":   os.path.join(BASE_DIR, "github_token.txt"),
    "jira_token_file":     os.path.join(BASE_DIR, "jira_token.txt"),
    "jira_base_url":       "https://yourcompany.atlassian.net/rest/api/3",
    "jira_project_key":    "PROJ",
    "jira_email":          "you@company.com",
    "github_owner":        "",
    "github_repo":         "",
    "github_api_url":      "https://api.github.com",
    "word_doc_output_dir": BASE_DIR,
    "ssl_cert_file":       os.path.join(BASE_DIR, "certs", "dummy.pem"),
    "jira_ssl_cert_file":  os.path.join(BASE_DIR, "certs", "jira_dummy.pem"),
    "jira_ssl_verify":     True,
    "jira_auth_method":    "basic",
    "github_ssl_verify":   True,
    "fields":              BUILTIN_FIELDS,
    "field_defaults":      {},   # key → default_value overrides
    "theme":               "dark",
    "auto_open_word_doc":  False,
    "attach_word_doc":     True,
    "attach_sql_file":     True,
}

# ─────────────────────────────────────────────────────────────────────────────
#  Config helpers
# ─────────────────────────────────────────────────────────────────────────────

def load_config():
    cfg = copy.deepcopy(DEFAULT_CONFIG)
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE) as f:
            saved = json.load(f)
        # apply non-field settings (tokens, urls, etc.)
        for k, v in saved.items():
            if k != "fields":
                cfg[k] = v
        # merge saved per-field overrides (enabled, default) into BUILTIN_FIELDS
        saved_by_key = {fd["key"]: fd for fd in saved.get("fields", [])}
        for fd in cfg["fields"]:
            if fd["key"] in saved_by_key:
                sf = saved_by_key[fd["key"]]
                if "enabled" in sf:
                    fd["enabled"] = sf["enabled"]
                if "default" in sf:
                    fd["default"] = sf["default"]
        # also apply field_defaults overrides
        for fd in cfg["fields"]:
            if fd["key"] in cfg.get("field_defaults", {}):
                fd["default"] = cfg["field_defaults"][fd["key"]]
    return cfg

def save_config(cfg):
    with open(CONFIG_FILE, "w") as f:
        json.dump(cfg, f, indent=2)

def load_field_values():
    if os.path.exists(FIELD_VALUES_FILE):
        try:
            with open(FIELD_VALUES_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}

def save_field_values(values):
    try:
        with open(FIELD_VALUES_FILE, "w", encoding="utf-8") as f:
            json.dump(values, f, indent=2, ensure_ascii=False)
    except Exception:
        pass

def load_keywords():
    if os.path.exists(KEYWORDS_FILE):
        try:
            with open(KEYWORDS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, list):
                    return data
        except Exception:
            pass
    return ["work_prod"]

def save_keywords(kws):
    try:
        with open(KEYWORDS_FILE, "w", encoding="utf-8") as f:
            json.dump(kws, f, indent=2, ensure_ascii=False)
    except Exception:
        pass

def read_token(path):
    path = os.path.expanduser(path)
    if not os.path.exists(path):
        raise FileNotFoundError(f"Token file not found:\n{path}")
    with open(path) as f:
        return f.read().strip()

def gh_headers(token):
    return {
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    }

def jira_auth_headers(email, token, method="basic"):
    if method in ("bearer", "oauth"):
        return {"Authorization": f"Bearer {token}",
                "Content-Type": "application/json", "Accept": "application/json"}
    creds = base64.b64encode(f"{email}:{token}".encode()).decode()
    return {"Authorization": f"Basic {creds}",
            "Content-Type": "application/json", "Accept": "application/json"}

def _ssl_verify(cfg):
    """GitHub SSL. Returns False (skip), cert path, or True (system CA)."""
    cfg = cfg or {}
    if not cfg.get("github_ssl_verify", True):
        return False
    p = cfg.get("ssl_cert_file", "")
    if p and os.path.isfile(p):
        return p
    return True

def _jira_ssl_verify(cfg):
    """Jira SSL. Returns False (skip), cert path, or True (system CA)."""
    cfg = cfg or {}
    if not cfg.get("jira_ssl_verify", True):
        return False
    p = cfg.get("jira_ssl_cert_file", "")
    if p and os.path.isfile(p):
        return p
    return True

# ─────────────────────────────────────────────────────────────────────────────
#  SQL helpers
# ─────────────────────────────────────────────────────────────────────────────

def _detect_sql_object(filename, patch):
    """Return a short label like 'TABLE users' or 'VIEW active_orders' from
    the patch/filename.  Used as the default step comment in the SQL popup."""
    # Gather added/changed lines from patch (first 60) for detection
    search = ""
    if patch:
        added = [l[1:] for l in patch.splitlines()
                 if l.startswith('+') and not l.startswith('+++')]
        search = "\n".join(added[:60])
    if not search:
        search = os.path.basename(filename)

    DDL = (r'(CREATE|ALTER|DROP)\s+(?:OR\s+REPLACE\s+)?(?:FORCE\s+)?'
           r'(?:EDITIONABLE\s+)?'
           r'(TABLE|VIEW|PROCEDURE|PROC|FUNCTION|TRIGGER|PACKAGE|SEQUENCE|'
           r'TYPE|INDEX|SCHEMA)\s+(?:\w+\.)?(\w+)')
    m = re.search(DDL, search, re.IGNORECASE | re.MULTILINE)
    if m:
        return f"{m.group(2).upper()} {m.group(3)}"

    DML = r'(?:INSERT\s+INTO|UPDATE|DELETE\s+FROM)\s+(?:\w+\.)?(\w+)'
    m = re.search(DML, search, re.IGNORECASE)
    if m:
        return m.group(1)

    base = os.path.splitext(os.path.basename(filename))[0]
    base = re.sub(r'^\d+[_\-]+', '', base)   # strip leading "01_"
    return base


# ─────────────────────────────────────────────────────────────────────────────
#  Topological file sorter
# ─────────────────────────────────────────────────────────────────────────────

_SQL_EXTS = {'.sql', '.prc', '.fnc', '.trg', '.pks', '.pkb', '.vw', '.typ', '.sps', '.spb'}

_DDL_DEFINE = re.compile(
    r'(?:CREATE|ALTER)\s+(?:OR\s+REPLACE\s+)?(?:FORCE\s+)?(?:EDITIONABLE\s+)?'
    r'(?:TABLE|VIEW|MATERIALIZED\s+VIEW|PROCEDURE|PROC|FUNCTION|TRIGGER|'
    r'PACKAGE(?:\s+BODY)?|SEQUENCE|TYPE(?:\s+BODY)?|SYNONYM|INDEX)\s+'
    r'(?:\w+\.)?(\w+)',
    re.IGNORECASE | re.MULTILINE,
)

# Word boundary: Oracle identifiers use [A-Za-z0-9_$#]
_WB  = r'(?<![A-Za-z0-9_$#])'
_WBA = r'(?![A-Za-z0-9_$#])'


def _file_current_text(f):
    """Return the 'new state' text of a file from its patch or full_content."""
    if f.get("full_content"):
        return f["full_content"]
    patch = f.get("patch", "") or ""
    lines = []
    for ln in patch.splitlines():
        if ln.startswith(("+++", "---", "@@")):
            continue
        if ln.startswith("-"):
            continue              # old content — skip
        lines.append(ln[1:] if ln.startswith("+") else ln)
    return "\n".join(lines)


def _topo_sort_files(files):
    """
    Sort PR files so dependencies come before dependants.

    Algorithm:
      1. For each SQL file detect which DB object it defines (via DDL regex).
      2. For each file scan its new-state content for references to objects
         defined by *other* files in this PR.
      3. Topological sort (Kahn's) — within each wave, order by _dep_score.
      4. Non-SQL files and any cycle remainder fall back to _dep_score.
    """
    import heapq

    n = len(files)
    if n <= 1:
        return list(files)

    is_sql = [
        os.path.splitext(f.get("filename", ""))[1].lower() in _SQL_EXTS
        for f in files
    ]

    # ── 1. detect defined object name (SQL files only) ─────────────────────
    texts = [_file_current_text(f) for f in files]

    obj_names = []   # None for non-SQL or undetectable
    for i, f in enumerate(files):
        if not is_sql[i]:
            obj_names.append(None)
            continue
        m = _DDL_DEFINE.search(texts[i])
        if m:
            obj_names.append(m.group(1).lower())
        else:
            # fallback: filename stem (stripped of leading numbers)
            base = os.path.splitext(os.path.basename(f.get("filename", "")))[0]
            obj_names.append(re.sub(r'^\d+[_\-]+', '', base).lower() or None)

    # obj → file index (only for SQL files with a name)
    obj_to_idx = {name: i for i, name in enumerate(obj_names) if name}

    # ── 2. build dependency edges (file i depends on file j) ───────────────
    deps = [set() for _ in range(n)]   # deps[i] = set of j that i needs first
    for i, text in enumerate(texts):
        if not is_sql[i]:
            continue
        for obj, j in obj_to_idx.items():
            if j == i:
                continue
            if re.search(_WB + re.escape(obj) + _WBA, text, re.IGNORECASE):
                deps[i].add(j)

    # ── 3. Kahn's topological sort with _dep_score tiebreaker ──────────────
    in_deg = [len(deps[i]) for i in range(n)]
    rdeps  = [[] for _ in range(n)]
    for i in range(n):
        for j in deps[i]:
            rdeps[j].append(i)

    heap = []
    for i in range(n):
        if in_deg[i] == 0:
            heapq.heappush(heap, (_dep_score(files[i].get("filename", "")), i))

    order = []
    while heap:
        _, i = heapq.heappop(heap)
        order.append(i)
        for j in rdeps[i]:
            in_deg[j] -= 1
            if in_deg[j] == 0:
                heapq.heappush(heap, (_dep_score(files[j].get("filename", "")), j))

    # ── 4. handle cycles — append remainder by _dep_score ──────────────────
    if len(order) < n:
        used = set(order)
        rest = sorted(
            [i for i in range(n) if i not in used],
            key=lambda i: (_dep_score(files[i].get("filename", "")),
                           files[i].get("filename", "")),
        )
        order.extend(rest)

    return [files[i] for i in order]


# ─────────────────────────────────────────────────────────────────────────────
#  GitHub-style diff renderer for Word docs
# ─────────────────────────────────────────────────────────────────────────────

def _word_diff_block(doc, patch, status, adds, dels):
    """Print-friendly diff: light backgrounds, dark text readable on white paper."""
    # colour palette (light backgrounds, dark foreground text)
    BG_ADD  = "E6FFED"; FG_ADD  = RGBColor(0x1a, 0x7f, 0x37)  # light green / dark green
    BG_DEL  = "FFEBE9"; FG_DEL  = RGBColor(0xcf, 0x22, 0x2e)  # light red   / dark red
    BG_MOD  = "FFFBDD"; FG_MOD  = RGBColor(0x7c, 0x4a, 0x00)  # light amber / dark amber
    BG_HUNK = "DDF4FF"; FG_HUNK = RGBColor(0x09, 0x69, 0xda)  # light blue  / dark blue
    BG_CTX  = "F6F8FA"; FG_CTX  = RGBColor(0x57, 0x60, 0x6a)  # light grey  / dark grey
    BG_NONE = "F6F8FA"; FG_NONE = RGBColor(0x57, 0x60, 0x6a)

    def _sp(text, bg_hex, fg_rgb, bold=False):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  bg_hex)
        pPr.append(shd)
        r = p.add_run(text)
        r.font.name      = "Consolas"
        r.font.size      = Pt(8)
        r.font.color.rgb = fg_rgb
        r.bold           = bold
        return p

    if status in ("removed", "deleted"):
        _sp(f"  −  FILE DELETED  —  {dels} line{'s' if dels != 1 else ''} removed",
            BG_DEL, FG_DEL, bold=True)
        if patch:
            old_ln = 1
            for raw in patch.splitlines():
                if raw.startswith('@@'):
                    m2 = re.match(r'@@ -(\d+)(?:,\d+)? \+(\d+)(?:,\d+)? @@', raw)
                    if m2:
                        old_ln = int(m2.group(1))
                    _sp(f"  {raw}", BG_HUNK, FG_HUNK)
                elif raw.startswith('-') and not raw.startswith('---'):
                    _sp(f"  {old_ln:>4}  −  {raw[1:]}", BG_DEL, FG_DEL)
                    old_ln += 1
        return

    if not patch:
        _sp("     (Diff not available — file too large or binary)", BG_NONE, FG_NONE)
        return

    if status == "added":
        _sp(f"  +  NEW FILE  —  {adds} line{'s' if adds != 1 else ''} added",
            BG_ADD, FG_ADD, bold=True)
        new_ln = 1
        for raw in patch.splitlines():
            if raw.startswith('+') and not raw.startswith('+++'):
                _sp(f"  {new_ln:>4}  +  {raw[1:]}", BG_ADD, FG_ADD)
                new_ln += 1
        return

    # modified / renamed
    old_ln = new_ln = 0
    has_changes = False
    lines = patch.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        if raw.startswith('@@'):
            m = re.match(r'@@ -(\d+)(?:,\d+)? \+(\d+)(?:,\d+)? @@(.*)', raw)
            if m:
                old_ln = int(m.group(1))
                new_ln = int(m.group(2))
                _sp(f"  {raw}", BG_HUNK, FG_HUNK)
            i += 1
        elif (raw.startswith('-') and not raw.startswith('---')) or \
             (raw.startswith('+') and not raw.startswith('+++')):
            block_dels, block_adds = [], []
            while i < len(lines) and lines[i].startswith('-') and not lines[i].startswith('---'):
                block_dels.append(lines[i][1:])
                old_ln += 1
                i += 1
            new_ln_start = new_ln
            while i < len(lines) and lines[i].startswith('+') and not lines[i].startswith('+++'):
                block_adds.append(lines[i][1:])
                new_ln += 1
                i += 1
            for j, add_text in enumerate(block_adds):
                ln = new_ln_start + j
                if j < len(block_dels):
                    _sp(f"  {ln:>4}  ~  {add_text}", BG_MOD, FG_MOD)
                else:
                    _sp(f"  {ln:>4}  +  {add_text}", BG_ADD, FG_ADD)
            has_changes = has_changes or bool(block_dels or block_adds)
        elif raw.startswith(' '):
            old_ln += 1; new_ln += 1
            i += 1
        else:
            i += 1

    if not has_changes:
        _sp("     (No changed lines detected)", BG_NONE, FG_NONE)


# ─────────────────────────────────────────────────────────────────────────────
#  Word document generator
# ─────────────────────────────────────────────────────────────────────────────

def generate_word_doc(field_values, field_defs, pr_data, file_comments, output_dir,
                      keyword_findings=None):
    # ── Colour palette ────────────────────────────────────────────────────────
    C_PRIMARY   = RGBColor(0x1F, 0x38, 0x64)   # Deep Navy
    C_SECONDARY = RGBColor(0x2E, 0x75, 0xB6)   # Steel Blue
    C_BODY      = RGBColor(0x40, 0x40, 0x40)   # Charcoal
    C_MUTED     = RGBColor(0x7F, 0x7F, 0x7F)   # Mid Gray
    C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
    HEX_PRIMARY = "1F3864"
    HEX_LIGHT   = "EAF0F8"   # alternating row tint
    HEX_WHITE   = "FFFFFF"
    HEX_RULE    = "BDD7EE"   # subtle blue-gray border

    doc = Document()

    # ── Page setup — 1 inch margins ───────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

    # ── Normal style — Calibri 11pt charcoal, 1.15 line spacing ──────────────
    normal = doc.styles["Normal"]
    normal.font.name  = "Calibri"
    normal.font.size  = Pt(11)
    normal.font.color.rgb = C_BODY
    normal.paragraph_format.space_after      = Pt(7)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing      = 1.15

    # ── XML helpers ───────────────────────────────────────────────────────────
    def _shade_para(p, hex_fill):
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        pPr.append(shd)

    def _shade_cell(cell, hex_fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        tcPr.append(shd)

    def _bottom_border(p, color=HEX_RULE, sz="4"):
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), sz)
        bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), color)
        pBdr.append(bot); pPr.append(pBdr)

    def _left_border(p, color, sz="18"):
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        lft  = OxmlElement("w:left")
        lft.set(qn("w:val"), "single"); lft.set(qn("w:sz"), sz)
        lft.set(qn("w:space"), "4");    lft.set(qn("w:color"), color)
        pBdr.append(lft); pPr.append(pBdr)

    def _fld_run(p, instr, color=C_MUTED):
        fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
        it  = OxmlElement("w:instrText"); it.text = instr
        fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
        r   = p.add_run(); r.font.size = Pt(9); r.font.color.rgb = color
        r._r.append(fc1); r._r.append(it); r._r.append(fc2)

    # ── Header — document title, muted gray ──────────────────────────────────
    doc_title_val = (field_values.get("doc_title") or "").strip() or \
                    (pr_data.get("title") if pr_data else None) or "Issue Report"
    hdr_p = section.header.paragraphs[0]
    hdr_p.clear()
    hdr_p.paragraph_format.space_after = Pt(0)
    _bottom_border(hdr_p, color=HEX_RULE, sz="4")
    rh = hdr_p.add_run(doc_title_val)
    rh.font.name = "Calibri"; rh.font.size = Pt(9); rh.font.color.rgb = C_MUTED

    # ── Footer — page N of M, centered ───────────────────────────────────────
    ftr_p = section.footer.paragraphs[0]
    ftr_p.clear()
    ftr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ftr_p.paragraph_format.space_before = Pt(0)
    rpg = ftr_p.add_run("Page "); rpg.font.size = Pt(9); rpg.font.color.rgb = C_MUTED
    _fld_run(ftr_p, " PAGE ")
    rof = ftr_p.add_run(" of "); rof.font.size = Pt(9); rof.font.color.rgb = C_MUTED
    _fld_run(ftr_p, " NUMPAGES ")

    # ── Section heading helpers ───────────────────────────────────────────────
    def _h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = "Calibri"; r.font.size = Pt(16); r.bold = True
        r.font.color.rgb = C_PRIMARY
        _bottom_border(p, color=HEX_RULE, sz="4")
        return p

    def _h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = "Calibri"; r.font.size = Pt(13); r.bold = True
        r.font.color.rgb = C_SECONDARY
        return p

    # ── Table styling — navy header + zebra rows ──────────────────────────────
    def _build_table(rows_data, col_headers=None):
        start = 0
        if col_headers:
            tbl = doc.add_table(rows=1, cols=len(col_headers))
            tbl.style = "Table Grid"
            for ci, hdr_txt in enumerate(col_headers):
                cell = tbl.rows[0].cells[ci]
                cell.text = hdr_txt
                _shade_cell(cell, HEX_PRIMARY)
                for p in cell.paragraphs:
                    if p.runs:
                        p.runs[0].font.color.rgb = C_WHITE
                        p.runs[0].bold = True
                        p.runs[0].font.size = Pt(10)
        else:
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Table Grid"
        for ri, (label, value) in enumerate(rows_data):
            row = tbl.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value) if value else "N/A"
            bg = HEX_LIGHT if ri % 2 == 0 else HEX_WHITE
            for cell in row.cells:
                _shade_cell(cell, bg)
            for p in row.cells[0].paragraphs:
                if p.runs:
                    p.runs[0].bold = True
                    p.runs[0].font.color.rgb = C_PRIMARY
                    p.runs[0].font.size = Pt(10)
            for p in row.cells[1].paragraphs:
                if p.runs:
                    p.runs[0].font.color.rgb = C_BODY
                    p.runs[0].font.size = Pt(10)
        return tbl

    # ── Title block ───────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(4)
    rt = title_p.add_run(doc_title_val)
    rt.font.name = "Calibri"; rt.font.size = Pt(26); rt.bold = True
    rt.font.color.rgb = C_PRIMARY

    sub = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(0); sub.paragraph_format.space_after = Pt(12)
    if sub.runs:
        sub.runs[0].font.size = Pt(10); sub.runs[0].italic = True
        sub.runs[0].font.color.rgb = C_MUTED

    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(0); rule_p.paragraph_format.space_after = Pt(16)
    _bottom_border(rule_p, color=HEX_PRIMARY, sz="12")

    # ── Issue Details ─────────────────────────────────────────────────────────
    _h1("Issue Details")
    enabled = [fd for fd in field_defs if fd.get("enabled", True)]
    if enabled:
        rows = [(fd["label"], field_values.get(fd["key"], fd.get("default", "")))
                for fd in enabled]
        _build_table(rows, col_headers=["Field", "Value"])

    # ── Description ──────────────────────────────────────────────────────────
    _h1("Description")
    desc_p = doc.add_paragraph(field_values.get("description") or "No description provided.")
    desc_p.paragraph_format.left_indent  = Cm(0.3)
    desc_p.paragraph_format.right_indent = Cm(0.3)

    # ── GitHub PR Details ─────────────────────────────────────────────────────
    if pr_data:
        _h1("GitHub PR Details")
        pr_rows = [
            ("PR Title", pr_data.get("title")),
            ("PR State", (pr_data.get("state") or "").upper()),
            ("Author",   pr_data.get("user", {}).get("login")),
            ("Created",  (pr_data.get("created_at") or "")[:10]),
            ("Updated",  (pr_data.get("updated_at") or "")[:10]),
            ("Branch",   pr_data.get("head", {}).get("ref")),
            ("Base",     pr_data.get("base", {}).get("ref")),
        ]
        _build_table(pr_rows, col_headers=["Field", "Value"])
        if pr_data.get("body"):
            _h2("PR Body")
            doc.add_paragraph(pr_data["body"])

    # ── Code Changes ──────────────────────────────────────────────────────────
    if file_comments:
        _h1("Code Changes")
        for idx, fc in enumerate(file_comments, 1):
            fname   = fc.get("filename", "unknown")
            status  = fc.get("status", "modified")
            comment = (fc.get("word_comment") or fc.get("comment", "")).strip()
            shots   = fc.get("screenshots", [])
            adds    = fc.get("additions", 0)
            dels    = fc.get("deletions", 0)
            patch   = fc.get("patch", "")

            # File header
            fh = doc.add_paragraph()
            fh.paragraph_format.space_before = Pt(10); fh.paragraph_format.space_after = Pt(2)
            r1 = fh.add_run(f"Step {idx}  ")
            r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = C_PRIMARY
            badge_lbl = STATUS_EMOJI.get(status, "~")
            r2 = fh.add_run(f"[{badge_lbl} {status.upper()}]  ")
            r2.bold = True; r2.font.size = Pt(10)
            if status == "added":
                r2.font.color.rgb = RGBColor(0x1a, 0x7f, 0x37)
            elif status in ("removed", "deleted"):
                r2.font.color.rgb = RGBColor(0xcf, 0x22, 0x2e)
            else:
                r2.font.color.rgb = RGBColor(0x7c, 0x4a, 0x00)
            r3 = fh.add_run(fname)
            r3.bold = True; r3.font.size = Pt(10); r3.font.name = "Consolas"
            r3.font.color.rgb = C_BODY

            if adds or dels:
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent  = Cm(0.5)
                sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(4)
                ra = sp.add_run(f"+{adds} additions  ")
                ra.font.size = Pt(9); ra.font.color.rgb = RGBColor(0x1a, 0x7f, 0x37)
                rd = sp.add_run(f"−{dels} deletions")
                rd.font.size = Pt(9); rd.font.color.rgb = RGBColor(0xcf, 0x22, 0x2e)

            # Callout box — change description with steel-blue left border
            if comment:
                ch = doc.add_paragraph()
                ch.paragraph_format.left_indent  = Cm(0.5)
                ch.paragraph_format.right_indent = Cm(0.5)
                ch.paragraph_format.space_before = Pt(6); ch.paragraph_format.space_after = Pt(6)
                _shade_para(ch, HEX_LIGHT)
                _left_border(ch, color="2E75B6", sz="18")
                rl = ch.add_run("Change Description:  ")
                rl.bold = True; rl.font.size = Pt(10); rl.font.color.rgb = C_SECONDARY
                rc = ch.add_run(comment)
                rc.font.size = Pt(11); rc.font.color.rgb = C_BODY

            # Diff label
            dh = doc.add_paragraph()
            dh.paragraph_format.left_indent  = Cm(0.5)
            dh.paragraph_format.space_before = Pt(4); dh.paragraph_format.space_after = Pt(2)
            rh = dh.add_run("Code Diff:")
            rh.bold = True; rh.font.size = Pt(9); rh.font.color.rgb = C_MUTED

            _word_diff_block(doc, patch, status, adds, dels)

            for si, shot in enumerate(shots, 1):
                if isinstance(shot, dict):
                    sp_path  = shot.get("path", "")
                    shot_cmt = shot.get("comment", "")
                else:
                    sp_path  = str(shot); shot_cmt = ""
                if not sp_path or not os.path.exists(sp_path):
                    continue
                doc.add_paragraph("")
                cap_lbl = shot_cmt or f"Screenshot {si}  —  {os.path.basename(sp_path)}"
                cap = doc.add_paragraph(f"  {cap_lbl}")
                if cap.runs:
                    cap.runs[0].font.size = Pt(9); cap.runs[0].italic = True
                    cap.runs[0].font.color.rgb = C_MUTED
                try:
                    doc.add_picture(sp_path, width=Inches(5.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as ex:
                    doc.add_paragraph(f"  [Cannot embed image: {ex}]")

            # Divider rule
            div = doc.add_paragraph("")
            div.paragraph_format.space_before = Pt(10); div.paragraph_format.space_after = Pt(2)
            _bottom_border(div, color=HEX_RULE, sz="4")

    # ── Keyword Findings / Testing Scenarios ─────────────────────────────────
    if keyword_findings:
        HEX_ALERT  = "FFF3CD"
        HEX_ALERT2 = "FFECB3"
        _h1("⚠  Keyword Findings — Testing Scenarios")
        warn_p = doc.add_paragraph()
        warn_p.paragraph_format.left_indent  = Cm(0.3)
        warn_p.paragraph_format.space_before = Pt(0)
        warn_p.paragraph_format.space_after  = Pt(8)
        rw = warn_p.add_run(
            "The following SQL keywords were flagged during generation. "
            "Verify that each occurrence is intentional and confirm test coverage.")
        rw.font.size = Pt(10); rw.italic = True
        rw.font.color.rgb = RGBColor(0x7c, 0x4a, 0x00)

        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        for ci, hdr_txt in enumerate(["File", "Keywords Found", "Test Notes"]):
            cell = tbl.rows[0].cells[ci]
            cell.text = hdr_txt
            _shade_cell(cell, HEX_PRIMARY)
            for p in cell.paragraphs:
                if p.runs:
                    p.runs[0].font.color.rgb = C_WHITE
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(10)
        for ri, (fn, kws) in enumerate(keyword_findings.items()):
            row = tbl.add_row()
            bg = HEX_ALERT if ri % 2 == 0 else HEX_ALERT2
            row.cells[0].text = fn
            row.cells[1].text = ", ".join(kws)
            row.cells[2].text = ""
            for cell in row.cells:
                _shade_cell(cell, bg)
            for p in row.cells[0].paragraphs:
                if p.runs:
                    p.runs[0].font.name = "Consolas"
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.color.rgb = RGBColor(0x7c, 0x4a, 0x00)
            for p in row.cells[1].paragraphs:
                if p.runs:
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(10)
                    p.runs[0].font.color.rgb = RGBColor(0xcf, 0x22, 0x2e)

    doc_name = field_values.get("doc_name") or \
               f"Issue_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    if not doc_name.endswith(".docx"):
        doc_name += ".docx"
    pr_num   = str(pr_data.get("number", "")) if pr_data else ""
    save_dir = os.path.join(output_dir, f"PR_{pr_num}") if pr_num else output_dir
    os.makedirs(save_dir, exist_ok=True)
    out = os.path.join(save_dir, doc_name)
    doc.save(out)
    return out

# ─────────────────────────────────────────────────────────────────────────────
#  HTML preview builder  (rendered in a tk.Text widget with embedded images)
# ─────────────────────────────────────────────────────────────────────────────

def build_preview_html(field_values, field_defs, pr_data, file_comments):
    """Return an HTML string that visually matches the Word doc layout."""
    e = html.escape

    def trow(label, value, odd=True):
        bg = "#EAF0F8" if odd else "#FFFFFF"
        return (f'<tr style="background:{bg}">'
                f'<td style="font-weight:700;padding:6px 12px 6px 8px;'
                f'color:#1F3864;white-space:nowrap;border:1px solid #BDD7EE">{e(label)}</td>'
                f'<td style="padding:6px 8px;color:#404040;border:1px solid #BDD7EE">'
                f'{e(str(value) if value else "N/A")}</td></tr>')

    def theader(cols):
        cells = "".join(
            f'<th style="background:#1F3864;color:#fff;font-weight:700;'
            f'padding:6px 10px;text-align:left;font-size:11px;'
            f'letter-spacing:.04em;border:1px solid #1F3864">{e(c)}</th>'
            for c in cols)
        return f'<thead><tr>{cells}</tr></thead>'

    css = """
    *{box-sizing:border-box}
    body{background:#F4F6FA;color:#404040;
         font-family:Calibri,"Segoe UI",Arial,sans-serif;
         font-size:13px;line-height:1.5;margin:0;padding:28px 36px;max-width:920px}
    h1{color:#1F3864;font-size:17px;font-weight:700;
       border-bottom:2px solid #BDD7EE;padding-bottom:7px;
       margin:20px 0 10px 0}
    h2{color:#2E75B6;font-size:14px;font-weight:700;margin:14px 0 6px 0}
    .doc-title{color:#1F3864;font-size:26px;font-weight:700;text-align:center;
               display:block;margin-bottom:4px;line-height:1.2;font-family:Calibri,"Segoe UI",sans-serif}
    .sub{text-align:center;color:#7F7F7F;font-size:11px;
         margin-bottom:8px;font-style:italic}
    .title-rule{border:none;border-top:3px solid #1F3864;margin:10px 0 22px 0}
    table{border-collapse:collapse;width:100%;max-width:720px;margin-bottom:16px}
    td,th{vertical-align:top}
    .section{background:#FFFFFF;border:1px solid #BDD7EE;border-radius:6px;
             padding:18px 22px;margin-bottom:20px;
             box-shadow:0 1px 3px rgba(31,56,100,.06)}
    .file-card{border-left:4px solid #2E75B6;margin:14px 0;padding:12px 16px;
               background:#FAFBFC;border-radius:0 6px 6px 0;
               border:1px solid #D9E4F0;border-left-width:4px}
    .file-added{border-left-color:#1a7f37}
    .file-modified{border-left-color:#2E75B6}
    .file-removed{border-left-color:#cf222e}
    .file-renamed{border-left-color:#7c4a00}
    .step-label{font-weight:700;font-size:14px;color:#1F3864}
    .badge{display:inline-block;padding:2px 9px;border-radius:3px;
           font-size:10px;font-weight:700;margin-right:8px;letter-spacing:.03em}
    .badge-A{background:#1a7f37;color:#fff}
    .badge-M{background:#1F3864;color:#fff}
    .badge-D{background:#cf222e;color:#fff}
    .badge-R{background:#7c4a00;color:#fff}
    .badge-C{background:#2E75B6;color:#fff}
    .fname{font-family:Consolas,monospace;font-size:12px;font-weight:700;color:#404040}
    .stats{font-size:11px;color:#7F7F7F;margin-top:5px}
    .adds{color:#1a7f37;font-weight:700} .dels{color:#cf222e;font-weight:700}
    .comment-box{background:#EAF0F8;border-left:4px solid #2E75B6;
                 border-radius:0 4px 4px 0;padding:10px 14px;
                 margin-top:10px;font-size:12px;color:#404040}
    .comment-label{font-size:10px;color:#2E75B6;font-weight:700;
                   margin-bottom:4px;text-transform:uppercase;letter-spacing:.05em}
    .jira-box{background:#F0F4FB;border-left:4px solid #1F3864}
    .jira-label{color:#1F3864}
    .screenshot-wrap{margin-top:12px;text-align:center}
    .screenshot-cap{font-size:10px;color:#7F7F7F;margin-bottom:6px;font-style:italic}
    img{max-width:100%;border-radius:4px;border:1px solid #BDD7EE}
    .divider{border:none;border-top:1px solid #BDD7EE;margin:18px 0}
    .pr-body{white-space:pre-wrap;font-size:11px;color:#404040;
             background:#F9FAFB;padding:10px 14px;border-radius:4px;
             border:1px solid #BDD7EE;line-height:1.5}
    .diff-block{margin-top:10px;background:#F9FAFB;border-radius:4px;
                border:1px solid #BDD7EE}
    .diff-block summary{padding:6px 10px;cursor:pointer;color:#7F7F7F;
                        font-size:11px;user-select:none}
    .diff-content{font-family:Consolas,monospace;font-size:11px;
                  padding:8px;overflow-x:auto;max-height:320px;overflow-y:auto}
    .diff-add{background:#E6FFED;color:#1a7f37;white-space:pre}
    .diff-del{background:#FFEBE9;color:#cf222e;white-space:pre}
    .diff-hunk{background:#DDF4FF;color:#0969da;white-space:pre}
    .diff-ctx{color:#57606a;white-space:pre}
    .new-file-banner{background:#E6FFED;border-radius:4px;padding:6px 12px;
                     margin-top:8px;font-family:Consolas,monospace;
                     font-size:12px;color:#1a7f37;font-weight:700}
    """

    doc_title_val = (field_values.get("doc_title") or "").strip() or \
                    (pr_data.get("title") if pr_data else None) or "Issue Report"

    body = [f"<style>{css}</style>"]
    body.append(f'<div class="doc-title">{e(doc_title_val)}</div>')
    body.append(f'<p class="sub">Generated: {datetime.now().strftime("%Y-%m-%d  %H:%M:%S")}</p>')
    body.append('<hr class="title-rule">')

    # Issue Details
    body.append('<div class="section"><h1>Issue Details</h1>')
    body.append(f'<table>{theader(["Field","Value"])}<tbody>')
    enabled = [fd for fd in field_defs if fd.get("enabled", True)]
    for ri, fd in enumerate(enabled):
        val = field_values.get(fd["key"], fd.get("default", ""))
        body.append(trow(fd["label"], val, odd=ri % 2 == 0))
    body.append('</tbody></table></div>')

    # Description
    desc = field_values.get("description") or "No description provided."
    body.append(f'<div class="section"><h1>Description</h1>'
                f'<p style="white-space:pre-wrap;margin:0">{e(desc)}</p></div>')

    # PR Details
    if pr_data:
        pr_fields = [
            ("PR Title", pr_data.get("title")),
            ("PR State", (pr_data.get("state") or "").upper()),
            ("Author",   pr_data.get("user", {}).get("login")),
            ("Created",  (pr_data.get("created_at") or "")[:10]),
            ("Updated",  (pr_data.get("updated_at") or "")[:10]),
            ("Branch",   pr_data.get("head", {}).get("ref")),
            ("Base",     pr_data.get("base", {}).get("ref")),
        ]
        body.append('<div class="section"><h1>GitHub PR Details</h1>')
        body.append(f'<table>{theader(["Field","Value"])}<tbody>')
        for ri, (label, value) in enumerate(pr_fields):
            body.append(trow(label, value, odd=ri % 2 == 0))
        body.append('</tbody></table>')
        if pr_data.get("body"):
            body.append(f'<h2>PR Body</h2>'
                        f'<div class="pr-body">{e(pr_data["body"])}</div>')
        body.append('</div>')

    # Code Changes
    if file_comments:
        body.append('<div class="section"><h1>Code Changes</h1>')
        for idx, fc in enumerate(file_comments, 1):
            fname        = fc.get("filename", "unknown")
            status       = fc.get("status", "modified")
            word_comment = (fc.get("word_comment") or fc.get("comment", "")).strip()
            jira_comment = fc.get("jira_comment", "").strip()
            shots        = fc.get("screenshots", [])
            adds         = fc.get("additions", 0)
            dels         = fc.get("deletions", 0)
            patch        = fc.get("patch", "")
            badge        = STATUS_EMOJI.get(status, "~")

            body.append(f'<div class="file-card file-{status}">')
            body.append(f'<span class="step-label">Step {idx}</span>&nbsp;&nbsp;'
                        f'<span class="badge badge-{badge}">{badge} {status.upper()}</span>'
                        f'<span class="fname">{e(fname)}</span>')
            body.append(f'<div class="stats">'
                        f'<span class="adds">+{adds}</span>&nbsp;&nbsp;'
                        f'<span class="dels">−{dels}</span></div>')

            if status == "added":
                body.append(f'<div class="new-file-banner">'
                            f'+ NEW FILE &nbsp;—&nbsp; {adds} line{"s" if adds != 1 else ""} added'
                            f'</div>')

            if word_comment:
                body.append(f'<div class="comment-box">'
                            f'<div class="comment-label">Change Description</div>'
                            f'{e(word_comment)}</div>')
            if jira_comment and jira_comment != word_comment:
                body.append(f'<div class="comment-box jira-box">'
                            f'<div class="comment-label jira-label">Jira Comment</div>'
                            f'{e(jira_comment)}</div>')

            if patch and status != "added":
                patch_lines = patch.splitlines()
                shown = patch_lines[:120]
                diff_rows = []
                for dl in shown:
                    if dl.startswith('+') and not dl.startswith('+++'):
                        diff_rows.append(
                            f'<div class="diff-add">'
                            f'<span style="opacity:.5;margin-right:6px">+</span>{e(dl[1:])}</div>')
                    elif dl.startswith('-') and not dl.startswith('---'):
                        diff_rows.append(
                            f'<div class="diff-del">'
                            f'<span style="opacity:.5;margin-right:6px">−</span>{e(dl[1:])}</div>')
                    elif dl.startswith('@'):
                        diff_rows.append(f'<div class="diff-hunk">{e(dl)}</div>')
                    else:
                        diff_rows.append(
                            f'<div class="diff-ctx">'
                            f'<span style="opacity:.3;margin-right:6px">&nbsp;</span>'
                            f'{e(dl[1:] if dl and dl[0]==" " else dl)}</div>')
                if len(patch_lines) > 120:
                    diff_rows.append(
                        f'<div class="diff-ctx" style="color:#7c4a00">'
                        f'⋯  {len(patch_lines)-120} more lines not shown</div>')
                body.append(
                    f'<details class="diff-block" open>'
                    f'<summary>Code Diff &nbsp;'
                    f'<span class="adds">+{adds}</span>&nbsp;'
                    f'<span class="dels">−{dels}</span></summary>'
                    f'<div class="diff-content">{"".join(diff_rows)}</div>'
                    f'</details>')

            for si, shot in enumerate(shots, 1):
                if isinstance(shot, dict):
                    sp_path  = shot.get("path", "")
                    shot_cmt = shot.get("comment", "")
                else:
                    sp_path  = str(shot); shot_cmt = ""
                if not sp_path or not os.path.exists(sp_path):
                    continue
                import base64 as b64
                try:
                    ext  = os.path.splitext(sp_path)[1].lower().lstrip(".")
                    mime = {"jpg":"jpeg","jpeg":"jpeg","png":"png",
                            "gif":"gif","webp":"webp","bmp":"bmp"}.get(ext, "png")
                    with open(sp_path, "rb") as fh:
                        data = b64.b64encode(fh.read()).decode()
                    cap_html = (f'<div class="screenshot-cap">{e(shot_cmt)}</div>'
                                if shot_cmt else
                                f'<div class="screenshot-cap">'
                                f'Screenshot {si} — {e(os.path.basename(sp_path))}</div>')
                    body.append(f'<div class="screenshot-wrap">'
                                f'{cap_html}'
                                f'<img src="data:image/{mime};base64,{data}" />'
                                f'</div>')
                except Exception:
                    body.append(f'<div class="stats">[Cannot load screenshot {si}]</div>')

            body.append('</div><hr class="divider">')
        body.append('</div>')

    return "<html><body>" + "".join(body) + "</body></html>"


# ─────────────────────────────────────────────────────────────────────────────
#  DocPreviewWindow
# ─────────────────────────────────────────────────────────────────────────────

class DocPreviewWindow(tk.Toplevel):
    """Preview window with left comment panel (per-file) and right HTML preview."""

    def __init__(self, parent, html_content, doc_path=None,
                 pr_files=None, file_comments=None, on_comments_saved=None,
                 build_html_func=None):
        super().__init__(parent)
        self.title(f"{APP_NAME}  —  Code Changes")
        self.geometry("1300x820")
        self.minsize(900, 560)
        self.configure(bg=C["bg"])
        self._doc_path         = doc_path
        self._html_content     = html_content
        self._pr_files         = pr_files or []
        self._on_save          = on_comments_saved
        self._build_html_func  = build_html_func
        self._html_frame       = None
        self._comment_rows     = []   # {filename, word_var, jira_var, status, ...}

        # pre-index existing comments by filename
        self._existing = {fc["filename"]: fc for fc in (file_comments or [])}

        self._build()

    # ── build ─────────────────────────────────────────────────────────────────

    def _build(self):
        # ── top accent stripe ─────────────────────────────────────────────────
        tk.Frame(self, bg=C["orange"], height=3).pack(fill="x")

        # ── toolbar ───────────────────────────────────────────────────────────
        self._bar = tk.Frame(self, bg=C["surface"])
        self._bar.pack(fill="x")
        title_f = tk.Frame(self._bar, bg=C["surface"])
        title_f.pack(side="left", padx=(16, 0), pady=10)
        tk.Label(title_f, text="◈", bg=C["surface"], fg=C["orange"],
                 font=("Segoe UI", 13)).pack(side="left", padx=(0, 8))
        tk.Label(title_f, text="Code Changes",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 12, "bold")).pack(side="left")
        tk.Label(title_f, text="  —  Add step comments for Word doc",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 9)).pack(side="left", pady=(3, 0))

        # Open in Word button — enabled only after doc is generated
        doc_exists = bool(self._doc_path and os.path.exists(self._doc_path))
        self._open_word_btn = tk.Button(
            self._bar, text="⊞  Open in Word",
            bg=C["green"] if doc_exists else C["card2"],
            fg="#ffffff" if doc_exists else C["muted"],
            relief="flat", font=("Segoe UI", 9, "bold"), padx=12, pady=6,
            activebackground="#2ea043", cursor="hand2" if doc_exists else "arrow",
            state="normal" if doc_exists else "disabled",
            command=self._open_word)
        self._open_word_btn.pack(side="right", padx=8, pady=8)

        if self._pr_files and self._on_save:
            self._status_lbl = tk.Label(self._bar, text="",
                                        bg=C["surface"], fg=C["green"],
                                        font=("Segoe UI", 9))
            self._status_lbl.pack(side="right", padx=8)
            tk.Button(self._bar, text="✦  Generate Doc",
                      bg=C["green"], fg="#ffffff", relief="flat",
                      font=("Segoe UI", 9, "bold"), padx=12, pady=6,
                      activebackground="#2ea043", cursor="hand2",
                      command=self._generate_doc).pack(side="right", padx=4, pady=8)

        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        # ── body: full-screen comment panel ───────────────────────────────────
        body = tk.Frame(self, bg=C["bg"])
        body.pack(fill="both", expand=True)

        if self._pr_files:
            self._build_comment_panel(body)

    def _build_comment_panel(self, parent):
        panel = tk.Frame(parent, bg=C["bg"])
        panel.pack(fill="both", expand=True)

        # header
        ph = tk.Frame(panel, bg=C["card2"])
        ph.pack(fill="x")
        tk.Frame(ph, bg=C["yellow"], width=4).pack(side="left", fill="y")
        ph_inner = tk.Frame(ph, bg=C["card2"])
        ph_inner.pack(side="left", fill="x", expand=True, padx=(10, 8), pady=8)
        tk.Label(ph_inner, text=f"Code Changes",
                 bg=C["card2"], fg=C["yellow"],
                 font=("Segoe UI", 10, "bold")).pack(side="left")
        tk.Label(ph_inner, text=f"  {len(self._pr_files)} files changed  •  Add step comments",
                 bg=C["card2"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(side="left", pady=(2, 0))

        tk.Frame(panel, bg=C["border"], height=1).pack(fill="x")

        # scrollable file cards
        wrap = tk.Frame(panel, bg=C["bg"]); wrap.pack(fill="both", expand=True)
        canvas = tk.Canvas(wrap, bg=C["bg"], highlightthickness=0)
        vsb = ttk.Scrollbar(wrap, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)
        lf = tk.Frame(canvas, bg=C["bg"])
        wid = canvas.create_window((0, 0), window=lf, anchor="nw")
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(wid, width=e.width))
        lf.bind("<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        def _mw(e): canvas.yview_scroll(int(-1*(e.delta/120)), "units")
        canvas.bind("<MouseWheel>", _mw)

        sorted_files = _topo_sort_files(self._pr_files)

        for idx, f in enumerate(sorted_files, 1):
            fname  = f.get("filename", "")
            status = f.get("status", "modified")
            adds   = f.get("additions", 0)
            dels   = f.get("deletions", 0)
            rbg    = STATUS_ROW.get(status, C["card"])
            bbg    = STATUS_BADGE.get(status, C["surface"])
            em     = STATUS_EMOJI.get(status, "~")
            ex     = self._existing.get(fname, {})

            card = tk.Frame(lf, bg=rbg,
                            highlightbackground=C["border"], highlightthickness=1)
            card.pack(fill="x", padx=6, pady=3)
            card.bind("<MouseWheel>", _mw)

            # file header
            fhdr = tk.Frame(card, bg=rbg); fhdr.pack(fill="x", padx=6, pady=(6, 2))
            fhdr.bind("<MouseWheel>", _mw)
            tk.Label(fhdr, text=f"Step {idx}",
                     bg=C["accent"], fg="#fff",
                     font=("Segoe UI", 8, "bold"), padx=4, pady=1
                     ).pack(side="left", padx=(0, 4))
            tk.Label(fhdr, text=f" {em} {status.upper()} ",
                     bg=bbg, fg="#fff",
                     font=("Segoe UI", 8, "bold"), padx=3, pady=1
                     ).pack(side="left", padx=(0, 6))
            sfg = C["green"] if adds > dels else C["red"] if dels > adds else C["muted"]
            tk.Label(fhdr, text=f"+{adds} −{dels}",
                     bg=rbg, fg=sfg,
                     font=("Segoe UI", 8, "bold")).pack(side="right")

            # filename + inline Show Code Change button
            parts = fname.split("/")
            pre   = "/".join(parts[:-2]) + "/" if len(parts) > 2 else ""
            tail  = "/".join(parts[-2:])
            nf    = tk.Frame(card, bg=rbg); nf.pack(fill="x", padx=6, pady=(0, 4))
            nf.bind("<MouseWheel>", _mw)
            if pre:
                tk.Label(nf, text=pre, bg=rbg, fg=C["muted"],
                         font=("Consolas", 8)).pack(side="left")
            tk.Label(nf, text=tail, bg=rbg, fg=C["text"],
                     font=("Consolas", 9, "bold")).pack(side="left")
            patch_str = f.get("patch", "")
            if patch_str:
                n_lines = len(patch_str.splitlines())
                tog_btn = tk.Button(
                    nf,
                    text=f"⊞ Code Change ({n_lines})",
                    bg=C["accent"], fg="#ffffff", relief="flat",
                    font=("Segoe UI", 8, "bold"),
                    padx=6, pady=1, cursor="hand2",
                    activebackground=C["hover"], activeforeground="#ffffff",
                    command=lambda fn=fname, p=patch_str, st=status, a=adds, d=dels:
                        DiffViewPopup(self, fn, p, st, a, d))
                tog_btn.pack(side="left", padx=(8, 0))
                tog_btn.bind("<MouseWheel>", _mw)

            # Change Description
            tk.Label(card, text="Change Description:",
                     bg=rbg, fg=C["yellow"],
                     font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=6)
            word_var = tk.StringVar(value=ex.get("word_comment", ex.get("comment", "")))
            word_ent = scrolledtext.ScrolledText(
                card, height=3, bg=C["inp"], fg=C["text"],
                font=("Segoe UI", 10), relief="flat",
                insertbackground=C["text"], wrap="word")
            word_ent.insert("1.0", word_var.get())
            word_ent.pack(fill="x", padx=6, pady=(2, 8))
            word_ent.bind("<MouseWheel>", _mw)

            self._comment_rows.append(dict(
                filename=fname, status=status,
                additions=adds, deletions=dels,
                patch=f.get("patch", ""),
                screenshots=ex.get("screenshots", []),
                word_ent=word_ent,
            ))

    def _build_preview(self, parent):
        pf = tk.Frame(parent, bg=C["bg"])
        pf.pack(side="left", fill="both", expand=True)
        try:
            from tkinterweb import HtmlFrame
            frame = HtmlFrame(pf, horizontal_scrollbar="auto")
            frame.pack(fill="both", expand=True)
            frame.load_html(self._html_content)
            self._html_frame = frame
        except ImportError:
            self._fallback_viewer(pf, self._html_content)

    def _fallback_viewer(self, parent, html_content):
        text = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL)
        text = re.sub(r'<[^>]+>', '', text)
        text = text.replace("&amp;", "&").replace("&lt;", "<").replace(
            "&gt;", ">").replace("&#39;", "'").replace("&quot;", '"')
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        text  = "\n".join(lines)
        tk.Label(parent,
                 text="Install  pip install tkinterweb  for full HTML preview  |  "
                      "Showing plain-text fallback",
                 bg=C["yellow"], fg="#0d1117",
                 font=("Segoe UI", 9), pady=4).pack(fill="x")
        st = scrolledtext.ScrolledText(parent, bg=C["card"], fg=C["text"],
                                       font=("Segoe UI", 10), relief="flat",
                                       wrap="word")
        st.pack(fill="both", expand=True, padx=4, pady=4)
        st.insert("1.0", text)
        st.config(state="disabled")

    # ── actions ───────────────────────────────────────────────────────────────

    def _generate_doc(self):
        result = []
        for r in self._comment_rows:
            wc = r["word_ent"].get("1.0", "end-1c").strip()
            result.append(dict(
                filename=r["filename"], status=r["status"],
                additions=r["additions"], deletions=r["deletions"],
                patch=r["patch"], screenshots=r["screenshots"],
                comment=wc, word_comment=wc, jira_comment="",
                include=True,
            ))
        n_c = sum(1 for r in result if r["comment"])
        if hasattr(self, "_status_lbl"):
            self._status_lbl.config(text=f"Generating…  {len(result)} files, {n_c} comments")
        self.update_idletasks()
        if self._on_save:
            self._on_save(result)

    def _open_word(self):
        if self._doc_path and os.path.exists(self._doc_path):
            import subprocess, sys
            if sys.platform == "win32":
                os.startfile(self._doc_path)
            elif sys.platform == "darwin":
                subprocess.run(["open", self._doc_path])
            else:
                subprocess.run(["xdg-open", self._doc_path])

    def _enable_open_word(self, path):
        self._doc_path = path
        self._open_word_btn.config(
            bg=C["green"], fg="#0d1117",
            activebackground="#2ea043", cursor="hand2",
            state="normal")


# ─────────────────────────────────────────────────────────────────────────────
#  FieldManagerDialog  – add / remove / reorder / set defaults for Jira fields
# ─────────────────────────────────────────────────────────────────────────────

class FieldManagerDialog(tk.Toplevel):
    """
    Shows all field definitions.  User can:
      • toggle enabled
      • change default value
      • move up / down
      • delete custom fields
      • add new custom fields
    Returns via .result  (list of field dicts) after apply.
    """

    # Column spec: (header_label, min_pixel_width, tooltip)
    _COLS = [
        ("",            28,  ""),
        ("On",          34,  "Enable / disable field"),
        ("Field Label", 155, "Name shown in UI and Word doc"),
        ("Type",        68,  "text / dropdown / number / date"),
        ("Default",     175, "Pre-filled value when creating a ticket"),
        ("Choices",     64,  "Edit the dropdown options list"),
        ("Jira API",    56,  "Map this field directly to a Jira API field"),
        ("Show Label",  72,  "Prefix 'FieldName: value' in Jira description"),
        ("",            92,  ""),
    ]

    def __init__(self, parent, fields):
        super().__init__(parent)
        self.title(f"{APP_NAME}  —  Manage Fields")
        self.geometry("980x660")
        self.minsize(820, 480)
        self.configure(bg=C["bg"])
        self.grab_set()
        self.transient(parent)

        self._fields     = copy.deepcopy(fields)
        self.result      = None
        self._row_frames = []

        self._build()

    # ── build ─────────────────────────────────────────────────────────────────

    def _build(self):
        # ── top accent stripe ─────────────────────────────────────────────────
        tk.Frame(self, bg=C["purple"], height=3).pack(fill="x")

        # ── title bar ────────────────────────────────────────────────────────
        hdr = tk.Frame(self, bg=C["surface"])
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=C["purple"], width=4).pack(side="left", fill="y")
        ti = tk.Frame(hdr, bg=C["surface"]); ti.pack(side="left", padx=14, pady=12)
        hdr_row = tk.Frame(ti, bg=C["surface"])
        hdr_row.pack(anchor="w")
        tk.Label(hdr_row, text="⊞", bg=C["surface"], fg=C["purple"],
                 font=("Segoe UI", 14)).pack(side="left", padx=(0, 8))
        tk.Label(hdr_row, text="Manage Fields",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 13, "bold")).pack(side="left", anchor="w")
        tk.Label(ti, text="Enable / disable fields  •  set defaults  •  control Jira mapping",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(2, 0))
        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        # ── column header (grid-based, matches _render_rows columns exactly) ──
        col_hdr = tk.Frame(self, bg=C["card"])
        col_hdr.pack(fill="x")
        # left indent strip to visually align with row content
        tk.Frame(col_hdr, bg=C["card"], width=2).pack(side="left", fill="y")
        hdr_inner = tk.Frame(col_hdr, bg=C["card"])
        hdr_inner.pack(side="left", fill="x", expand=True)
        for c, (txt, minw, tip) in enumerate(self._COLS):
            hdr_inner.columnconfigure(c, minsize=minw, weight=1 if c == 2 else 0)
        for c, (txt, minw, tip) in enumerate(self._COLS):
            cell = tk.Label(hdr_inner, text=txt, bg=C["card"], fg=C["accent"],
                            font=("Segoe UI", 8, "bold"), anchor="w", padx=6)
            cell.grid(row=0, column=c, sticky="ew", ipady=6)
            if tip:
                Tooltip(cell, tip)
        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        # ── scrollable list ───────────────────────────────────────────────────
        outer = tk.Frame(self, bg=C["bg"])
        outer.pack(fill="both", expand=True)
        self._canvas = tk.Canvas(outer, bg=C["bg"], highlightthickness=0)
        vbar = ttk.Scrollbar(outer, orient="vertical", command=self._canvas.yview)
        self._canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side="right", fill="y")
        self._canvas.pack(side="left", fill="both", expand=True)
        self._list_frame = tk.Frame(self._canvas, bg=C["bg"])
        self._win = self._canvas.create_window((0, 0), window=self._list_frame, anchor="nw")
        self._canvas.bind("<Configure>",
                          lambda e: self._canvas.itemconfig(self._win, width=e.width))
        self._list_frame.bind("<Configure>",
                              lambda e: self._canvas.configure(
                                  scrollregion=self._canvas.bbox("all")))
        self._canvas.bind("<MouseWheel>",
                          lambda e: self._canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        self._render_rows()

        # ── bottom action bar ─────────────────────────────────────────────────
        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")
        bot = tk.Frame(self, bg=C["surface"])
        bot.pack(fill="x", side="bottom", pady=0)
        inner_bot = tk.Frame(bot, bg=C["surface"])
        inner_bot.pack(fill="x", padx=14, pady=10)

        tk.Button(inner_bot, text="⊕  Add Custom Field",
                  bg=C["accent_dim"], fg=C["accent"], relief="flat",
                  font=("Segoe UI", 9, "bold"), padx=12, pady=6,
                  activebackground=C["border"], cursor="hand2",
                  command=self._add_field).pack(side="left")

        tk.Button(inner_bot, text="✦  Apply & Close",
                  bg=C["green"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=16, pady=6,
                  activebackground="#2ea043", cursor="hand2",
                  command=self._apply).pack(side="right")
        tk.Button(inner_bot, text="✕  Cancel",
                  bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground="#c0392b", cursor="hand2",
                  command=self.destroy).pack(side="right", padx=(0, 8))

    def _render_rows(self):
        for w in self._list_frame.winfo_children():
            w.destroy()
        self._row_frames = []

        # configure grid columns to match header exactly
        for c, (_, minw, _) in enumerate(self._COLS):
            self._list_frame.columnconfigure(c, minsize=minw, weight=1 if c == 2 else 0)

        type_clrs = {"text": C["muted"], "dropdown": C["accent"],
                     "number": C["yellow"], "date": C["green"]}

        for i, fd in enumerate(self._fields):
            r      = i * 2          # data on even rows, separator on odd
            bg     = C["card"] if i % 2 == 0 else C["surface"]
            enabled = fd.get("enabled", True)
            fg_txt = C["text"] if enabled else C["muted"]
            ftype  = fd.get("type", "text")

            def _cell(text="", col=0, fg=None, font=("Segoe UI", 9), anchor="w"):
                tk.Label(self._list_frame, text=text, bg=bg, fg=fg or fg_txt,
                         font=font, anchor=anchor, padx=6
                         ).grid(row=r, column=col, sticky="nsew", ipady=7)

            # col 0 — index
            _cell(str(i + 1), col=0, fg=C["muted"], font=("Segoe UI", 8))

            # col 1 — enabled checkbox
            en_var = tk.BooleanVar(value=enabled)
            tk.Checkbutton(self._list_frame, variable=en_var, bg=bg, fg=C["text"],
                           selectcolor=C["cb_sel"], activebackground=bg,
                           activeforeground=C["text"],
                           relief="flat", bd=0, cursor="hand2"
                           ).grid(row=r, column=1, sticky="ns", pady=4)

            # col 2 — field label
            lbl_var = tk.StringVar(value=fd["label"])
            if fd.get("custom"):
                e = ttk.Entry(self._list_frame, textvariable=lbl_var)
                e.grid(row=r, column=2, sticky="ew", padx=(6, 4), pady=5)
            else:
                hint = FIELD_HINTS.get(fd.get("key", ""), "")
                lw = tk.Label(self._list_frame, textvariable=lbl_var, bg=bg, fg=fg_txt,
                              anchor="w", padx=6,
                              font=("Segoe UI", 9, "bold" if fd.get("required") else "normal"))
                lw.grid(row=r, column=2, sticky="nsew", ipady=7)
                if hint:
                    Tooltip(lw, hint)

            # col 3 — type
            _cell(ftype, col=3, fg=type_clrs.get(ftype, C["muted"]),
                  font=("Segoe UI", 8, "italic"))

            # col 4 — default value
            def_var    = tk.StringVar(value=fd.get("default", ""))
            choices_list = fd.get("choices", [])
            choices_cb   = None
            if choices_list:
                choices_cb = ttk.Combobox(self._list_frame, textvariable=def_var,
                                          values=choices_list, state="normal")
                choices_cb.grid(row=r, column=4, sticky="ew", padx=(6, 4), pady=5)
            else:
                ttk.Entry(self._list_frame, textvariable=def_var).grid(
                    row=r, column=4, sticky="ew", padx=(6, 4), pady=5)

            # col 5 — edit choices (dropdown only)
            if ftype == "dropdown":
                def _open_edit(fd=fd, cb=choices_cb, dv=def_var):
                    self._edit_choices_dialog(fd, cb, dv)
                tk.Button(self._list_frame, text="✎", bg=bg, fg=C["accent"],
                          relief="flat", font=("Segoe UI", 10), padx=6, pady=0,
                          activebackground=C["border"], cursor="hand2",
                          command=_open_edit).grid(row=r, column=5, sticky="ns", pady=5)
            else:
                tk.Label(self._list_frame, bg=bg).grid(row=r, column=5, sticky="nsew")

            # col 6 — Jira API checkbox
            jira_var = tk.BooleanVar(value=bool(fd.get("jira_key")))
            tk.Checkbutton(self._list_frame, variable=jira_var, bg=bg, fg=C["text"],
                           selectcolor=C["cb_sel"], activebackground=bg,
                           activeforeground=C["text"],
                           relief="flat", bd=0, cursor="hand2"
                           ).grid(row=r, column=6, sticky="ns", pady=4)

            # col 7 — show label checkbox
            show_lbl_var = tk.BooleanVar(value=fd.get("show_label_in_jira", True))
            tk.Checkbutton(self._list_frame, variable=show_lbl_var, bg=bg, fg=C["text"],
                           selectcolor=C["cb_sel"], activebackground=bg,
                           activeforeground=C["text"],
                           relief="flat", bd=0, cursor="hand2"
                           ).grid(row=r, column=7, sticky="ns", pady=4)

            # col 8 — ▲▼✕ actions
            af = tk.Frame(self._list_frame, bg=bg)
            af.grid(row=r, column=8, sticky="ns", padx=4, pady=4)

            def _up(idx=i):
                if idx > 0:
                    self._fields[idx], self._fields[idx-1] = \
                        self._fields[idx-1], self._fields[idx]
                    self._render_rows()
            def _down(idx=i):
                if idx < len(self._fields) - 1:
                    self._fields[idx], self._fields[idx+1] = \
                        self._fields[idx+1], self._fields[idx]
                    self._render_rows()
            def _del(idx=i, fd=fd):
                if not fd.get("custom") and fd.get("required"):
                    messagebox.showwarning("Cannot delete",
                        f'"{fd["label"]}" is required — disable via checkbox instead.',
                        parent=self)
                    return
                if messagebox.askyesno("Remove field",
                        f'Remove "{fd["label"]}"?', parent=self):
                    del self._fields[idx]
                    self._render_rows()

            for txt, cmd, fg in [("▲", _up, C["muted"]),
                                  ("▼", _down, C["muted"]),
                                  ("✕", _del, C["red"])]:
                tk.Button(af, text=txt, bg=bg, fg=fg,
                          relief="flat", font=("Segoe UI", 9), padx=4,
                          activebackground=C["border"],
                          command=cmd, cursor="hand2").pack(side="left")

            # thin separator row
            tk.Frame(self._list_frame, bg=C["border"], height=1).grid(
                row=r + 1, column=0, columnspan=len(self._COLS), sticky="ew")

            self._row_frames.append(dict(
                fd=fd, en_var=en_var, lbl_var=lbl_var,
                def_var=def_var, jira_var=jira_var,
                show_lbl_var=show_lbl_var))

    def _add_field(self):
        """Dialog to create a new custom field."""
        dlg = tk.Toplevel(self)
        dlg.title("Add Custom Field")
        dlg.geometry("540x440")
        dlg.configure(bg=C["bg"])
        dlg.grab_set()

        tk.Frame(dlg, bg=C["accent"], height=3).pack(fill="x")
        hdr = tk.Frame(dlg, bg=C["surface"])
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=C["accent"], width=4).pack(side="left", fill="y")
        hdr_inner = tk.Frame(hdr, bg=C["surface"])
        hdr_inner.pack(side="left", padx=12, pady=10)
        tk.Label(hdr_inner, text="⊕  New Custom Field",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 12, "bold")).pack(anchor="w")
        tk.Label(hdr_inner, text="Define a new field for Jira tickets and Word documents",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(2, 0))
        tk.Frame(dlg, bg=C["border"], height=1).pack(fill="x")

        def row(parent, label, idx=None):
            row_bg = C["card"] if (idx is not None and idx % 2 == 0) else C["bg"]
            f = tk.Frame(parent, bg=row_bg)
            f.pack(fill="x", padx=14, pady=0)
            inner = tk.Frame(f, bg=row_bg)
            inner.pack(fill="x", padx=6, pady=5)
            tk.Label(inner, text=label, bg=row_bg, fg=C["muted"],
                     width=18, anchor="w",
                     font=("Segoe UI", 9)).pack(side="left")
            return inner

        # label
        lbl_var = tk.StringVar()
        f1 = row(dlg, "Field Label *", idx=0)
        ttk.Entry(f1, textvariable=lbl_var, width=28).pack(side="left", fill="x", expand=True)

        # type
        type_var = tk.StringVar(value="text")
        f2 = row(dlg, "Field Type", idx=1)
        ttk.Combobox(f2, textvariable=type_var,
                     values=["text", "dropdown", "number", "date"],
                     state="readonly", width=14).pack(side="left")

        # default
        def_var = tk.StringVar()
        f3 = row(dlg, "Default Value", idx=2)
        ttk.Entry(f3, textvariable=def_var, width=28).pack(side="left", fill="x", expand=True)

        # dropdown choices (shown when type=dropdown)
        choices_var = tk.StringVar()
        f4 = row(dlg, "Choices (csv)", idx=3)
        choices_entry = ttk.Entry(f4, textvariable=choices_var, width=28)

        def _on_type(*_):
            if type_var.get() == "dropdown":
                choices_entry.pack(side="left")
            else:
                choices_entry.pack_forget()
        type_var.trace_add("write", _on_type)
        _on_type()

        # jira key
        jira_key_var = tk.StringVar()
        f5 = row(dlg, "Jira Field Key", idx=4)
        ttk.Entry(f5, textvariable=jira_key_var, width=28).pack(side="left", fill="x", expand=True)

        # show label in Jira
        show_lbl_new = tk.BooleanVar(value=True)
        f6 = row(dlg, "Show Label in Jira", idx=5)
        tk.Checkbutton(f6, variable=show_lbl_new, bg=C["bg"] if 5 % 2 else C["card"],
                       fg=C["text"], selectcolor=C["cb_sel"],
                       activebackground=C["bg"], activeforeground=C["text"],
                       relief="flat").pack(side="left")
        tk.Label(f6, text="show 'Label: value' in Jira description",
                 bg=C["bg"] if 5 % 2 else C["card"],
                 fg=C["muted"], font=("Segoe UI", 8)).pack(side="left", padx=6)

        def _create():
            label = lbl_var.get().strip()
            if not label:
                messagebox.showwarning("Required", "Field Label is required", parent=dlg)
                return
            key     = label.lower().replace(" ", "_")
            choices = [c.strip() for c in choices_var.get().split(",") if c.strip()] \
                      if type_var.get() == "dropdown" else []
            new_fd  = dict(
                label             = label,
                key               = key,
                type              = type_var.get(),
                choices           = choices,
                default           = def_var.get().strip(),
                jira_key          = jira_key_var.get().strip() or None,
                show_label_in_jira= show_lbl_new.get(),
                required          = False,
                enabled           = True,
                custom            = True,
            )
            self._fields.append(new_fd)
            self._render_rows()
            dlg.destroy()

        tk.Frame(dlg, bg=C["border"], height=1).pack(fill="x")
        brow = tk.Frame(dlg, bg=C["surface"])
        brow.pack(fill="x", padx=16, pady=12)
        tk.Button(brow, text="⊕  Add Field",
                  bg=C["green"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=14, pady=6,
                  cursor="hand2", activebackground="#2ea043",
                  command=_create).pack(side="right")
        tk.Button(brow, text="✕  Cancel",
                  bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground="#c0392b", cursor="hand2",
                  command=dlg.destroy).pack(side="right", padx=(0, 8))

    def _edit_choices_dialog(self, fd, combobox, def_var):
        dlg = tk.Toplevel(self)
        dlg.title(f"Edit choices — {fd['label']}")
        dlg.geometry("400x380")
        dlg.configure(bg=C["bg"])
        dlg.grab_set()
        dlg.transient(self)

        tk.Label(dlg, text=f"Choices for \"{fd['label']}\"",
                 bg=C["bg"], fg=C["accent"],
                 font=("Segoe UI", 11, "bold")).pack(anchor="w", padx=16, pady=(14, 2))
        tk.Label(dlg, text="One value per line  •  Order here = order in dropdown",
                 bg=C["bg"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", padx=16, pady=(0, 8))

        from tkinter import scrolledtext as st_mod
        txt = st_mod.ScrolledText(dlg, height=12,
                                  bg=C["inp"], fg=C["text"],
                                  insertbackground=C["text"],
                                  relief="flat", font=("Segoe UI", 10),
                                  highlightthickness=1,
                                  highlightbackground=C["border"])
        txt.pack(fill="both", expand=True, padx=16, pady=(0, 10))
        txt.insert("1.0", "\n".join(fd.get("choices", [])))

        add_var = tk.StringVar()
        add_row = tk.Frame(dlg, bg=C["bg"]); add_row.pack(fill="x", padx=16, pady=(0, 6))
        ttk.Entry(add_row, textvariable=add_var).pack(side="left", fill="x", expand=True)
        def _add_value():
            v = add_var.get().strip()
            if v:
                txt.insert("end", ("\n" if txt.get("1.0", "end-1c") else "") + v)
                add_var.set("")
        tk.Button(add_row, text="+ Add", bg=C["accent"], fg="#fff",
                  relief="flat", font=("Segoe UI", 9, "bold"), padx=10, pady=3,
                  cursor="hand2", command=_add_value).pack(side="left", padx=(6, 0))

        def _save():
            lines = [l.strip() for l in txt.get("1.0", "end-1c").splitlines() if l.strip()]
            fd["choices"] = lines
            if combobox:
                combobox.config(values=lines)
                if def_var.get() not in lines and lines:
                    def_var.set(lines[0])
            dlg.destroy()

        br = tk.Frame(dlg, bg=C["bg"]); br.pack(fill="x", padx=16, pady=(0, 14))
        tk.Button(br, text="✦  Save", bg=C["green"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=14, pady=6,
                  cursor="hand2", command=_save).pack(side="right")
        tk.Button(br, text="✕  Cancel", bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground="#c0392b", cursor="hand2",
                  command=dlg.destroy).pack(side="right", padx=(0, 8))

    def _apply(self):
        # read current widget values back into field dicts
        for meta in self._row_frames:
            fd = meta["fd"]
            fd["enabled"] = meta["en_var"].get()
            fd["default"] = meta["def_var"].get()
            fd["show_label_in_jira"] = meta["show_lbl_var"].get()
            if fd.get("custom"):
                fd["label"] = meta["lbl_var"].get().strip()
            if not meta["jira_var"].get():
                fd["jira_key"] = None
        self.result = self._fields
        self.destroy()


# ─────────────────────────────────────────────────────────────────────────────
#  FieldToggleDialog  – simple enable/disable list for a filtered field set
# ─────────────────────────────────────────────────────────────────────────────

class FieldToggleDialog(tk.Toplevel):
    """
    Popup showing a filtered list of fields with enable/disable checkboxes.
    filter_fn: callable(fd) → bool to select which fields to display.
    Supports adding custom Jira API fields.
    Returns updated fields list via .result after apply.
    """

    def __init__(self, parent, fields, title, subtitle, filter_fn, accent_color=None):
        super().__init__(parent)
        self.title(f"{APP_NAME}  —  {title}")
        self.geometry("560x600")
        self.minsize(460, 400)
        self.configure(bg=C["bg"])
        self.grab_set()
        self.transient(parent)

        self._all_fields = fields
        self._filter_fn  = filter_fn
        self._accent     = accent_color or C["accent"]
        self.result      = None
        self._vars       = {}
        self._canvas     = None
        self._list_frame = None

        self._build(title, subtitle)

    def _build(self, title, subtitle):
        tk.Frame(self, bg=self._accent, height=3).pack(fill="x")
        hdr = tk.Frame(self, bg=C["surface"])
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=self._accent, width=4).pack(side="left", fill="y")
        ti = tk.Frame(hdr, bg=C["surface"]); ti.pack(side="left", padx=14, pady=12)
        tk.Label(ti, text=title, bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 13, "bold")).pack(anchor="w")
        tk.Label(ti, text=subtitle, bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(2, 0))
        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        # column headers
        ch = tk.Frame(self, bg=C["card"])
        ch.pack(fill="x")
        tk.Label(ch, text="On", bg=C["card"], fg=C["accent"],
                 font=("Segoe UI", 8, "bold"), width=4).pack(side="left", padx=(14, 0), pady=5)
        tk.Label(ch, text="Field", bg=C["card"], fg=C["accent"],
                 font=("Segoe UI", 8, "bold")).pack(side="left", padx=8, pady=5)
        tk.Label(ch, text="Jira API Key", bg=C["card"], fg=C["accent"],
                 font=("Segoe UI", 8, "bold")).pack(side="right", padx=14, pady=5)
        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        # scrollable list
        outer = tk.Frame(self, bg=C["bg"])
        outer.pack(fill="both", expand=True)
        self._canvas = tk.Canvas(outer, bg=C["bg"], highlightthickness=0)
        vbar = ttk.Scrollbar(outer, orient="vertical", command=self._canvas.yview)
        self._canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side="right", fill="y")
        self._canvas.pack(side="left", fill="both", expand=True)
        self._list_frame = tk.Frame(self._canvas, bg=C["bg"])
        win = self._canvas.create_window((0, 0), window=self._list_frame, anchor="nw")
        self._canvas.bind("<Configure>", lambda e: self._canvas.itemconfig(win, width=e.width))
        self._list_frame.bind("<Configure>",
                              lambda e: self._canvas.configure(
                                  scrollregion=self._canvas.bbox("all")))
        self._canvas.bind("<MouseWheel>",
                          lambda e: self._canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        self._render_rows()

        # bottom bar
        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")
        bot = tk.Frame(self, bg=C["surface"])
        bot.pack(fill="x", side="bottom")
        inner_bot = tk.Frame(bot, bg=C["surface"])
        inner_bot.pack(fill="x", padx=14, pady=10)

        tk.Button(inner_bot, text="⊕  Add Custom Jira Field",
                  bg=C["accent_dim"], fg=C["accent"], relief="flat",
                  font=("Segoe UI", 9, "bold"), padx=12, pady=6,
                  activebackground=C["border"], cursor="hand2",
                  command=self._add_custom_jira_field).pack(side="left")

        tk.Button(inner_bot, text="✦  Apply & Close",
                  bg=C["green"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=16, pady=6,
                  activebackground="#2ea043", cursor="hand2",
                  command=self._apply).pack(side="right")
        tk.Button(inner_bot, text="✕  Cancel",
                  bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground="#c0392b", cursor="hand2",
                  command=self.destroy).pack(side="right", padx=(0, 8))

    def _render_rows(self):
        for w in self._list_frame.winfo_children():
            w.destroy()
        self._vars = {}

        filtered = [fd for fd in self._all_fields if self._filter_fn(fd)]
        for i, fd in enumerate(filtered):
            bg  = C["card"] if i % 2 == 0 else C["surface"]
            req = fd.get("required", False)

            row = tk.Frame(self._list_frame, bg=bg)
            row.pack(fill="x")

            en_var = tk.BooleanVar(value=fd.get("enabled", True))
            self._vars[fd["key"]] = en_var

            cb = tk.Checkbutton(row, variable=en_var, bg=bg, fg=C["text"],
                                selectcolor=C["cb_sel"], activebackground=bg,
                                activeforeground=C["text"],
                                relief="flat", bd=0, cursor="hand2")
            cb.pack(side="left", padx=(14, 4), pady=8)

            lbl_fg = C["text"] if fd.get("enabled", True) else C["muted"]
            tk.Label(row, text=fd["label"],
                     bg=bg, fg=lbl_fg,
                     font=("Segoe UI", 10, "bold" if req else "normal"),
                     anchor="w").pack(side="left", fill="x", expand=True, pady=8)

            if req:
                tk.Label(row, text="req", bg=bg, fg=C["purple"],
                         font=("Segoe UI", 7, "italic")).pack(side="right", padx=(4, 2), pady=8)

            if fd.get("custom") and not req:
                def _del(fd=fd):
                    if messagebox.askyesno("Remove", f'Remove "{fd["label"]}"?', parent=self):
                        self._all_fields.remove(fd)
                        self._render_rows()
                tk.Button(row, text="✕", bg=bg, fg=C["red"], relief="flat",
                          font=("Segoe UI", 9), padx=4,
                          activebackground=C["border"], cursor="hand2",
                          command=_del).pack(side="right", padx=(2, 6), pady=4)

            jk = fd.get("jira_key", "") or ""
            tk.Label(row, text=jk, bg=bg, fg=C["muted"],
                     font=("Segoe UI", 8, "italic"),
                     anchor="e", padx=12).pack(side="right", pady=8)

            tk.Frame(self._list_frame, bg=C["border"], height=1).pack(fill="x")

    def _add_custom_jira_field(self):
        dlg = tk.Toplevel(self)
        dlg.title("Add Custom Jira Field")
        dlg.geometry("480x340")
        dlg.configure(bg=C["bg"])
        dlg.grab_set()
        dlg.transient(self)

        tk.Frame(dlg, bg=C["accent"], height=3).pack(fill="x")
        hdr = tk.Frame(dlg, bg=C["surface"]); hdr.pack(fill="x")
        tk.Frame(hdr, bg=C["accent"], width=4).pack(side="left", fill="y")
        hdr_inner = tk.Frame(hdr, bg=C["surface"])
        hdr_inner.pack(side="left", padx=12, pady=10)
        tk.Label(hdr_inner, text="⊕  New Custom Jira API Field",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 12, "bold")).pack(anchor="w")
        tk.Label(hdr_inner, text="Field sent directly to the Jira API when creating a ticket",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(2, 0))
        tk.Frame(dlg, bg=C["border"], height=1).pack(fill="x")

        def _row(label, idx):
            row_bg = C["card"] if idx % 2 == 0 else C["bg"]
            f = tk.Frame(dlg, bg=row_bg); f.pack(fill="x", padx=14, pady=0)
            inner = tk.Frame(f, bg=row_bg); inner.pack(fill="x", padx=6, pady=6)
            tk.Label(inner, text=label, bg=row_bg, fg=C["muted"],
                     width=18, anchor="w", font=("Segoe UI", 9)).pack(side="left")
            return inner

        lbl_var = tk.StringVar()
        f1 = _row("Field Label *", 0)
        ttk.Entry(f1, textvariable=lbl_var, width=28).pack(side="left", fill="x", expand=True)

        jira_key_var = tk.StringVar()
        f2 = _row("Jira API Key *", 1)
        ttk.Entry(f2, textvariable=jira_key_var, width=28).pack(side="left", fill="x", expand=True)
        tk.Label(f2, text="e.g. assignee, customfield_10001",
                 bg=C["bg"], fg=C["muted"], font=("Segoe UI", 7)).pack(side="left", padx=6)

        type_var = tk.StringVar(value="text")
        f3 = _row("Field Type", 2)
        ttk.Combobox(f3, textvariable=type_var,
                     values=["text", "dropdown", "number", "date"],
                     state="readonly", width=14).pack(side="left")

        def_var = tk.StringVar()
        f4 = _row("Default Value", 3)
        ttk.Entry(f4, textvariable=def_var, width=28).pack(side="left", fill="x", expand=True)

        choices_var = tk.StringVar()
        f5 = _row("Choices (csv)", 4)
        choices_entry = ttk.Entry(f5, textvariable=choices_var, width=28)

        def _on_type(*_):
            if type_var.get() == "dropdown":
                choices_entry.pack(side="left")
            else:
                choices_entry.pack_forget()
        type_var.trace_add("write", _on_type)

        def _create():
            label = lbl_var.get().strip()
            jk    = jira_key_var.get().strip()
            if not label:
                messagebox.showwarning("Required", "Field Label is required.", parent=dlg)
                return
            if not jk:
                messagebox.showwarning("Required", "Jira API Key is required.", parent=dlg)
                return
            key = re.sub(r'\W+', '_', label.lower()).strip('_') + "_jira_custom"
            existing_keys = {fd["key"] for fd in self._all_fields}
            if key in existing_keys:
                key = key + "_2"
            choices = [c.strip() for c in choices_var.get().split(",") if c.strip()] \
                      if type_var.get() == "dropdown" else []
            new_fd = dict(
                label             = label,
                key               = key,
                type              = type_var.get(),
                choices           = choices,
                default           = def_var.get().strip(),
                jira_key          = jk,
                jira_field        = True,
                show_label_in_jira= False,
                required          = False,
                enabled           = True,
                custom            = True,
            )
            # insert after last jira_field entry
            last_jira_idx = max(
                (i for i, fd in enumerate(self._all_fields) if fd.get("jira_field")),
                default=-1)
            self._all_fields.insert(last_jira_idx + 1, new_fd)
            self._render_rows()
            dlg.destroy()

        tk.Frame(dlg, bg=C["border"], height=1).pack(fill="x")
        brow = tk.Frame(dlg, bg=C["surface"]); brow.pack(fill="x", padx=16, pady=12)
        tk.Button(brow, text="⊕  Add Field",
                  bg=C["green"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=14, pady=6,
                  cursor="hand2", activebackground="#2ea043",
                  command=_create).pack(side="right")
        tk.Button(brow, text="✕  Cancel",
                  bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground="#c0392b", cursor="hand2",
                  command=dlg.destroy).pack(side="right", padx=(0, 8))

    def _apply(self):
        for fd in self._all_fields:
            if fd["key"] in self._vars:
                fd["enabled"] = self._vars[fd["key"]].get()
        self.result = self._all_fields
        self.destroy()


# ─────────────────────────────────────────────────────────────────────────────
#  DescInfoDialog  – manage description-body fields with value inputs
# ─────────────────────────────────────────────────────────────────────────────

class DescInfoDialog(tk.Toplevel):
    """
    Shows all non-jira_field fields.  Each row has:
      • enable/disable checkbox
      • field label
      • actual input widget (dropdown or text) for editing the value/default
    On apply: persists enabled state + current value as new default.
    """

    def __init__(self, parent, fields):
        super().__init__(parent)
        self.title(f"{APP_NAME}  —  Description Info")
        self.geometry("600x580")
        self.minsize(480, 400)
        self.configure(bg=C["bg"])
        self.grab_set()
        self.transient(parent)

        self._fields = copy.deepcopy(fields)
        self.result  = None
        self._rows   = []

        self._build()

    def _build(self):
        tk.Frame(self, bg=C["purple"], height=3).pack(fill="x")
        hdr = tk.Frame(self, bg=C["surface"])
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=C["purple"], width=4).pack(side="left", fill="y")
        ti = tk.Frame(hdr, bg=C["surface"]); ti.pack(side="left", padx=14, pady=12)
        hdr_row = tk.Frame(ti, bg=C["surface"]); hdr_row.pack(anchor="w")
        tk.Label(hdr_row, text="✐", bg=C["surface"], fg=C["purple"],
                 font=("Segoe UI", 14)).pack(side="left", padx=(0, 8))
        tk.Label(hdr_row, text="Description Info",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 13, "bold")).pack(side="left")
        tk.Label(ti, text="Fields written into the Jira description body as  Label: Value",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(2, 0))
        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        # column headers
        ch = tk.Frame(self, bg=C["card"]); ch.pack(fill="x")
        for txt, anchor, side, pad in [
            ("On",          "w", "left", (14, 4)),
            ("Field Label", "w", "left", (4, 8)),
            ("Value",       "w", "left", (0, 0)),
        ]:
            tk.Label(ch, text=txt, bg=C["card"], fg=C["purple"],
                     font=("Segoe UI", 8, "bold"), anchor=anchor
                     ).pack(side=side, padx=pad, pady=5)
        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        # scrollable list
        outer = tk.Frame(self, bg=C["bg"]); outer.pack(fill="both", expand=True)
        canvas = tk.Canvas(outer, bg=C["bg"], highlightthickness=0)
        vbar = ttk.Scrollbar(outer, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)
        self._list_frame = tk.Frame(canvas, bg=C["bg"])
        win = canvas.create_window((0, 0), window=self._list_frame, anchor="nw")
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(win, width=e.width))
        self._list_frame.bind("<Configure>",
                              lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<MouseWheel>",
                    lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        self._render_rows()

        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")
        bot = tk.Frame(self, bg=C["surface"]); bot.pack(fill="x", side="bottom")
        inner_bot = tk.Frame(bot, bg=C["surface"]); inner_bot.pack(fill="x", padx=14, pady=10)

        tk.Button(inner_bot, text="⊕  Add Custom Field",
                  bg=C["accent_dim"], fg=C["accent"], relief="flat",
                  font=("Segoe UI", 9, "bold"), padx=12, pady=6,
                  activebackground=C["border"], cursor="hand2",
                  command=self._add_field).pack(side="left")

        tk.Button(inner_bot, text="✦  Apply & Close",
                  bg=C["green"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=16, pady=6,
                  activebackground="#2ea043", cursor="hand2",
                  command=self._apply).pack(side="right")
        tk.Button(inner_bot, text="✕  Cancel",
                  bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground="#c0392b", cursor="hand2",
                  command=self.destroy).pack(side="right", padx=(0, 8))

    def _render_rows(self):
        for w in self._list_frame.winfo_children():
            w.destroy()
        self._rows = []

        desc_fields = [fd for fd in self._fields if not fd.get("jira_field")]
        for i, fd in enumerate(desc_fields):
            bg      = C["card"] if i % 2 == 0 else C["surface"]
            ftype   = fd.get("type", "text")
            choices = fd.get("choices", [])

            row = tk.Frame(self._list_frame, bg=bg)
            row.pack(fill="x", padx=2, pady=0)

            # enable checkbox
            en_var = tk.BooleanVar(value=fd.get("enabled", True))
            tk.Checkbutton(row, variable=en_var, bg=bg, fg=C["text"],
                           selectcolor=C["cb_sel"], activebackground=bg,
                           activeforeground=C["text"],
                           relief="flat", bd=0, cursor="hand2"
                           ).pack(side="left", padx=(12, 4), pady=6)

            # label column (fixed width)
            lf = tk.Frame(row, bg=bg, width=152); lf.pack(side="left", fill="y")
            lf.pack_propagate(False)
            lbl_w = tk.Label(lf, text=fd["label"], bg=bg,
                             fg=C["text"] if fd.get("enabled", True) else C["muted"],
                             font=("Segoe UI", 9, "bold" if fd.get("required") else "normal"),
                             anchor="w")
            lbl_w.pack(anchor="w", pady=6)
            hint = FIELD_HINTS.get(fd.get("key", ""), "")
            if hint:
                Tooltip(lbl_w, hint)

            # value input
            val_var = tk.StringVar(value=fd.get("default", ""))
            if ftype == "dropdown" and choices:
                w = ttk.Combobox(row, textvariable=val_var, values=choices, state="normal")
            else:
                w = ttk.Entry(row, textvariable=val_var)
            w.pack(side="left", fill="x", expand=True, padx=(4, 4), pady=5)

            # edit choices button for dropdown fields
            if ftype == "dropdown":
                def _edit_cb(fd=fd, cb=w, v=val_var):
                    self._edit_choices(fd, cb, v)
                tk.Button(row, text="✎", bg=bg, fg=C["purple"],
                          relief="flat", font=("Segoe UI", 10), padx=5, pady=1,
                          activebackground=C["border"], cursor="hand2",
                          command=_edit_cb).pack(side="left", padx=(0, 4))

            # delete button for custom fields
            if fd.get("custom"):
                def _del(fd=fd):
                    if messagebox.askyesno("Remove", f'Remove "{fd["label"]}"?', parent=self):
                        self._fields.remove(fd)
                        self._render_rows()
                tk.Button(row, text="✕", bg=bg, fg=C["red"],
                          relief="flat", font=("Segoe UI", 9), padx=4, pady=1,
                          activebackground=C["border"], cursor="hand2",
                          command=_del).pack(side="left", padx=(0, 8))

            tk.Frame(self._list_frame, bg=C["border"], height=1).pack(fill="x")
            self._rows.append(dict(fd=fd, en_var=en_var, val_var=val_var))

    def _edit_choices(self, fd, combobox, val_var):
        dlg = tk.Toplevel(self)
        dlg.title(f"Edit choices — {fd['label']}")
        dlg.geometry("360x320")
        dlg.configure(bg=C["bg"])
        dlg.grab_set()
        dlg.transient(self)

        tk.Label(dlg, text=f"Choices for \"{fd['label']}\"",
                 bg=C["bg"], fg=C["accent"],
                 font=("Segoe UI", 11, "bold")).pack(anchor="w", padx=16, pady=(14, 2))
        tk.Label(dlg, text="One value per line",
                 bg=C["bg"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", padx=16, pady=(0, 8))

        txt = scrolledtext.ScrolledText(dlg, height=10,
                                        bg=C["inp"], fg=C["text"],
                                        insertbackground=C["text"],
                                        relief="flat", font=("Segoe UI", 10),
                                        highlightthickness=1,
                                        highlightbackground=C["border"])
        txt.pack(fill="both", expand=True, padx=16, pady=(0, 10))
        txt.insert("1.0", "\n".join(fd.get("choices", [])))

        def _save():
            lines = [l.strip() for l in txt.get("1.0", "end-1c").splitlines() if l.strip()]
            fd["choices"] = lines
            if combobox:
                combobox.config(values=lines)
                if val_var.get() not in lines and lines:
                    val_var.set(lines[0])
            dlg.destroy()

        br = tk.Frame(dlg, bg=C["bg"]); br.pack(fill="x", padx=16, pady=(0, 14))
        tk.Button(br, text="✦  Save", bg=C["green"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=14, pady=6,
                  cursor="hand2", command=_save).pack(side="right")
        tk.Button(br, text="✕  Cancel", bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground="#c0392b", cursor="hand2",
                  command=dlg.destroy).pack(side="right", padx=(0, 8))

    def _add_field(self):
        dlg = tk.Toplevel(self)
        dlg.title("Add Custom Description Field")
        dlg.geometry("460x340")
        dlg.configure(bg=C["bg"])
        dlg.grab_set()
        dlg.transient(self)

        tk.Frame(dlg, bg=C["purple"], height=3).pack(fill="x")
        hdr = tk.Frame(dlg, bg=C["surface"]); hdr.pack(fill="x")
        tk.Frame(hdr, bg=C["purple"], width=4).pack(side="left", fill="y")
        hdr_inner = tk.Frame(hdr, bg=C["surface"]); hdr_inner.pack(side="left", padx=12, pady=10)
        tk.Label(hdr_inner, text="⊕  New Description Field",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 12, "bold")).pack(anchor="w")
        tk.Label(hdr_inner, text="Appears in Jira description as  Label: Value",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(2, 0))
        tk.Frame(dlg, bg=C["border"], height=1).pack(fill="x")

        def _row(lbl_text, idx):
            bg = C["card"] if idx % 2 == 0 else C["bg"]
            f = tk.Frame(dlg, bg=bg); f.pack(fill="x", padx=14)
            inner = tk.Frame(f, bg=bg); inner.pack(fill="x", padx=6, pady=5)
            tk.Label(inner, text=lbl_text, bg=bg, fg=C["muted"],
                     width=16, anchor="w", font=("Segoe UI", 9)).pack(side="left")
            return inner

        lbl_var = tk.StringVar()
        ttk.Entry(_row("Field Label *", 0), textvariable=lbl_var, width=26
                  ).pack(side="left", fill="x", expand=True)

        type_var = tk.StringVar(value="text")
        r2 = _row("Input Type", 1)
        ttk.Combobox(r2, textvariable=type_var,
                     values=["text", "dropdown", "number", "date"],
                     state="readonly", width=14).pack(side="left")

        def_var = tk.StringVar()
        ttk.Entry(_row("Default Value", 2), textvariable=def_var, width=26
                  ).pack(side="left", fill="x", expand=True)

        choices_var = tk.StringVar()
        r4 = _row("Choices (csv)", 3)
        choices_entry = ttk.Entry(r4, textvariable=choices_var, width=26)

        def _on_type(*_):
            if type_var.get() == "dropdown":
                choices_entry.pack(side="left", fill="x", expand=True)
            else:
                choices_entry.pack_forget()
        type_var.trace_add("write", _on_type); _on_type()

        def _create():
            label = lbl_var.get().strip()
            if not label:
                messagebox.showwarning("Required", "Field Label is required", parent=dlg)
                return
            key = label.lower().replace(" ", "_")
            choices = [c.strip() for c in choices_var.get().split(",") if c.strip()] \
                      if type_var.get() == "dropdown" else []
            self._fields.append(dict(
                label=label, key=key, type=type_var.get(), choices=choices,
                default=def_var.get().strip(), jira_key=None,
                show_label_in_jira=True, required=False, enabled=True, custom=True,
            ))
            self._render_rows()
            dlg.destroy()

        tk.Frame(dlg, bg=C["border"], height=1).pack(fill="x")
        brow = tk.Frame(dlg, bg=C["surface"]); brow.pack(fill="x", padx=16, pady=12)
        tk.Button(brow, text="⊕  Add Field", bg=C["green"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=14, pady=6,
                  cursor="hand2", activebackground="#2ea043",
                  command=_create).pack(side="right")
        tk.Button(brow, text="✕  Cancel", bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground="#c0392b", cursor="hand2",
                  command=dlg.destroy).pack(side="right", padx=(0, 8))

    def _apply(self):
        for meta in self._rows:
            meta["fd"]["enabled"] = meta["en_var"].get()
            meta["fd"]["default"] = meta["val_var"].get()
        # copy updated desc fields back into full fields list
        updated_keys = {fd["key"] for fd in self._fields if not fd.get("jira_field")}
        updated_map  = {fd["key"]: fd for fd in self._fields}
        # merge: keep jira fields from original order, replace desc fields with updated
        merged = []
        seen = set()
        for fd in self._fields:
            if fd["key"] not in seen:
                merged.append(fd)
                seen.add(fd["key"])
        self.result = merged
        self.destroy()


# ─────────────────────────────────────────────────────────────────────────────
#  DiffViewPopup  — GitHub-style split diff viewer
# ─────────────────────────────────────────────────────────────────────────────

class DiffViewPopup(tk.Toplevel):

    def __init__(self, parent, filename, patch, status, adds=0, dels=0):
        super().__init__(parent)
        self.title(f"{APP_NAME}  —  {os.path.basename(filename)}")
        self.geometry("1300x720")
        self.minsize(900, 480)
        self.configure(bg=C["bg"])
        self.transient(parent)
        self._build(filename, patch, status, adds, dels)

    # ── patch parser ──────────────────────────────────────────────────────────

    @staticmethod
    def _split_rows(patch):
        """Convert unified diff patch to split-diff row list.

        Each row: (l_type, l_ln, l_text, r_type, r_ln, r_text)
        l/r_type: 'del' | 'add' | 'ctx' | 'hunk' | 'empty'
        """
        unified = []
        old_ln = new_ln = 0
        for line in patch.splitlines():
            if line.startswith("@@"):
                m = re.match(r"@@ -(\d+)(?:,\d+)? \+(\d+)(?:,\d+)? @@(.*)", line)
                if m:
                    old_ln = int(m.group(1))
                    new_ln = int(m.group(2))
                unified.append(("hunk", None, None, line))
            elif line.startswith("+") and not line.startswith("+++"):
                unified.append(("add", None, new_ln, line[1:]))
                new_ln += 1
            elif line.startswith("-") and not line.startswith("---"):
                unified.append(("del", old_ln, None, line[1:]))
                old_ln += 1
            elif line and line[0] not in ("\\",):
                text = line[1:] if line.startswith(" ") else line
                unified.append(("ctx", old_ln, new_ln, text))
                old_ln += 1; new_ln += 1

        rows = []
        i = 0
        while i < len(unified):
            rtype, a, b, text = unified[i]
            if rtype == "hunk":
                rows.append(("hunk", None, text, "hunk", None, text))
                i += 1
            elif rtype == "ctx":
                rows.append(("ctx", a, text, "ctx", b, text))
                i += 1
            else:
                dels, adds = [], []
                while i < len(unified) and unified[i][0] in ("del", "add"):
                    t, oa, nb, tx = unified[i]
                    (dels if t == "del" else adds).append(
                        (oa if t == "del" else nb, tx))
                    i += 1
                for j in range(max(len(dels), len(adds))):
                    lv = dels[j] if j < len(dels) else None
                    rv = adds[j] if j < len(adds) else None
                    rows.append((
                        "del"   if lv else "empty", lv[0] if lv else None, lv[1] if lv else "",
                        "add"   if rv else "empty", rv[0] if rv else None, rv[1] if rv else "",
                    ))
        return rows

    # ── UI ────────────────────────────────────────────────────────────────────

    def _build(self, filename, patch, status, adds, dels):
        # ── top accent stripe ─────────────────────────────────────────────────
        tk.Frame(self, bg=C["accent"], height=3).pack(fill="x")

        # ── header bar ───────────────────────────────────────────────────────
        hdr = tk.Frame(self, bg=C["surface"])
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=C["accent"], width=4).pack(side="left", fill="y")
        hdr_inner = tk.Frame(hdr, bg=C["surface"])
        hdr_inner.pack(side="left", padx=(10, 0), pady=10)
        badge_bg = STATUS_BADGE.get(status, C["surface"])
        tk.Label(hdr_inner, text=f" {STATUS_EMOJI.get(status,'~')} {status.upper()} ",
                 bg=badge_bg, fg="#ffffff",
                 font=("Segoe UI", 8, "bold"), padx=4).pack(side="left", padx=(0, 10))
        tk.Label(hdr_inner, text=f"{filename}",
                 bg=C["surface"], fg=C["text"],
                 font=("Consolas", 10, "bold")).pack(side="left")
        sfg = C["green"] if adds > dels else C["red"] if dels > adds else C["muted"]
        tk.Label(hdr, text=f"  +{adds}  −{dels}",
                 bg=C["surface"], fg=sfg,
                 font=("Segoe UI", 9, "bold")).pack(side="left", padx=4)

        # ── column labels ─────────────────────────────────────────────────────
        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")
        col_hdr = tk.Frame(self, bg="#161b22")
        col_hdr.pack(fill="x")
        tk.Label(col_hdr, text="  Before (old)",
                 bg="#161b22", fg=C["muted"],
                 font=("Segoe UI", 9, "bold"), anchor="w",
                 ).pack(side="left", fill="x", expand=True, padx=8, pady=4)
        tk.Frame(col_hdr, bg=C["border"], width=1).pack(side="left", fill="y")
        tk.Label(col_hdr, text="  After (new)",
                 bg="#161b22", fg=C["muted"],
                 font=("Segoe UI", 9, "bold"), anchor="w",
                 ).pack(side="left", fill="x", expand=True, padx=8, pady=4)
        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        # ── pane area ─────────────────────────────────────────────────────────
        pane = tk.Frame(self, bg="#0d1117")
        pane.pack(fill="both", expand=True)

        vsb = ttk.Scrollbar(pane, orient="vertical")
        vsb.pack(side="right", fill="y")

        xsb = ttk.Scrollbar(self, orient="horizontal")
        xsb.pack(fill="x", side="bottom")

        lf = tk.Frame(pane, bg="#0d1117"); lf.pack(side="left", fill="both", expand=True)
        tk.Frame(pane, bg=C["border"], width=1).pack(side="left", fill="y")
        rf = tk.Frame(pane, bg="#0d1117"); rf.pack(side="left", fill="both", expand=True)

        txt_kw = dict(
            bg="#0d1117", fg=C["text"],
            font=("Consolas", 9), relief="flat",
            state="disabled", wrap="none",
            highlightthickness=0, insertbackground=C["text"],
        )

        self._lt = tk.Text(lf, **txt_kw)
        self._rt = tk.Text(rf, **txt_kw)

        _TAG = {
            "del":   dict(background="#200e0e", foreground="#f85149"),
            "add":   dict(background="#0e2318", foreground="#3fb950"),
            "ctx":   dict(foreground="#8b949e"),
            "hunk":  dict(background="#0c1c2c", foreground="#58a6ff"),
            "empty": dict(background="#161b22"),
            "ln":    dict(foreground="#484f58"),
        }
        for t in (self._lt, self._rt):
            for tag, cfg in _TAG.items():
                t.tag_config(tag, **cfg)

        # ── populate ──────────────────────────────────────────────────────────
        rows = self._split_rows(patch)
        self._lt.config(state="normal")
        self._rt.config(state="normal")

        for l_type, l_ln, l_text, r_type, r_ln, r_text in rows:
            if l_type == "hunk":
                self._lt.insert("end", f"  {l_text}\n", "hunk")
                self._rt.insert("end", f"  {r_text}\n", "hunk")
            else:
                l_ln_s = f"{l_ln:>5} " if l_ln is not None else "       "
                r_ln_s = f"{r_ln:>5} " if r_ln is not None else "       "
                sym_l  = "−" if l_type == "del" else (" " if l_type == "ctx" else " ")
                sym_r  = "+" if r_type == "add" else (" " if r_type == "ctx" else " ")

                if l_type == "empty":
                    self._lt.insert("end", "\n", "empty")
                else:
                    self._lt.insert("end", l_ln_s, "ln")
                    self._lt.insert("end", f"{sym_l} {l_text}\n", l_type)

                if r_type == "empty":
                    self._rt.insert("end", "\n", "empty")
                else:
                    self._rt.insert("end", r_ln_s, "ln")
                    self._rt.insert("end", f"{sym_r} {r_text}\n", r_type)

        self._lt.config(state="disabled")
        self._rt.config(state="disabled")

        # ── scrollbar wiring ─────────────────────────────────────────────────
        def _yscroll(*args):
            self._lt.yview(*args)
            self._rt.yview(*args)

        def _on_scroll_l(first, last):
            vsb.set(first, last)
            self._rt.yview_moveto(first)

        def _on_scroll_r(first, last):
            vsb.set(first, last)
            self._lt.yview_moveto(first)

        vsb.config(command=_yscroll)
        self._lt.config(yscrollcommand=_on_scroll_l, xscrollcommand=xsb.set)
        self._rt.config(yscrollcommand=_on_scroll_r, xscrollcommand=xsb.set)
        xsb.config(command=lambda *a: [self._lt.xview(*a), self._rt.xview(*a)])

        self._lt.pack(fill="both", expand=True)
        self._rt.pack(fill="both", expand=True)

        def _mw(e):
            self._lt.yview_scroll(int(-1*(e.delta/120)), "units")
            self._rt.yview_scroll(int(-1*(e.delta/120)), "units")
            return "break"
        self._lt.bind("<MouseWheel>", _mw)
        self._rt.bind("<MouseWheel>", _mw)


# ─────────────────────────────────────────────────────────────────────────────
#  FileChangesPopup  (unchanged from v2, just cleaner)
# ─────────────────────────────────────────────────────────────────────────────

class FileChangesPopup(tk.Toplevel):

    def __init__(self, parent, files, pr_title):
        super().__init__(parent)
        self.title(f"{APP_NAME}  —  {pr_title[:70]}")
        self.geometry("1120x760")
        self.minsize(900, 560)
        self.configure(bg=C["bg"])
        self.grab_set()
        self.transient(parent)
        self._files  = _topo_sort_files(files)
        self._rows   = []
        self._result = None
        self._build()

    def _build(self):
        # ── top accent stripe ─────────────────────────────────────────────────
        tk.Frame(self, bg=C["accent"], height=3).pack(fill="x")

        bar = tk.Frame(self, bg=C["surface"])
        bar.pack(fill="x")
        tk.Frame(bar, bg=C["accent"], width=4).pack(side="left", fill="y")
        hdr_f = tk.Frame(bar, bg=C["surface"])
        hdr_f.pack(side="left", padx=14, pady=12, fill="x", expand=True)
        hdr_row = tk.Frame(hdr_f, bg=C["surface"])
        hdr_row.pack(anchor="w")
        tk.Label(hdr_row, text="◈", bg=C["surface"], fg=C["accent"],
                 font=("Segoe UI", 14)).pack(side="left", padx=(0, 8))
        tk.Label(hdr_row,
                 text=f"{len(self._files)} Files Changed",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 13, "bold")).pack(side="left")
        tk.Label(hdr_f,
                 text="Sorted by dependency  •  Change Description → Word doc  •  Jira Comment → Jira ticket",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(2, 0))

        for lbl, st, abg, afg, aabg in [
            ("✓  Select All",   True,  C["accent"],      "#ffffff",   C["hover"]),
            ("✕  Deselect All", False, C["accent_dim"],  C["accent"], C["border"]),
        ]:
            tk.Button(bar, text=lbl,
                      bg=abg, fg=afg, relief="flat",
                      font=("Segoe UI", 9, "bold"), padx=10, pady=6,
                      activebackground=aabg, cursor="hand2",
                      command=lambda s=st: self._toggle_all(s)
                      ).pack(side="right", padx=4, pady=10)

        tk.Button(bar, text="↺  Reset Comments",
                  bg=C["card2"], fg=C["yellow"], relief="flat",
                  font=("Segoe UI", 9, "bold"), padx=10, pady=6,
                  activebackground=C["border"], cursor="hand2",
                  command=self._reset_step_comments
                  ).pack(side="right", padx=4, pady=10)

        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        wrap = tk.Frame(self, bg=C["bg"])
        wrap.pack(fill="both", expand=True)
        self._canvas = tk.Canvas(wrap, bg=C["bg"], highlightthickness=0)
        vbar = ttk.Scrollbar(wrap, orient="vertical", command=self._canvas.yview)
        hbar = ttk.Scrollbar(wrap, orient="horizontal", command=self._canvas.xview)
        self._canvas.configure(yscrollcommand=vbar.set, xscrollcommand=hbar.set)
        hbar.pack(side="bottom", fill="x")
        vbar.pack(side="right",  fill="y")
        self._canvas.pack(side="left", fill="both", expand=True)
        self._lf = tk.Frame(self._canvas, bg=C["bg"])
        wid = self._canvas.create_window((0, 0), window=self._lf, anchor="nw")
        self._canvas.bind("<Configure>",
                          lambda e: self._canvas.itemconfig(wid, width=e.width))
        self._lf.bind("<Configure>",
                      lambda e: self._canvas.configure(
                          scrollregion=self._canvas.bbox("all")))
        self._canvas.bind("<MouseWheel>",
                          lambda e: self._canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        self._build_rows()

        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")
        bot = tk.Frame(self, bg=C["surface"], height=52)
        bot.pack(fill="x", side="bottom"); bot.pack_propagate(False)
        self._sv = tk.StringVar(value="")
        tk.Label(bot, textvariable=self._sv, bg=C["surface"], fg=C["green"],
                 font=("Segoe UI", 9)).pack(side="left", padx=16)
        tk.Button(bot, text="Cancel",
                  bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=5,
                  activebackground="#c0392b", cursor="hand2",
                  command=self.destroy).pack(side="right", padx=8, pady=10)
        tk.Button(bot, text="  Save & Close  ",
                  bg=C["green"], fg="#0d1117", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=14, pady=5,
                  activebackground="#2ea043", cursor="hand2",
                  command=self._save).pack(side="right", padx=4, pady=10)

    @staticmethod
    def _default_step_comment(step, filename, status):
        verb = STATUS_VERB.get(status, "Modify")
        base = os.path.basename(filename)
        return f"Step {step}: {verb} {base}"

    def _build_rows(self):
        total = len(self._files)
        for idx, f in enumerate(self._files, 1):
            fname   = f.get("filename", "")
            status  = f.get("status", "modified")
            adds    = f.get("additions", 0)
            dels    = f.get("deletions", 0)
            rbg     = STATUS_ROW.get(status, C["card"])
            bbg     = STATUS_BADGE.get(status, C["surface"])
            em      = STATUS_EMOJI.get(status, "~")

            card = tk.Frame(self._lf, bg=rbg,
                            highlightbackground=C["border"], highlightthickness=1)
            card.pack(fill="x", padx=10, pady=3)

            def _sw(e, c=self._canvas):
                c.yview_scroll(int(-1*(e.delta/120)), "units")
            card.bind("<MouseWheel>", _sw)

            hdr = tk.Frame(card, bg=rbg)
            hdr.pack(fill="x", padx=8, pady=(8, 4))

            inc = tk.BooleanVar(value=True)
            tk.Checkbutton(hdr, variable=inc, bg=rbg, fg=C["text"],
                           selectcolor=C["cb_sel"],
                           activebackground=rbg, activeforeground=C["text"],
                           text="Include", font=("Segoe UI", 8),
                           bd=0, relief="flat", cursor="hand2").pack(side="left")

            # step badge
            tk.Label(hdr, text=f"  Step {idx}/{total}  ",
                     bg=C["accent"], fg="#ffffff",
                     font=("Segoe UI", 8, "bold"), padx=4, pady=2
                     ).pack(side="left", padx=(8, 6))

            tk.Label(hdr, text=f"  {em}  {status.upper()}  ",
                     bg=bbg, fg="#ffffff",
                     font=("Segoe UI", 8, "bold"), padx=4, pady=2
                     ).pack(side="left", padx=(0, 10))

            parts = fname.split("/")
            pre   = "/".join(parts[:-2]) + "/" if len(parts) > 2 else ""
            tail  = "/".join(parts[-2:])
            pf    = tk.Frame(hdr, bg=rbg); pf.pack(side="left", fill="x", expand=True)
            if pre:
                tk.Label(pf, text=pre, bg=rbg, fg=C["muted"],
                         font=("Consolas", 10)).pack(side="left")
            tk.Label(pf, text=tail, bg=rbg, fg=C["text"],
                     font=("Consolas", 10, "bold")).pack(side="left")

            sfg = C["green"] if adds > dels else C["red"] if dels > adds else C["muted"]
            tk.Label(hdr, text=f"+{adds}  -{dels}", bg=rbg, fg=sfg,
                     font=("Segoe UI", 9, "bold")).pack(side="right", padx=8)

            body = tk.Frame(card, bg=rbg); body.pack(fill="x", padx=8, pady=(0, 8))

            # ── GitHub link + diff toggle row ────────────────────────────────
            gh_row = tk.Frame(body, bg=rbg)
            gh_row.pack(fill="x", pady=(0, 4))

            gh_url = f.get("blob_url") or f.get("html_url", "")
            if gh_url:
                tk.Button(gh_row, text="View on GitHub ↗",
                          bg=rbg, fg=C["accent"], relief="flat",
                          font=("Segoe UI", 8), padx=6, pady=2, cursor="hand2",
                          activebackground=C["border"],
                          command=lambda u=gh_url: webbrowser.open(u)
                          ).pack(side="left")

            patch = f.get("patch", "")
            if patch:
                n_lines = len(patch.splitlines())
                tk.Button(
                    gh_row,
                    text=f"⊞ View Changes  ({n_lines} lines)",
                    bg=rbg, fg=C["accent"], relief="flat",
                    font=("Segoe UI", 8, "bold"), padx=6, pady=2, cursor="hand2",
                    activebackground=C["border"],
                    command=lambda fn=fname, p=patch, st=status, a=adds, d=dels:
                        DiffViewPopup(self, fn, p, st, a, d)
                ).pack(side="left", padx=(6, 0))

            # ── two comment columns ──────────────────────────────────────────
            comment_row = tk.Frame(body, bg=rbg)
            comment_row.pack(fill="x", pady=(0, 4))

            # Word comment (left)
            word_col = tk.Frame(comment_row, bg=rbg)
            word_col.pack(side="left", fill="both", expand=True, padx=(0, 6))
            tk.Label(word_col, text="Change Description (Word doc):", bg=rbg, fg=C["yellow"],
                     font=("Segoe UI", 9, "bold")).pack(anchor="w")
            word_cbox = scrolledtext.ScrolledText(
                word_col, height=3, bg=C["inp"], fg=C["text"],
                insertbackground=C["text"], selectbackground=C["hover"],
                font=("Segoe UI", 10), relief="flat", wrap="word")
            word_cbox.pack(fill="both", expand=True, pady=(2, 0))
            word_cbox.bind("<MouseWheel>", _sw)
            word_cbox.insert("1.0", self._default_step_comment(idx, fname, status))

            # Jira comment (right)
            jira_col = tk.Frame(comment_row, bg=rbg)
            jira_col.pack(side="left", fill="both", expand=True)
            tk.Label(jira_col, text="Jira Comment:", bg=rbg, fg=C["accent"],
                     font=("Segoe UI", 9, "bold")).pack(anchor="w")
            jira_cbox = scrolledtext.ScrolledText(
                jira_col, height=3, bg=C["inp"], fg=C["text"],
                insertbackground=C["text"], selectbackground=C["hover"],
                font=("Segoe UI", 10), relief="flat", wrap="word")
            jira_cbox.pack(fill="both", expand=True, pady=(2, 0))
            jira_cbox.bind("<MouseWheel>", _sw)
            jira_cbox.insert("1.0", self._default_step_comment(idx, fname, status))

            # ── screenshots with per-screenshot comments ─────────────────────
            shots      = []   # list of {"path": str, "var": StringVar, "frame": Frame}
            shots_cv   = tk.StringVar(value="No screenshots")
            shots_area = tk.Frame(body, bg=rbg)
            shots_area.pack(fill="x", pady=(6, 0))

            def _add_shot_row(path, sa=shots_area, s=shots, sv=shots_cv, _rbg=rbg):
                var = tk.StringVar()
                fr  = tk.Frame(sa, bg=C["surface"],
                               highlightbackground=C["border"], highlightthickness=1)
                fr.pack(fill="x", pady=2)
                item = {"path": path, "var": var, "frame": fr}
                s.append(item)
                # thumbnail
                try:
                    from PIL import Image, ImageTk
                    img = Image.open(path); img.thumbnail((56, 42))
                    ph  = ImageTk.PhotoImage(img)
                    lb  = tk.Label(fr, image=ph, bg=C["surface"], cursor="hand2")
                    lb.pack(side="left", padx=(4, 6), pady=4)
                    lb.image = ph
                    lb.bind("<Button-1>", lambda e, p=path: webbrowser.open(p))
                except Exception:
                    tk.Label(fr, text="IMG", bg=C["surface"], fg=C["accent"],
                             font=("Consolas", 9), padx=6).pack(
                                 side="left", padx=(4, 6), pady=4)
                # filename label
                tk.Label(fr, text=os.path.basename(path)[:22],
                         bg=C["surface"], fg=C["muted"],
                         font=("Consolas", 8)).pack(side="left")
                # comment entry
                ent = ttk.Entry(fr, textvariable=var)
                ent.pack(side="left", fill="x", expand=True, padx=(6, 4), pady=4)
                ent.insert(0, "Screenshot caption...")
                ent.bind("<FocusIn>",
                         lambda e, w=ent: w.delete(0, "end")
                         if w.get() == "Screenshot caption..." else None)
                ent.bind("<MouseWheel>", _sw)
                # remove button
                def _rm(i=item, f=fr, s=s, sv=sv):
                    s.remove(i)
                    f.destroy()
                    n = len(s)
                    sv.set(f"{n} screenshot{'s' if n!=1 else ''}" if n else "No screenshots")
                tk.Button(fr, text="✕", bg=C["surface"], fg=C["red"], relief="flat",
                          font=("Segoe UI", 9, "bold"), padx=6, pady=2, cursor="hand2",
                          command=_rm).pack(side="right", padx=4, pady=4)
                n = len(s)
                sv.set(f"{n} screenshot{'s' if n!=1 else ''}" if n else "No screenshots")

            srow = tk.Frame(body, bg=rbg); srow.pack(fill="x")
            tk.Label(srow, textvariable=shots_cv, bg=rbg, fg=C["muted"],
                     font=("Segoe UI", 9)).pack(side="left", padx=(0, 8))

            def _add_shots(sa=shots_area, s=shots, sv=shots_cv):
                paths = filedialog.askopenfilenames(
                    title="Select screenshots",
                    filetypes=[("Images","*.png *.jpg *.jpeg *.gif *.bmp *.webp *.tiff"),
                               ("All","*.*")])
                for p in paths:
                    if p and not any(x["path"] == p for x in s):
                        _add_shot_row(p, sa, s, sv)

            def _clr_shots(s=shots, sa=shots_area, sv=shots_cv):
                for x in s:
                    x["frame"].destroy()
                s.clear()
                sv.set("No screenshots")

            tk.Button(srow, text="+ Add Screenshot",
                      bg=C["green"], fg="#0d1117", relief="flat",
                      font=("Segoe UI", 9, "bold"), padx=10, pady=3,
                      activebackground="#2ea043", cursor="hand2",
                      command=_add_shots).pack(side="left", padx=(0, 6))
            tk.Button(srow, text="Clear",
                      bg=C["orange"], fg="#0d1117", relief="flat",
                      font=("Segoe UI", 9, "bold"), padx=8, pady=3,
                      activebackground=C["yellow"], cursor="hand2",
                      command=_clr_shots).pack(side="left")

            self._rows.append(dict(filename=fname, status=status,
                                   additions=adds, deletions=dels,
                                   word_cbox=word_cbox, jira_cbox=jira_cbox,
                                   shots=shots, inc=inc))

    def _toggle_all(self, state):
        for r in self._rows: r["inc"].set(state)

    def _reset_step_comments(self):
        """Re-fill both comment boxes with default step comments."""
        for idx, r in enumerate(self._rows, 1):
            default = self._default_step_comment(idx, r["filename"], r["status"])
            for box in (r["word_cbox"], r["jira_cbox"]):
                box.delete("1.0", "end")
                box.insert("1.0", default)

    def _save(self):
        res = []
        for r in self._rows:
            if not r["inc"].get(): continue
            word_comment = r["word_cbox"].get("1.0", "end-1c").strip()
            jira_comment = r["jira_cbox"].get("1.0", "end-1c").strip()
            shots_out = [
                {"path": x["path"],
                 "comment": x["var"].get().strip()
                            if x["var"].get().strip() != "Screenshot caption..." else ""}
                for x in r["shots"]
            ]
            res.append(dict(filename=r["filename"], status=r["status"],
                            additions=r["additions"], deletions=r["deletions"],
                            word_comment=word_comment,
                            jira_comment=jira_comment,
                            comment=word_comment,   # backward compat
                            screenshots=shots_out))
        self._result = res
        n_wc = sum(1 for x in res if x["word_comment"])
        n_jc = sum(1 for x in res if x["jira_comment"])
        n_s  = sum(1 for x in res if x["screenshots"])
        self._sv.set(
            f"Saved {len(res)} files  —  {n_wc} word comments  |  "
            f"{n_jc} jira comments  |  {n_s} screenshots")
        self.after(400, self.destroy)

    def get_result(self):
        return self._result if self._result is not None else []


# ─────────────────────────────────────────────────────────────────────────────
#  SQL PR File popup  –  combine changed files into a single runnable script
# ─────────────────────────────────────────────────────────────────────────────

class SqlFilePopup(tk.Toplevel):

    def __init__(self, parent, files, pr_data, config_data, file_comments=None):
        super().__init__(parent)
        self.title(f"{APP_NAME}  —  SQL File Generator")
        self.geometry("1020x700")
        self.minsize(820, 520)
        self.configure(bg=C["bg"])
        self.grab_set()
        self.transient(parent)

        self._files   = _topo_sort_files(files)
        self._pr      = pr_data
        self._cfg     = config_data
        self._rows    = []
        self._result  = None
        # lookup: filename → latest GitHub review comment
        self._gh_comments = {
            fc["filename"]: (fc.get("jira_comment") or fc.get("comment", "")).strip()
            for fc in (file_comments or [])
        }

        pr_num = pr_data.get("number", "combined")
        out_dir = config_data.get("word_doc_output_dir", BASE_DIR)
        pr_dir  = os.path.join(out_dir, f"PR_{pr_num}")
        self._out_var = tk.StringVar(
            value=os.path.join(pr_dir, f"PR_{pr_num}.sql"))

        self._build()

    # ── UI ────────────────────────────────────────────────────────────────────

    def _build(self):
        # ── top accent stripe ─────────────────────────────────────────────────
        tk.Frame(self, bg=C["yellow"], height=3).pack(fill="x")

        bar = tk.Frame(self, bg=C["surface"])
        bar.pack(fill="x")
        tk.Frame(bar, bg=C["yellow"], width=4).pack(side="left", fill="y")
        hf = tk.Frame(bar, bg=C["surface"])
        hf.pack(side="left", padx=14, pady=12, fill="x", expand=True)
        hf_row = tk.Frame(hf, bg=C["surface"])
        hf_row.pack(anchor="w")
        tk.Label(hf_row, text="⊞", bg=C["surface"], fg=C["yellow"],
                 font=("Segoe UI", 14)).pack(side="left", padx=(0, 8))
        tk.Label(hf_row, text="SQL PR File",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 13, "bold")).pack(side="left")
        tk.Label(hf,
                 text="Files sorted by dependency  •  Edit SQL comment per file  "
                      "•  Combined into one runnable .sql script",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(2, 0))
        for lbl, st, abg, afg, aabg in [
            ("✓  Select All",   True,  C["accent"],      "#ffffff", C["hover"]),
            ("✕  Deselect All", False, C["accent_dim"],  C["accent"], C["border"]),
        ]:
            tk.Button(bar, text=lbl, bg=abg, fg=afg,
                      relief="flat", font=("Segoe UI", 9, "bold"), padx=10, pady=6,
                      activebackground=aabg, cursor="hand2",
                      command=lambda s=st: self._toggle_all(s)
                      ).pack(side="right", padx=4, pady=10)

        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        nb = ttk.Notebook(self)
        nb.pack(fill="both", expand=True)

        # ── Tab 1 : Files ─────────────────────────────────────────────────────
        files_tab = tk.Frame(nb, bg=C["bg"])
        nb.add(files_tab, text="  📄  Files  ")

        wrap = tk.Frame(files_tab, bg=C["bg"]); wrap.pack(fill="both", expand=True)
        self._canvas = tk.Canvas(wrap, bg=C["bg"], highlightthickness=0)
        vsb = ttk.Scrollbar(wrap, orient="vertical", command=self._canvas.yview)
        self._canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        self._canvas.pack(side="left", fill="both", expand=True)
        self._lf = tk.Frame(self._canvas, bg=C["bg"])
        wid = self._canvas.create_window((0, 0), window=self._lf, anchor="nw")
        self._canvas.bind("<Configure>",
                          lambda e: self._canvas.itemconfig(wid, width=e.width))
        self._lf.bind("<Configure>",
                      lambda e: self._canvas.configure(
                          scrollregion=(0, 0, e.width, e.height)))
        self._canvas.bind("<MouseWheel>",
                          lambda e: self._canvas.yview_scroll(
                              int(-1*(e.delta/120)), "units"))

        self._build_rows()

        # ── Tab 2 : Keywords ──────────────────────────────────────────────────
        kw_tab = tk.Frame(nb, bg=C["bg"])
        nb.add(kw_tab, text="  🔑  Keywords  ")
        self._build_keywords_tab(kw_tab)

        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        bot = tk.Frame(self, bg=C["surface"]); bot.pack(fill="x", side="bottom")
        out_row = tk.Frame(bot, bg=C["surface"]); out_row.pack(fill="x", padx=16, pady=(10, 4))
        tk.Label(out_row, text="Output file:",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 6))
        ttk.Entry(out_row, textvariable=self._out_var).pack(side="left", fill="x", expand=True)
        tk.Button(out_row, text="Browse",
                  bg=C["purple"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 9, "bold"), padx=8, pady=3, cursor="hand2",
                  activebackground="#9966cc",
                  command=self._browse).pack(side="left", padx=(6, 0))

        btn_row = tk.Frame(bot, bg=C["surface"]); btn_row.pack(fill="x", padx=16, pady=(4, 14))
        self._sv = tk.StringVar(value="")
        tk.Label(btn_row, textvariable=self._sv,
                 bg=C["surface"], fg=C["green"],
                 font=("Segoe UI", 9)).pack(side="left")
        tk.Button(btn_row, text="✕  Cancel",
                  bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6, cursor="hand2",
                  activebackground=C["border"],
                  command=self.destroy).pack(side="right", padx=8)
        tk.Button(btn_row, text="✦  Generate SQL File",
                  bg=C["green"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=14, pady=6, cursor="hand2",
                  activebackground="#2ea043",
                  command=self._generate).pack(side="right", padx=4)

    def _build_rows(self):
        for f in self._files:
            fname    = f.get("filename", "")
            status   = f.get("status", "modified")
            adds     = f.get("additions", 0)
            dels     = f.get("deletions", 0)
            patch    = f.get("patch", "")
            obj_name = _detect_sql_object(fname, patch)
            gh_cmt   = self._gh_comments.get(fname, "")
            cmt_default = gh_cmt if gh_cmt else obj_name
            self._rows.append(dict(
                filename=fname, status=status,
                additions=adds, deletions=dels,
                patch=patch, obj_name=obj_name,
                inc=tk.BooleanVar(value=(status != "removed")),
                cmt_var=tk.StringVar(value=cmt_default),
                raw_url=f.get("raw_url", ""),
                contents_url=f.get("contents_url", ""),
                full_content=f.get("full_content"),
                gh_url=f.get("blob_url") or f.get("html_url", ""),
            ))
        self._render_rows()

    def _render_rows(self):
        for w in self._lf.winfo_children():
            w.destroy()
        total = len(self._rows)

        def _sw(e, c=self._canvas):
            c.yview_scroll(int(-1*(e.delta/120)), "units")

        for i, r in enumerate(self._rows):
            fname    = r["filename"]
            status   = r["status"]
            adds     = r["additions"]
            dels     = r["deletions"]
            inc      = r["inc"]
            cmt_var  = r["cmt_var"]
            obj_name = r["obj_name"]
            gh_url   = r["gh_url"]
            rbg      = STATUS_ROW.get(status, C["card"])
            bbg      = STATUS_BADGE.get(status, C["surface"])
            em       = STATUS_EMOJI.get(status, "~")
            idx      = i + 1

            card = tk.Frame(self._lf, bg=rbg,
                            highlightbackground=C["border"], highlightthickness=1)
            card.pack(fill="x", padx=10, pady=3)
            card.bind("<MouseWheel>", _sw)

            hdr = tk.Frame(card, bg=rbg); hdr.pack(fill="x", padx=8, pady=(8, 4))

            # ▲▼ reorder buttons
            nav = tk.Frame(hdr, bg=rbg); nav.pack(side="left", padx=(0, 4))
            tk.Button(nav, text="▲", bg=C["card"], fg=C["text"], relief="flat",
                      font=("Segoe UI", 7), padx=3, pady=0, cursor="hand2",
                      activebackground=C["border"],
                      state="normal" if i > 0 else "disabled",
                      command=lambda ii=i: self._move_row(ii, -1)
                      ).pack(side="top")
            tk.Button(nav, text="▼", bg=C["card"], fg=C["text"], relief="flat",
                      font=("Segoe UI", 7), padx=3, pady=0, cursor="hand2",
                      activebackground=C["border"],
                      state="normal" if i < total - 1 else "disabled",
                      command=lambda ii=i: self._move_row(ii, 1)
                      ).pack(side="top")

            tk.Checkbutton(hdr, variable=inc, bg=rbg, fg=C["text"],
                           selectcolor=C["cb_sel"],
                           activebackground=rbg, activeforeground=C["text"],
                           text="Include", font=("Segoe UI", 8),
                           bd=0, relief="flat", cursor="hand2").pack(side="left")

            tk.Label(hdr, text=f"  Step {idx}/{total}  ",
                     bg=C["accent"], fg="#ffffff",
                     font=("Segoe UI", 8, "bold"), padx=4, pady=2
                     ).pack(side="left", padx=(8, 6))
            tk.Label(hdr, text=f"  {em}  {status.upper()}  ",
                     bg=bbg, fg="#ffffff",
                     font=("Segoe UI", 8, "bold"), padx=4, pady=2
                     ).pack(side="left", padx=(0, 10))

            parts = fname.split("/")
            pre   = "/".join(parts[:-2]) + "/" if len(parts) > 2 else ""
            tail  = "/".join(parts[-2:])
            pf    = tk.Frame(hdr, bg=rbg); pf.pack(side="left", fill="x", expand=True)
            if pre:
                tk.Label(pf, text=pre, bg=rbg, fg=C["muted"],
                         font=("Consolas", 10)).pack(side="left")
            tk.Label(pf, text=tail, bg=rbg, fg=C["text"],
                     font=("Consolas", 10, "bold")).pack(side="left")
            # Show Code Change — compact, next to filename
            patch = r["patch"]
            if patch:
                n_lines = len(patch.splitlines())
                tk.Button(pf,
                          text=f"⊞ Code Change ({n_lines})",
                          bg=C["accent"], fg="#ffffff", relief="flat",
                          font=("Segoe UI", 8, "bold"),
                          padx=6, pady=1, cursor="hand2",
                          activebackground=C["hover"], activeforeground="#ffffff",
                          command=lambda fn=fname, p=patch, st=status, a=adds, d=dels:
                              DiffViewPopup(self, fn, p, st, a, d)
                          ).pack(side="left", padx=(8, 0))

            sfg = C["green"] if adds > dels else C["red"] if dels > adds else C["muted"]
            tk.Label(hdr, text=f"+{adds}  −{dels}", bg=rbg, fg=sfg,
                     font=("Segoe UI", 9, "bold")).pack(side="right", padx=8)

            body = tk.Frame(card, bg=rbg); body.pack(fill="x", padx=8, pady=(0, 8))

            # SQL comment row — "Step N:" prefix is read-only label, entry holds description only
            cmt_row = tk.Frame(body, bg=rbg); cmt_row.pack(fill="x", pady=(0, 4))
            tk.Label(cmt_row, text="SQL Comment:",
                     bg=rbg, fg=C["yellow"], font=("Segoe UI", 9, "bold")
                     ).pack(side="left", padx=(0, 6))
            tk.Label(cmt_row, text=f"Step {idx}:",
                     bg=rbg, fg=C["accent"], font=("Segoe UI", 9, "bold")
                     ).pack(side="left", padx=(0, 4))
            cmt_ent = ttk.Entry(cmt_row, textvariable=cmt_var)
            cmt_ent.pack(side="left", fill="x", expand=True)
            cmt_ent.bind("<MouseWheel>", _sw)

            info_row = tk.Frame(body, bg=rbg); info_row.pack(fill="x")
            tk.Label(info_row, text=f"Detected: {obj_name}",
                     bg=rbg, fg=C["muted"],
                     font=("Segoe UI", 8, "italic")).pack(side="left")
            if gh_url:
                tk.Button(info_row, text="View on GitHub ↗",
                          bg=rbg, fg=C["accent"], relief="flat",
                          font=("Segoe UI", 8), padx=6, pady=1, cursor="hand2",
                          activebackground=C["border"],
                          command=lambda u=gh_url: webbrowser.open(u)
                          ).pack(side="right")

    def _build_keywords_tab(self, parent):
        tk.Frame(parent, bg=C["yellow"], height=3).pack(fill="x")
        hf = tk.Frame(parent, bg=C["surface"]); hf.pack(fill="x", padx=14, pady=10)
        tk.Label(hf, text="🔑  Keyword Alerts",
                 bg=C["surface"], fg=C["yellow"],
                 font=("Segoe UI", 12, "bold")).pack(anchor="w")
        tk.Label(hf, text="SQL file generation is BLOCKED if any keyword is found in file content."
                          "  Remove or fix the match before generating.",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(2, 0))
        tk.Frame(parent, bg=C["border"], height=1).pack(fill="x")

        list_frame = tk.Frame(parent, bg=C["bg"]); list_frame.pack(fill="both", expand=True, padx=12, pady=8)

        self._kw_listbox = tk.Listbox(list_frame, bg=C["card"], fg=C["text"],
                                      selectbackground=C["accent"], selectforeground="#fff",
                                      font=("Consolas", 10), relief="flat",
                                      highlightthickness=1, highlightcolor=C["border"],
                                      highlightbackground=C["border"])
        kw_sb = ttk.Scrollbar(list_frame, orient="vertical",
                              command=self._kw_listbox.yview)
        self._kw_listbox.configure(yscrollcommand=kw_sb.set)
        kw_sb.pack(side="right", fill="y")
        self._kw_listbox.pack(side="left", fill="both", expand=True)

        for kw in load_keywords():
            self._kw_listbox.insert("end", kw)

        tk.Frame(parent, bg=C["border"], height=1).pack(fill="x")
        ctrl = tk.Frame(parent, bg=C["surface"]); ctrl.pack(fill="x", padx=12, pady=8)

        self._kw_entry_var = tk.StringVar()
        kw_entry = ttk.Entry(ctrl, textvariable=self._kw_entry_var, width=24)
        kw_entry.pack(side="left", padx=(0, 6))

        def _add_kw():
            val = self._kw_entry_var.get().strip()
            if not val:
                return
            existing = list(self._kw_listbox.get(0, "end"))
            if val.lower() in [e.lower() for e in existing]:
                return
            self._kw_listbox.insert("end", val)
            self._kw_entry_var.set("")
            save_keywords(list(self._kw_listbox.get(0, "end")))

        def _del_kw():
            sel = self._kw_listbox.curselection()
            if not sel:
                return
            self._kw_listbox.delete(sel[0])
            save_keywords(list(self._kw_listbox.get(0, "end")))

        kw_entry.bind("<Return>", lambda e: _add_kw())
        tk.Button(ctrl, text="Add", bg=C["accent"], fg="#fff", relief="flat",
                  font=("Segoe UI", 9, "bold"), padx=10, pady=4, cursor="hand2",
                  activebackground=C["hover"], command=_add_kw).pack(side="left", padx=(0, 4))
        tk.Button(ctrl, text="Remove Selected", bg=C["red"], fg="#fff", relief="flat",
                  font=("Segoe UI", 9), padx=10, pady=4, cursor="hand2",
                  activebackground=C["border"], command=_del_kw).pack(side="left")

    def _move_row(self, i, direction):
        j = i + direction
        if 0 <= j < len(self._rows):
            self._rows[i], self._rows[j] = self._rows[j], self._rows[i]
        self._render_rows()

    # ── actions ───────────────────────────────────────────────────────────────

    def _toggle_all(self, state):
        for r in self._rows: r["inc"].set(state)

    def _browse(self):
        p = filedialog.asksaveasfilename(
            defaultextension=".sql",
            filetypes=[("SQL files", "*.sql"), ("All files", "*.*")],
            initialfile=os.path.basename(self._out_var.get()),
            initialdir=os.path.dirname(self._out_var.get()) or BASE_DIR,
            parent=self)
        if p:
            self._out_var.set(p)

    def _read_token(self):
        tok_file = self._cfg.get("github_token_file", "")
        if tok_file and os.path.exists(tok_file):
            with open(tok_file) as fh:
                return fh.read().strip()
        return ""

    def _gh_headers(self, token):
        if not token:
            return {"Accept": "application/vnd.github+json"}
        return {
            "Authorization": f"Bearer {token}",
            "Accept": "application/vnd.github+json",
            "X-GitHub-Api-Version": "2022-11-28",
        }

    def _fetch_content(self, raw_url, contents_url="", filename=""):
        """Fetch full file content at PR HEAD. Uses head SHA first to guarantee latest commit."""
        token  = self._read_token()
        hdrs   = self._gh_headers(token)
        verify = _ssl_verify(self._cfg)
        errors = []

        # 1. Contents API at PR head.sha — always the latest committed state
        head_sha  = (self._pr.get("head") or {}).get("sha", "")
        head_repo = ((self._pr.get("head") or {}).get("repo") or
                     (self._pr.get("base") or {}).get("repo") or {})
        repo_full = head_repo.get("full_name", "")
        api_base  = self._cfg.get("github_api_url", "https://api.github.com").rstrip("/")

        if head_sha and repo_full and filename:
            head_url = f"{api_base}/repos/{repo_full}/contents/{filename}?ref={head_sha}"
            try:
                r = requests.get(head_url, headers=hdrs, timeout=15, verify=verify)
                if r.ok:
                    data = r.json()
                    if isinstance(data, dict) and data.get("encoding") == "base64":
                        return base64.b64decode(data["content"]).decode("utf-8", errors="replace"), None
                    errors.append(f"Contents API (head) unexpected format")
                else:
                    errors.append(f"Contents API (head) HTTP {r.status_code}")
            except Exception as e:
                errors.append(f"Contents API (head) error: {e}")

        # 2. raw_url (raw.githubusercontent.com) — uses SHA from PR files API response
        if raw_url:
            try:
                r = requests.get(raw_url, headers=hdrs, timeout=15, verify=verify)
                if r.ok:
                    return r.text, None
                errors.append(f"raw_url HTTP {r.status_code}")
            except Exception as e:
                errors.append(str(e))
        else:
            errors.append("no raw_url")

        # 3. contents_url from PR files API (may carry older ref)
        if contents_url:
            try:
                api_url = contents_url.split("?")[0]
                r2 = requests.get(api_url, headers=hdrs, timeout=15, verify=verify)
                if r2.ok:
                    data = r2.json()
                    if isinstance(data, dict) and data.get("encoding") == "base64":
                        return base64.b64decode(data["content"]).decode("utf-8", errors="replace"), None
                    errors.append("contents_url unexpected format")
                else:
                    errors.append(f"contents_url HTTP {r2.status_code}: {r2.text[:120]}")
            except Exception as e:
                errors.append(f"contents_url error: {e}")

        return None, " | ".join(errors)

    def get_sql_comments(self):
        return getattr(self, "_sql_comments_dict", {})

    def get_keyword_findings(self):
        return getattr(self, "_keyword_findings", {})

    def _resolve_content(self, r):
        """Fetch and return file content string (or None) for a row dict."""
        if r.get("full_content"):
            return r["full_content"], None
        content, err = self._fetch_content(
            r.get("raw_url", ""), r.get("contents_url", ""), r["filename"])
        if content is None:
            patch = r.get("patch", "")
            if patch and r.get("status") == "added":
                content = "\n".join(
                    l[1:] for l in patch.splitlines()
                    if l.startswith('+') and not l.startswith('+++'))
                err = None
        return content, err

    def _generate(self):
        included = [r for r in self._rows if r["inc"].get()]
        if not included:
            messagebox.showwarning("No files", "Select at least one file.", parent=self)
            return
        out = self._out_var.get().strip()
        if not out:
            messagebox.showwarning("No path", "Set an output file path.", parent=self)
            return
        if not out.lower().endswith(".sql"):
            out += ".sql"

        self._sv.set(f"Fetching {len(included)} file(s)…")
        self.update_idletasks()

        # ── Phase 1: fetch all content ────────────────────────────────────────
        fetched = {}   # filename → (content, fetch_err)
        for r in included:
            fetched[r["filename"]] = self._resolve_content(r)

        # ── Phase 2: keyword scan ─────────────────────────────────────────────
        keywords = [kw.strip() for kw in
                    (self._kw_listbox.get(0, "end") if hasattr(self, "_kw_listbox") else load_keywords())
                    if kw.strip()]
        findings = {}   # filename → [matched keywords]
        if keywords:
            for r in included:
                fname   = r["filename"]
                content = fetched[fname][0] or ""
                patch   = r.get("patch", "")
                text    = (content + "\n" + patch).lower()
                hits    = [kw for kw in keywords if kw.lower() in text]
                if hits:
                    findings[fname] = hits

        self._keyword_findings = findings

        if findings:
            lines = ["The following files contain blocked keywords:\n"]
            for fn, kws in findings.items():
                lines.append(f"  •  {fn}")
                lines.append(f"     Keywords: {', '.join(kws)}\n")
            lines.append("Fix or remove these keywords before generating the SQL file.")
            messagebox.showerror("⚠ Keyword Alert — Generation Blocked",
                                 "\n".join(lines), parent=self)
            self._sv.set("Blocked — keyword(s) found")
            return

        # ── Phase 3: build SQL file ───────────────────────────────────────────
        self._sql_comments_dict = {
            r["filename"]: f"Step {idx}: {r['cmt_var'].get().strip() or os.path.basename(r['filename'])}"
            for idx, r in enumerate(included, 1)
        }

        pr_num   = self._pr.get("number", "")
        pr_title = self._pr.get("title", "")

        blk = []
        blk.append("-- " + "=" * 68)
        blk.append(f"-- PR #{pr_num}: {pr_title}")
        blk.append(f"-- Generated : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        blk.append(f"-- Files     : {len(included)}")
        blk.append("-- " + "=" * 68)
        blk.append("")

        for idx, r in enumerate(included, 1):
            fname   = r["filename"]
            status  = r["status"]
            adds    = r["additions"]
            dels    = r["deletions"]
            desc    = r["cmt_var"].get().strip() or os.path.basename(fname)
            comment = f"Step {idx}: {desc}"
            content, fetch_err = fetched[fname]

            blk.append("-- " + "-" * 68)
            blk.append(f"-- {comment}")
            blk.append(f"-- File  : {fname}")
            blk.append(f"-- Status: {status.upper()}  (+{adds} / -{dels})")
            blk.append("-- " + "-" * 68)
            blk.append("")

            if content:
                lines = content.rstrip().splitlines()
                while lines and lines[-1].strip() == "/":
                    lines.pop()
                blk.append("\n".join(lines).rstrip())
            else:
                blk.append(f"-- *** CONTENT UNAVAILABLE: {fname}")
                if fetch_err:
                    blk.append(f"-- *** Fetch error: {fetch_err}")
                blk.append("-- *** Fix: ensure GitHub token is set in Settings and has 'repo' scope.")

            blk.append("")
            blk.append("")

        try:
            os.makedirs(os.path.dirname(os.path.abspath(out)), exist_ok=True)
            with open(out, "w", encoding="utf-8") as fh:
                fh.write("\n".join(blk))
            self._sv.set(f"Saved: {os.path.basename(out)}")
            self._result = out
            if messagebox.askyesno("Done",
                                   f"SQL file saved:\n{out}\n\nOpen now?",
                                   parent=self):
                webbrowser.open(out)
        except Exception as ex:
            self._sv.set(f"Error: {ex}")
            messagebox.showerror("Error", str(ex), parent=self)

    def get_result(self):
        return self._result


# ─────────────────────────────────────────────────────────────────────────────
#  GitHub PR Preview Popup
# ─────────────────────────────────────────────────────────────────────────────

class GitHubPRPreviewPopup(tk.Toplevel):
    """Renders a GitHub-style PR preview: header, description, files, comments."""

    GH = dict(
        bg="#0d1117", surface="#161b22", card="#21262d",
        border="#30363d", accent="#58a6ff", text="#e6edf3",
        muted="#8b949e", green="#3fb950", red="#f85149",
        purple="#bc8cff", yellow="#d29922", orange="#e3b341",
        merged="#8957e5",
        st=dict(added="#1f6931", modified="#1c4a6e", removed="#8a1f1f",
                renamed="#5a2d82"),
    )

    def __init__(self, parent, pr, files, review_comments=None, issue_comments=None):
        super().__init__(parent)
        self.title(f"GitHub PR #{pr.get('number','?')}  —  Preview")
        self.geometry("1020x760")
        self.minsize(820, 560)
        self.configure(bg=self.GH["bg"])
        self.grab_set()
        self.transient(parent)

        self._pr      = pr
        self._files   = files or []
        self._rev_cmt = review_comments or []
        self._iss_cmt = issue_comments or []

        self._build()

    def _build(self):
        G = self.GH
        pr = self._pr

        state   = pr.get("state", "open")
        merged  = bool(pr.get("merged_at"))
        st_bg   = G["merged"] if merged else (G["green"] if state == "open" else G["muted"])
        st_lbl  = "Merged" if merged else state.capitalize()

        # ── top bar ───────────────────────────────────────────────────────────
        topbar = tk.Frame(self, bg="#161b22", height=52)
        topbar.pack(fill="x"); topbar.pack_propagate(False)
        tk.Label(topbar, text="◈  GitHub", bg="#161b22", fg=G["text"],
                 font=("Segoe UI", 11, "bold")).pack(side="left", padx=14, pady=12)
        tk.Label(topbar, text=f"{pr.get('base',{}).get('repo',{}).get('full_name','') or 'example/repo'}",
                 bg="#161b22", fg=G["muted"],
                 font=("Segoe UI", 9)).pack(side="left", pady=12)
        tk.Frame(topbar, bg=G["border"], width=1).pack(side="left", fill="y", pady=8, padx=8)
        tk.Label(topbar, text=f"Pull Request #{pr.get('number','?')}",
                 bg="#161b22", fg=G["accent"],
                 font=("Segoe UI", 9)).pack(side="left", pady=12)
        ttk.Button(topbar, text="✕  Close", style="Ghost.TButton",
                   command=self.destroy).pack(side="right", padx=14, pady=10)
        tk.Frame(self, bg=G["border"], height=1).pack(fill="x")

        # ── PR title row ──────────────────────────────────────────────────────
        title_f = tk.Frame(self, bg=G["bg"])
        title_f.pack(fill="x", padx=24, pady=(14, 6))
        tk.Label(title_f, text=pr.get("title", "(no title)"),
                 bg=G["bg"], fg=G["text"],
                 font=("Segoe UI", 15, "bold"),
                 wraplength=860, justify="left", anchor="w").pack(side="left", anchor="w")

        meta_f = tk.Frame(self, bg=G["bg"])
        meta_f.pack(fill="x", padx=24, pady=(0, 10))
        tk.Label(meta_f, text=f"  {st_lbl}  ",
                 bg=st_bg, fg="#ffffff",
                 font=("Segoe UI", 9, "bold"), padx=6, pady=3).pack(side="left")
        author = (pr.get("user") or {}).get("login", "unknown")
        base_b = (pr.get("base") or {}).get("ref", "main")
        head_b = (pr.get("head") or {}).get("ref", "?")
        tk.Label(meta_f,
                 text=f"  {author}  wants to merge  {head_b}  →  {base_b}   ·   "
                      f"opened {(pr.get('created_at') or '')[:10]}",
                 bg=G["bg"], fg=G["muted"],
                 font=("Segoe UI", 9)).pack(side="left", padx=8)
        tk.Frame(self, bg=G["border"], height=1).pack(fill="x", padx=20)

        # ── Notebook-style tabs ───────────────────────────────────────────────
        n_rev = len([c for c in self._rev_cmt])
        n_iss = len(self._iss_cmt)
        tab_bar = tk.Frame(self, bg=G["bg"])
        tab_bar.pack(fill="x", padx=20)
        self._tab_frames = {}
        self._active_tab = tk.StringVar(value="description")

        body_host = tk.Frame(self, bg=G["bg"])
        body_host.pack(fill="both", expand=True, padx=20, pady=(0, 10))

        def _show_tab(name):
            for n, f in self._tab_frames.items():
                f.pack_forget()
            self._tab_frames[name].pack(fill="both", expand=True)
            self._active_tab.set(name)
            for btn in tab_btns:
                nm, b = btn
                b.config(bg=G["accent"] if nm == name else G["bg"],
                         fg="#ffffff" if nm == name else G["muted"])

        tab_btns = []
        tabs = [
            ("description", f"📋 Description"),
            ("files",       f"📁 Files ({len(self._files)})"),
            ("comments",    f"💬 Comments ({n_rev + n_iss})"),
        ]
        for name, label in tabs:
            b = tk.Button(tab_bar, text=label,
                          bg=G["bg"], fg=G["muted"],
                          relief="flat", font=("Segoe UI", 9, "bold"),
                          padx=14, pady=7, cursor="hand2",
                          activebackground=G["card"],
                          command=lambda n=name: _show_tab(n))
            b.pack(side="left")
            tab_btns.append((name, b))
            self._tab_frames[name] = self._make_scrollable(body_host)

        tk.Frame(self, bg=G["border"], height=1).pack(fill="x", padx=20)

        self._build_description(self._tab_frames["description"])
        self._build_files(self._tab_frames["files"])
        self._build_comments(self._tab_frames["comments"])
        _show_tab("description")

    def _make_scrollable(self, parent):
        outer = tk.Frame(parent, bg=self.GH["bg"])
        canvas = tk.Canvas(outer, bg=self.GH["bg"], highlightthickness=0)
        vsb = ttk.Scrollbar(outer, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)
        inner = tk.Frame(canvas, bg=self.GH["bg"])
        wid = canvas.create_window((0, 0), window=inner, anchor="nw")
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(wid, width=e.width))
        inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
        outer._inner = inner
        return outer

    def _build_description(self, host):
        G = self.GH
        f = host._inner
        pr_body = (self._pr.get("body") or "").strip() or "(No description provided)"

        card = tk.Frame(f, bg=G["surface"],
                        highlightbackground=G["border"], highlightthickness=1)
        card.pack(fill="x", padx=4, pady=8)

        hdr = tk.Frame(card, bg=G["card"])
        hdr.pack(fill="x")
        author = (self._pr.get("user") or {}).get("login", "author")
        tk.Label(hdr, text=f"  {author}  commented  ·  {(self._pr.get('created_at') or '')[:10]}",
                 bg=G["card"], fg=G["muted"],
                 font=("Segoe UI", 9), pady=7, padx=8).pack(side="left")
        tk.Frame(card, bg=G["border"], height=1).pack(fill="x")

        body_f = tk.Frame(card, bg=G["surface"])
        body_f.pack(fill="x", padx=16, pady=12)

        for line in pr_body.split("\n"):
            stripped = line.strip()
            if stripped.startswith("## "):
                tk.Label(body_f, text=stripped[3:],
                         bg=G["surface"], fg=G["text"],
                         font=("Segoe UI", 11, "bold"), anchor="w").pack(fill="x", pady=(8,2))
            elif stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
                done = stripped.startswith("- [x] ")
                txt  = stripped[6:]
                row  = tk.Frame(body_f, bg=G["surface"]); row.pack(fill="x")
                tk.Label(row, text="☑" if done else "☐",
                         bg=G["surface"], fg=G["green"] if done else G["muted"],
                         font=("Segoe UI", 10)).pack(side="left", padx=(4,6))
                tk.Label(row, text=txt,
                         bg=G["surface"], fg=G["muted"] if done else G["text"],
                         font=("Segoe UI", 9), anchor="w").pack(side="left", fill="x")
            elif stripped.startswith("- "):
                row = tk.Frame(body_f, bg=G["surface"]); row.pack(fill="x")
                tk.Label(row, text="•",
                         bg=G["surface"], fg=G["accent"],
                         font=("Segoe UI", 9)).pack(side="left", padx=(4,6))
                tk.Label(row, text=stripped[2:],
                         bg=G["surface"], fg=G["text"],
                         font=("Segoe UI", 9), anchor="w",
                         wraplength=760, justify="left").pack(side="left", fill="x")
            elif stripped:
                tk.Label(body_f, text=stripped,
                         bg=G["surface"], fg=G["text"],
                         font=("Segoe UI", 9), anchor="w",
                         wraplength=800, justify="left").pack(fill="x")

    def _build_files(self, host):
        G  = self.GH
        f  = host._inner
        st_colors = G["st"]

        stats_row = tk.Frame(f, bg=G["bg"])
        stats_row.pack(fill="x", padx=4, pady=(8, 4))
        total = len(self._files)
        adds  = sum(f_.get("additions", 0) for f_ in self._files)
        dels  = sum(f_.get("deletions", 0) for f_ in self._files)
        tk.Label(stats_row,
                 text=f"{total} files changed   +{adds} additions   −{dels} deletions",
                 bg=G["bg"], fg=G["muted"],
                 font=("Segoe UI", 9)).pack(side="left")

        # group review comments by file
        rev_by = {}
        for c in self._rev_cmt:
            path = c.get("path","")
            rev_by.setdefault(path, []).append(c)

        for file_ in self._files:
            fname  = file_.get("filename","")
            status = file_.get("status","modified")
            a      = file_.get("additions",0)
            d      = file_.get("deletions",0)
            fbg    = st_colors.get(status, G["card"])
            em     = STATUS_EMOJI.get(status,"~")

            card = tk.Frame(f, bg=G["surface"],
                            highlightbackground=G["border"], highlightthickness=1)
            card.pack(fill="x", padx=4, pady=3)

            lborder = tk.Frame(card, bg=fbg, width=4)
            lborder.pack(side="left", fill="y")

            inner = tk.Frame(card, bg=G["surface"])
            inner.pack(side="left", fill="x", expand=True, padx=10, pady=7)

            hrow = tk.Frame(inner, bg=G["surface"])
            hrow.pack(fill="x")
            tk.Label(hrow, text=f"{em}  {status.upper()}",
                     bg=fbg, fg="#ffffff",
                     font=("Segoe UI", 8, "bold"), padx=6, pady=2).pack(side="left")
            tk.Label(hrow, text=f"  {fname}",
                     bg=G["surface"], fg=G["accent"],
                     font=("Consolas", 9, "bold"), anchor="w").pack(side="left")
            sfg = G["green"] if a >= d else G["red"]
            tk.Label(hrow, text=f"  +{a}  −{d}",
                     bg=G["surface"], fg=sfg,
                     font=("Segoe UI", 9, "bold")).pack(side="right")

            # review comments for this file
            file_revs = rev_by.get(fname, [])
            for rc in file_revs:
                rc_user = (rc.get("user") or {}).get("login","")
                rc_body = rc.get("body","").strip()
                rc_date = (rc.get("created_at") or "")[:10]
                cmt_f = tk.Frame(inner, bg=G["card"],
                                 highlightbackground=G["border"], highlightthickness=1)
                cmt_f.pack(fill="x", pady=(4,0))
                tk.Label(cmt_f,
                         text=f"  💬  {rc_user}  ·  {rc_date}",
                         bg=G["card"], fg=G["muted"],
                         font=("Segoe UI", 8), pady=4, padx=6).pack(anchor="w")
                tk.Frame(cmt_f, bg=G["border"], height=1).pack(fill="x")
                tk.Label(cmt_f, text=f"    {rc_body}",
                         bg=G["card"], fg=G["text"],
                         font=("Segoe UI", 9), anchor="w", padx=6, pady=5,
                         wraplength=820, justify="left").pack(anchor="w")

    def _build_comments(self, host):
        G = self.GH
        f = host._inner
        all_comments = []
        for c in self._iss_cmt:
            all_comments.append(("issue", c))
        for c in self._rev_cmt:
            all_comments.append(("review", c))
        all_comments.sort(key=lambda x: x[1].get("created_at",""))

        if not all_comments:
            tk.Label(f, text="No comments on this PR.",
                     bg=G["bg"], fg=G["muted"],
                     font=("Segoe UI", 10)).pack(padx=20, pady=20)
            return

        for kind, c in all_comments:
            user = (c.get("user") or {}).get("login","")
            body = c.get("body","").strip()
            date = (c.get("created_at") or "")[:10]
            path = c.get("path","")

            card = tk.Frame(f, bg=G["surface"],
                            highlightbackground=G["border"], highlightthickness=1)
            card.pack(fill="x", padx=4, pady=4)

            hdr = tk.Frame(card, bg=G["card"])
            hdr.pack(fill="x")
            type_lbl = f"📝 Review — {path}" if kind == "review" and path else "💬 Comment"
            tk.Label(hdr, text=f"  {user}  ·  {date}  ·  {type_lbl}",
                     bg=G["card"], fg=G["muted"],
                     font=("Segoe UI", 8), pady=6, padx=8).pack(side="left")
            tk.Frame(card, bg=G["border"], height=1).pack(fill="x")
            tk.Label(card, text=body,
                     bg=G["surface"], fg=G["text"],
                     font=("Segoe UI", 9), anchor="w", padx=14, pady=8,
                     wraplength=860, justify="left").pack(anchor="w")


# ─────────────────────────────────────────────────────────────────────────────
#  Main Application
# ─────────────────────────────────────────────────────────────────────────────

class App(tk.Tk):

    def __init__(self):
        super().__init__()
        self.config_data    = load_config()
        self.pr_cache       = None
        self.pr_files       = []
        self.file_comments  = []
        self._jira_url        = ""
        self._last_doc_path   = None
        self._last_sql_path   = None
        self._last_issue_key  = None
        self._field_widgets = {}
        self._theme_name    = self.config_data.get("theme", "dark")
        C.update(THEMES[self._theme_name])
        STATUS_BADGE.clear(); STATUS_BADGE.update(C["st_badge"])
        STATUS_ROW.clear();   STATUS_ROW.update(C["st_row"])

        self.title(APP_NAME)
        self.geometry("1140x960")
        self.minsize(920, 740)
        self.configure(bg=C["bg"])
        self._apply_theme()
        self._load_app_icon()
        self._build_ui()

    # ── icon ──────────────────────────────────────────────────────────────────

    def _load_app_icon(self):
        icon_path = os.path.join(BASE_DIR, "prism_icon.ico")

        # Regenerate if missing or zero-byte
        if not os.path.exists(icon_path) or os.path.getsize(icon_path) == 0:
            _, icon_path2 = _generate_app_assets()
            if icon_path2:
                icon_path = icon_path2

        if os.path.exists(icon_path) and os.path.getsize(icon_path) > 0:
            try:
                # default= propagates to taskbar AND all child windows on Windows
                self.iconbitmap(default=icon_path)
                return
            except Exception:
                pass

        # Fallback: iconphoto via PIL (non-Windows or .ico unavailable)
        try:
            from PIL import Image, ImageTk
            logo = os.path.join(BASE_DIR, "prism_logo.png")
            src  = logo if os.path.exists(logo) else icon_path
            if os.path.exists(src):
                im = Image.open(src).resize((64, 64), Image.LANCZOS)
                self._app_photo = ImageTk.PhotoImage(im)
                self.iconphoto(True, self._app_photo)
        except Exception:
            pass

    # ── theme ─────────────────────────────────────────────────────────────────

    def _apply_theme(self):
        s = ttk.Style(self)
        s.theme_use("clam")
        s.configure(".",
                    background=C["bg"], foreground=C["text"],
                    fieldbackground=C["inp"], bordercolor=C["border"],
                    darkcolor=C["surface"], lightcolor=C["surface"],
                    troughcolor=C["surface"],
                    selectbackground=C["hover"], selectforeground=C["text"],
                    font=("Segoe UI", 10))
        s.configure("TFrame",  background=C["bg"])
        s.configure("TLabel",  background=C["bg"], foreground=C["text"])
        s.configure("TEntry",
                    fieldbackground=C["inp"], foreground=C["text"],
                    insertcolor=C["accent"], bordercolor=C["border"],
                    padding=(8, 6), relief="flat")
        s.map("TEntry",
              fieldbackground=[("focus", C["inp"])],
              bordercolor=[("focus", C["accent"])])
        s.configure("TCombobox",
                    fieldbackground=C["inp"], foreground=C["text"],
                    selectbackground=C["accent_dim"], arrowcolor=C["accent"],
                    padding=(7, 5), relief="flat")
        s.map("TCombobox",
              fieldbackground=[("readonly", C["inp"]), ("focus", C["inp"])],
              foreground=[("readonly", C["text"])],
              bordercolor=[("focus", C["accent"])])
        s.configure("TScrollbar",
                    background=C["border"], troughcolor=C["bg"],
                    bordercolor=C["bg"], arrowcolor=C["bg"],
                    arrowsize=8, relief="flat", width=8)
        s.map("TScrollbar",
              background=[("active", C["accent"]), ("pressed", C["hover"])])

        for name, bg, fg, active in [
            ("Accent",   C["accent"],  "#ffffff", C["hover"]),
            ("Success",  C["green"],   "#ffffff", "#2ea043"),
            ("Purple",   C["purple"],  "#ffffff", "#a070e0"),
            ("Orange",   C["orange"],  "#ffffff", "#c9952e"),
            ("Ghost",    C["card2"],   C["muted"], C["border"]),
            ("Danger",   C["red"],     "#ffffff", "#d03a30"),
            ("Subtle",   C["accent_dim"], C["accent"], C["border"]),
        ]:
            s.configure(f"{name}.TButton",
                        background=bg, foreground=fg,
                        padding=(16, 8), relief="flat",
                        font=("Segoe UI", 10, "bold"),
                        borderwidth=0)
            s.map(f"{name}.TButton",
                  background=[("active", active), ("pressed", active)],
                  foreground=[("active", fg)])

        s.configure("TNotebook", background=C["bg"], borderwidth=0)
        s.configure("TNotebook.Tab",
                    background=C["surface"], foreground=C["muted"],
                    padding=[18, 8], font=("Segoe UI", 10, "bold"),
                    borderwidth=0)
        s.map("TNotebook.Tab",
              background=[("selected", C["bg"])],
              foreground=[("selected", C["accent"])],
              expand=[("selected", [0, 0, 0, 2])])
        s.configure("TProgressbar",
                    troughcolor=C["border"],
                    background=C["accent"],
                    bordercolor=C["border"],
                    thickness=4)

    # ── theme toggle ──────────────────────────────────────────────────────────

    def _toggle_theme(self):
        state = self._save_app_state()
        self._theme_name = "light" if self._theme_name == "dark" else "dark"
        self.config_data["theme"] = self._theme_name
        save_config(self.config_data)
        C.update(THEMES[self._theme_name])
        STATUS_BADGE.clear(); STATUS_BADGE.update(C["st_badge"])
        STATUS_ROW.clear();   STATUS_ROW.update(C["st_row"])
        for w in self.winfo_children():
            w.destroy()
        self.configure(bg=C["bg"])
        self._apply_theme()
        self._build_ui()
        self._restore_app_state(state)

    def _save_app_state(self):
        state = {
            "pr_url":        getattr(self, "pr_url_var", tk.StringVar()).get(),
            "pr_cache":      self.pr_cache,
            "pr_files":      self.pr_files,
            "file_comments": self.file_comments,
            "_jira_url":     self._jira_url,
            "_last_doc_path":self._last_doc_path,
            "field_values":  {k: v.get() for k, v in self._field_widgets.items()},
            "description":   "",
        }
        if hasattr(self, "desc_text"):
            d = self.desc_text.get("1.0", "end-1c")
            if d != "Describe the issue in detail...":
                state["description"] = d
        return state

    def _restore_app_state(self, state):
        if state.get("pr_url") and hasattr(self, "pr_url_var"):
            self.pr_url_var.set(state["pr_url"])
        self.pr_cache       = state.get("pr_cache")
        self.pr_files       = state.get("pr_files", [])
        self.file_comments  = state.get("file_comments", [])
        self._jira_url      = state.get("_jira_url", "")
        self._last_doc_path = state.get("_last_doc_path")
        if self.pr_cache and hasattr(self, "pr_info_var"):
            pr = self.pr_cache
            self.pr_info_var.set(
                f"  #{pr['number']}  {pr['title']}   "
                f"{pr['state'].upper()}   by {pr['user']['login']}")
        if self.pr_files and hasattr(self, "sql_btn"):
            self.sql_btn.config(state="normal")
        for k, v in state.get("field_values", {}).items():
            if k in self._field_widgets:
                self._field_widgets[k].set(v)
        if state.get("description") and hasattr(self, "desc_text"):
            self.desc_text.delete("1.0", "end")
            self.desc_text.insert("1.0", state["description"])
            self.desc_text.config(fg=C["text"])

    # ── UI root ───────────────────────────────────────────────────────────────

    def _build_ui(self):
        # ── top accent stripe ─────────────────────────────────────────────────
        top_stripe = tk.Frame(self, bg=C["accent"], height=3)
        top_stripe.pack(fill="x")
        # subtle multi-color right edge of stripe
        stripe_r = tk.Frame(top_stripe, height=3, bg=C["purple"])
        stripe_r.place(relx=0.7, rely=0, relwidth=0.15, relheight=1)
        stripe_r2 = tk.Frame(top_stripe, height=3, bg=C["green"])
        stripe_r2.place(relx=0.85, rely=0, relwidth=0.15, relheight=1)

        # ── header bar ────────────────────────────────────────────────────────
        hdr = tk.Frame(self, bg=C["surface"])
        hdr.pack(fill="x")

        # logo / title
        title_f = tk.Frame(hdr, bg=C["surface"])
        title_f.pack(side="left", padx=(14, 0), pady=9)

        # canvas-drawn prism icon
        ico = tk.Canvas(title_f, width=42, height=42, bg=C["surface"],
                        highlightthickness=0)
        ico.pack(side="left", padx=(0, 10))
        ico.create_oval(2, 2, 40, 40, fill=C["card"], outline=C["border"], width=1)
        pts = [11, 7,  11, 35,  36, 21]
        ico.create_polygon(pts, fill=C["accent"], outline="#aad4ff", width=1)
        for dy, col in [(-9, C["green"]), (0, C["accent"]), (9, C["yellow"])]:
            ico.create_line(36, 21, 41, 21 + dy, fill=col, width=2)

        # "PRism" — "PR" in accent, "ism" in text
        name_f = tk.Frame(title_f, bg=C["surface"])
        name_f.pack(side="left")
        tk.Label(name_f, text="PR",
                 bg=C["surface"], fg=C["accent"],
                 font=("Segoe UI", 17, "bold")).pack(side="left")
        tk.Label(name_f, text="ism",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 17, "bold")).pack(side="left")
        tk.Label(name_f, text=f"  —  {APP_SUBTITLE}",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 9)).pack(side="left", pady=(5, 0))

        # right-side controls
        ctrl_f = tk.Frame(hdr, bg=C["surface"])
        ctrl_f.pack(side="right", padx=14, pady=10)

        theme_lbl = "  Light" if self._theme_name == "dark" else "  Dark"
        theme_ico = "☀" if self._theme_name == "dark" else "🌙"
        tk.Button(ctrl_f, text=f"{theme_ico}{theme_lbl}",
                  bg=C["card2"], fg=C["text"],
                  relief="flat", font=("Segoe UI", 9, "bold"),
                  padx=10, pady=6, cursor="hand2",
                  activebackground=C["border"], activeforeground=C["text"],
                  command=self._toggle_theme).pack(side="right", padx=(6, 0))

        for txt, cmd, bg, fg, abg in [
            ("⊞  Jira Fields",  self._open_jira_fields_mgr, C["accent_dim"], C["accent"],  C["border"]),
            ("⊞  Desc Info",    self._open_desc_info_mgr,   C["card2"],      C["purple"],  C["border"]),
            ("⚙  Settings",     self._open_settings,        C["card2"],      C["muted"],   C["border"]),
        ]:
            tk.Button(ctrl_f, text=txt, bg=bg, fg=fg, relief="flat",
                      font=("Segoe UI", 9, "bold"), padx=10, pady=6,
                      activebackground=abg, activeforeground=C["text"],
                      cursor="hand2", command=cmd).pack(side="right", padx=3)

        tk.Frame(self, bg=C["border"], height=1).pack(fill="x")

        # ── main content ──────────────────────────────────────────────────────
        self.tab_create = ttk.Frame(self)
        self.tab_create.pack(fill="both", expand=True)

        self._build_create_tab()

        # ── status bar ────────────────────────────────────────────────────────
        self.status_var = tk.StringVar(value="Ready  —  fetch a PR to begin")
        sbar = tk.Frame(self, bg=C["surface"])
        sbar.pack(fill="x", side="bottom")
        tk.Frame(sbar, bg=C["border"], height=1).pack(fill="x")
        sbar_inner = tk.Frame(sbar, bg=C["surface"])
        sbar_inner.pack(fill="x")
        self._status_dot = tk.Label(sbar_inner, text="●",
                 bg=C["surface"], fg=C["green"],
                 font=("Segoe UI", 8))
        self._status_dot.pack(side="left", padx=(12, 4), pady=6)
        tk.Label(sbar_inner, textvariable=self.status_var,
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 9)).pack(side="left")
        tk.Frame(sbar_inner, bg=C["border"], width=1).pack(side="right", fill="y", pady=3)
        tk.Label(sbar_inner, text=f"PRism  •  {self._theme_name.capitalize()}",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(side="right", padx=14, pady=6)

    # ── scrollable frame helper ───────────────────────────────────────────────

    def _scrollable(self, parent):
        c = tk.Canvas(parent, bg=C["bg"], highlightthickness=0)
        v = ttk.Scrollbar(parent, orient="vertical", command=c.yview)
        c.configure(yscrollcommand=v.set)
        v.pack(side="right", fill="y")
        c.pack(side="left", fill="both", expand=True)
        f = tk.Frame(c, bg=C["bg"])
        w = c.create_window((0, 0), window=f, anchor="nw")
        c.bind("<Configure>", lambda e: c.itemconfig(w, width=e.width))
        f.bind("<Configure>", lambda e: c.configure(scrollregion=(0, 0, e.width, e.height)))
        c.bind_all("<MouseWheel>",
                   lambda e: c.yview_scroll(int(-1*(e.delta/120)), "units"))
        return f

    def _card(self, parent, title="", icon="", accent_color=None):
        accent_color = accent_color or C["accent"]
        # outer shadow frame for depth effect
        shadow_f = tk.Frame(parent, bg=C["shadow"])
        shadow_f.pack(fill="both", expand=True, padx=(10, 8), pady=(0, 10))
        border_f = tk.Frame(shadow_f, bg=C["border"])
        border_f.pack(fill="both", expand=True, padx=0, pady=0)
        inner_wrap = tk.Frame(border_f, bg=C["card"])
        inner_wrap.pack(fill="both", expand=True, padx=1, pady=1)
        accent_strip = tk.Frame(inner_wrap, bg=accent_color, width=4)
        accent_strip.pack(side="left", fill="y")
        content = tk.Frame(inner_wrap, bg=C["card"])
        content.pack(side="left", fill="both", expand=True, padx=14, pady=12)
        if title:
            th = tk.Frame(content, bg=C["card"])
            th.pack(fill="x", pady=(0, 6))
            if icon:
                tk.Label(th, text=icon, bg=C["card"], fg=accent_color,
                         font=("Segoe UI", 12)).pack(side="left", padx=(0, 7))
            tk.Label(th, text=title, bg=C["card"], fg=C["text"],
                     font=("Segoe UI", 11, "bold")).pack(side="left")
            tk.Frame(content, bg=C["border"], height=1).pack(fill="x", pady=(0, 10))
        return content

    # ── Create Ticket tab ────────────────────────────────────────────────────

    def _build_create_tab(self):
        outer = tk.Frame(self.tab_create, bg=C["bg"])
        outer.pack(fill="both", expand=True)

        # ── Right sidebar: actions ────────────────────────────────────────────
        sidebar = tk.Frame(outer, bg=C["surface"], width=210)
        sidebar.pack(side="right", fill="y")
        sidebar.pack_propagate(False)
        tk.Frame(sidebar, bg=C["border"], width=1).pack(side="left", fill="y")
        rp = tk.Frame(sidebar, bg=C["surface"])
        rp.pack(side="left", fill="both", expand=True, padx=14, pady=16)

        # section header
        sh_f = tk.Frame(rp, bg=C["surface"])
        sh_f.pack(fill="x", pady=(0, 12))
        tk.Frame(sh_f, bg=C["accent"], width=3).pack(side="left", fill="y", padx=(0, 7))
        tk.Label(sh_f, text="ACTIONS", bg=C["surface"], fg=C["accent"],
                 font=("Segoe UI", 8, "bold")).pack(side="left", anchor="w")

        prev_btn = ttk.Button(rp, text="◈  Preview Doc", style="Orange.TButton",
                              command=self._show_preview)
        prev_btn.pack(fill="x", pady=(0, 2))
        Tooltip(prev_btn, "Add change comments and generate the Word document.  •  Ctrl+Shift+D")
        tk.Label(rp, text="Generate Word doc  •  Ctrl+Shift+D", bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 7)).pack(anchor="w", padx=2, pady=(0, 8))

        self.sql_btn = ttk.Button(rp, text="⊞  SQL PR File", style="Accent.TButton",
                                  state="disabled", command=self._open_sql_popup)
        self.sql_btn.pack(fill="x", pady=(0, 2))
        Tooltip(self.sql_btn, "Combine all PR SQL files into a single runnable script.")
        tk.Label(rp, text="Combine SQL scripts", bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 7)).pack(anchor="w", padx=2, pady=(0, 8))

        # ── Attach options ────────────────────────────────────────────────────
        self._attach_word_var = tk.BooleanVar(
            value=bool(self.config_data.get("attach_word_doc", True)))
        self._attach_sql_var  = tk.BooleanVar(
            value=bool(self.config_data.get("attach_sql_file", True)))

        def _save_attach_prefs(*_):
            self.config_data["attach_word_doc"] = self._attach_word_var.get()
            self.config_data["attach_sql_file"] = self._attach_sql_var.get()
            save_config(self.config_data)

        att_frame = tk.Frame(rp, bg=C["surface"])
        att_frame.pack(fill="x", pady=(0, 4))
        tk.Label(att_frame, text="Attach to ticket:", bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 7)).pack(anchor="w", padx=2)
        att_row = tk.Frame(att_frame, bg=C["surface"]); att_row.pack(fill="x", padx=2)
        ttk.Checkbutton(att_row, text="Word Doc",
                        variable=self._attach_word_var,
                        command=_save_attach_prefs).pack(side="left")
        ttk.Checkbutton(att_row, text="SQL File",
                        variable=self._attach_sql_var,
                        command=_save_attach_prefs).pack(side="left", padx=(8, 0))

        jira_btn = ttk.Button(rp, text="✦  Create Jira Ticket", style="Success.TButton",
                              command=self._create_jira_thread)
        jira_btn.pack(fill="x", pady=(0, 2))
        Tooltip(jira_btn, "Create Jira ticket and attach Word doc + SQL file.  •  Ctrl+Shift+J")
        tk.Label(rp, text="Post to Jira  •  Ctrl+Shift+J", bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 7)).pack(anchor="w", padx=2, pady=(0, 6))

        jira_prev_btn = ttk.Button(rp, text="◈  Jira Preview", style="Accent.TButton",
                                   command=self._show_jira_preview)
        jira_prev_btn.pack(fill="x", pady=(0, 2))
        Tooltip(jira_prev_btn, "Preview the Jira ticket before creating — validate fields and content.  •  Ctrl+Shift+P")
        tk.Label(rp, text="Validate before submitting  •  Ctrl+Shift+P", bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 7)).pack(anchor="w", padx=2, pady=(0, 14))

        tk.Frame(rp, bg=C["border"], height=1).pack(fill="x", pady=(0, 12))

        self.progress = ttk.Progressbar(rp, mode="indeterminate")
        self.progress.pack(fill="x", pady=(0, 8))

        self.result_var = tk.StringVar()
        result_card = tk.Frame(rp, bg=C["card"],
                               highlightthickness=1,
                               highlightbackground=C["border"])
        result_card.pack(fill="x", pady=(4, 0))
        result_inner = tk.Frame(result_card, bg=C["card"])
        result_inner.pack(fill="x", padx=8, pady=6)
        rl = tk.Label(result_inner, textvariable=self.result_var,
                      bg=C["card"], fg=C["green"],
                      font=("Segoe UI", 8, "bold"), wraplength=170,
                      justify="left", cursor="hand2")
        rl.pack(anchor="w")
        rl.bind("<Button-1>", lambda e: self._jira_url and webbrowser.open(self._jira_url))

        self._attach_btn = ttk.Button(rp, text="⊞  Attach Files to Ticket",
                                      style="Accent.TButton", state="disabled",
                                      command=self._attach_files_thread)
        self._attach_btn.pack(fill="x", pady=(6, 0))
        Tooltip(self._attach_btn,
                "Attach Word doc and/or SQL file to the last created Jira ticket.")

        # ── Left content pane ─────────────────────────────────────────────────
        lp = tk.Frame(outer, bg=C["bg"])
        lp.pack(side="left", fill="both", expand=True)

        def _sep():
            tk.Frame(lp, bg=C["border"], height=1).pack(fill="x", padx=14, pady=(6, 0))
            tk.Frame(lp, bg=C["bg"], height=4).pack(fill="x")

        _SH_ICONS = {
            "GITHUB": "⎇", "DESCRIPTION": "✐", "JIRA": "⊡",
            "ISSUE": "≋", "SQL": "⊞", "CODE": "◈",
        }
        def _sh(parent, text, accent=None):
            clr = accent or C["accent"]
            icon = next((v for k, v in _SH_ICONS.items() if k in text), "")
            f = tk.Frame(parent, bg=C["bg"])
            f.pack(fill="x", padx=10, pady=(10, 3))
            tk.Frame(f, bg=clr, width=3).pack(side="left", fill="y",
                                               padx=(2, 8))
            if icon:
                tk.Label(f, text=icon, bg=C["bg"], fg=clr,
                         font=("Segoe UI", 9, "bold")).pack(side="left",
                                                            padx=(0, 5))
            tk.Label(f, text=text, bg=C["bg"], fg=clr,
                     font=("Segoe UI", 8, "bold")).pack(side="left", anchor="w")

        _frow_count = [0]
        def _frow(parent, fd):
            key = fd["key"]; label = fd["label"]
            ftype = fd.get("type","text"); choices = fd.get("choices",[])
            req = fd.get("required", False)
            if key == "project":
                default = self.config_data.get("jira_project_key", fd.get("default",""))
            else:
                default = fd.get("default","")
            _frow_count[0] += 1
            row_bg = C["card"] if _frow_count[0] % 2 == 0 else C["bg"]
            fr = tk.Frame(parent, bg=row_bg)
            fr.pack(fill="x", padx=10, pady=0)
            pad_f = tk.Frame(fr, bg=row_bg)
            pad_f.pack(fill="x", padx=4, pady=(2, 2))
            lf = tk.Frame(pad_f, bg=row_bg, width=132)
            lf.pack(side="left", fill="y")
            lf.pack_propagate(False)
            lbl_txt = label
            lbl_fg  = C["text"] if req else C["muted"]
            lbl_w = tk.Label(lf, text=lbl_txt,
                             bg=row_bg, fg=lbl_fg,
                             font=("Segoe UI", 9, "bold" if req else "normal"),
                             anchor="w")
            lbl_w.pack(side="left", anchor="w", pady=4)
            if req:
                tk.Label(lf, text=" ●", bg=row_bg, fg=C["accent"],
                         font=("Segoe UI", 7)).pack(side="left", anchor="w")
            hint = FIELD_HINTS.get(key, "")
            if hint:
                Tooltip(lbl_w, hint)
            var = tk.StringVar(value=default)
            if ftype == "dropdown" and choices:
                w = ttk.Combobox(pad_f, textvariable=var, values=choices,
                                 state="readonly")
            else:
                w = ttk.Entry(pad_f, textvariable=var)
            w.pack(side="left", fill="x", expand=True)
            if key == "doc_name":
                tk.Button(pad_f, text="⊞", bg=C["purple"], fg="#ffffff",
                          relief="flat", font=("Segoe UI", 9, "bold"), padx=7, pady=4,
                          activebackground="#a060cc", cursor="hand2",
                          command=self._choose_doc_dir).pack(side="left", padx=(4, 0))
            if ftype == "dropdown":
                def _edit(fd=fd, cb=w, v=var):
                    self._edit_dropdown_choices(fd, cb, v)
                tk.Button(pad_f, text="✎", bg=row_bg, fg=C["accent"], relief="flat",
                          font=("Segoe UI", 9), padx=5, pady=4,
                          activebackground=C["border"],
                          cursor="hand2", command=_edit).pack(side="left", padx=(3, 0))
            return var

        # ── PR URL ────────────────────────────────────────────────────────────
        _sh(lp, "GITHUB PULL REQUEST")
        pr_row = tk.Frame(lp, bg=C["bg"]); pr_row.pack(fill="x", padx=12, pady=(0, 3))
        lf = tk.Frame(pr_row, bg=C["bg"], width=115); lf.pack(side="left", fill="y")
        lf.pack_propagate(False)
        lbl = tk.Label(lf, text="PR URL / Number", bg=C["bg"], fg=C["muted"],
                       font=("Segoe UI", 9), anchor="w")
        lbl.pack(anchor="w", pady=4)
        Tooltip(lbl, "Paste full GitHub PR URL or just the number (e.g. 42)")
        self.pr_url_var = tk.StringVar()
        ttk.Entry(pr_row, textvariable=self.pr_url_var).pack(
            side="left", fill="x", expand=True, padx=(0, 6))
        ttk.Button(pr_row, text="↓  Fetch PR", style="Accent.TButton",
                   command=self._fetch_pr_thread).pack(side="left", padx=(0, 4))
        ttk.Button(pr_row, text="⚗  Test Data", style="Ghost.TButton",
                   command=self._load_mock_data).pack(side="left")

        info_outer = tk.Frame(lp, bg=C["bg"])
        info_outer.pack(fill="x", padx=12, pady=(3, 2))
        info_f = tk.Frame(info_outer, bg=C["card"],
                          highlightthickness=1,
                          highlightbackground=C["border"])
        info_f.pack(fill="x")
        info_inner = tk.Frame(info_f, bg=C["card"])
        info_inner.pack(fill="x", padx=10, pady=6)
        self.pr_info_var = tk.StringVar(
            value="No PR loaded — enter a PR URL or number above")
        tk.Label(info_inner, textvariable=self.pr_info_var,
                 bg=C["card"], fg=C["muted"],
                 font=("Segoe UI", 8), anchor="w").pack(anchor="w")
        self.files_sum_var = tk.StringVar(value="")
        tk.Label(info_inner, textvariable=self.files_sum_var,
                 bg=C["card"], fg=C["green"],
                 font=("Segoe UI", 8, "bold"), anchor="w").pack(anchor="w")

        _sep()

        # ── Description (compact) ─────────────────────────────────────────────
        _sh(lp, "DESCRIPTION")
        self.desc_text = scrolledtext.ScrolledText(
            lp, height=4, bg=C["inp"], fg=C["muted"],
            insertbackground=C["text"], relief="flat",
            font=("Segoe UI", 9), wrap="word", selectbackground=C["hover"],
            highlightthickness=1, highlightbackground=C["border"],
            highlightcolor=C["accent"], padx=8, pady=6)
        self.desc_text.pack(fill="x", padx=12, pady=(0, 4))
        _PLACEHOLDER = "Describe the issue in detail..."
        self.desc_text.insert("1.0", _PLACEHOLDER)
        def _clr(e):
            if self.desc_text.get("1.0", "end-1c") == _PLACEHOLDER:
                self.desc_text.delete("1.0", "end")
                self.desc_text.config(fg=C["text"],
                                      highlightbackground=C["accent"])
        def _restore(e):
            if not self.desc_text.get("1.0", "end-1c").strip():
                self.desc_text.insert("1.0", _PLACEHOLDER)
                self.desc_text.config(fg=C["muted"],
                                      highlightbackground=C["border"])
        self.desc_text.bind("<FocusIn>", _clr)
        self.desc_text.bind("<FocusOut>", _restore)

        _sep()

        # ── Two columns: Jira Fields | Issue Details ──────────────────────────
        mid = tk.Frame(lp, bg=C["bg"]); mid.pack(fill="x", expand=False)
        jf_col = tk.Frame(mid, bg=C["bg"]); jf_col.pack(side="left", fill="both", expand=True)
        tk.Frame(mid, bg=C["border"], width=1).pack(side="left", fill="y", pady=6)
        id_col = tk.Frame(mid, bg=C["bg"]); id_col.pack(side="left", fill="both", expand=True)

        # Jira Fields — section header stays fixed; rows are rebuildable
        _sh(jf_col, "JIRA FIELDS", accent=C["accent"])
        self._jira_fields_parent = jf_col
        self._jira_fields_card   = None
        self._rebuild_jira_fields_section()

        # Issue Details
        _frow_count[0] = 0
        _sh(id_col, "ISSUE DETAILS", accent=C["purple"])
        self._fields_card_parent = id_col
        self._fields_card = None
        self._rebuild_fields_card()

        # ── Keyboard shortcuts ────────────────────────────────────────────────
        self.bind("<Control-Return>",       lambda e: self._fetch_pr_thread())
        self.bind("<Control-Shift-J>",      lambda e: self._create_jira_thread())
        self.bind("<Control-Shift-D>",      lambda e: self._show_preview())
        self.bind("<Control-Shift-P>",      lambda e: self._show_jira_preview())

    # PR fetch card
    def _build_pr_card(self, parent):
        card = self._card(parent, "GitHub Pull Request", "")

        row = tk.Frame(card, bg=C["card"]); row.pack(fill="x", pady=(0, 6))
        lf = tk.Frame(row, bg=C["card"], width=148); lf.pack(side="left")
        lf.pack_propagate(False)
        lbl = tk.Label(lf, text="PR URL or Number",
                       bg=C["card"], fg=C["text"],
                       font=("Segoe UI", 10))
        lbl.pack(anchor="w", pady=4)
        Tooltip(lbl, "Paste the full GitHub PR URL  or  just the PR number (e.g. 42).\n"
                     "GitHub Owner and Repo must be set in Settings if using a number.")
        self.pr_url_var = tk.StringVar()
        ttk.Entry(row, textvariable=self.pr_url_var).pack(side="left", fill="x", expand=True)
        ttk.Button(row, text="  Fetch PR  ", style="Accent.TButton",
                   command=self._fetch_pr_thread).pack(side="left", padx=(8, 0))
        ttk.Button(row, text="⚗ Test Data", style="Orange.TButton",
                   command=self._load_mock_data).pack(side="left", padx=(4, 0))
        self._gh_preview_btn = ttk.Button(row, text="⊞ PR Preview", style="Ghost.TButton",
                   state="disabled", command=self._show_gh_preview)
        self._gh_preview_btn.pack(side="left", padx=(4, 0))

        info_f = tk.Frame(card, bg=C["card"]); info_f.pack(fill="x", pady=(6, 0))
        self.pr_info_var = tk.StringVar(value="No PR loaded — enter a PR URL or number above")
        tk.Label(info_f, textvariable=self.pr_info_var,
                 bg=C["card"], fg=C["muted"],
                 font=("Segoe UI", 9), wraplength=820, justify="left"
                 ).pack(anchor="w")
        self.files_sum_var = tk.StringVar(value="")
        tk.Label(info_f, textvariable=self.files_sum_var,
                 bg=C["card"], fg=C["green"],
                 font=("Segoe UI", 9, "bold"), wraplength=820, justify="left"
                 ).pack(anchor="w")

    # Jira Fields card — fixed set of 4 fields sent directly to Jira API
    def _build_jira_fields_card(self, parent):
        card = self._card(parent, "Jira Fields", "")
        self._jira_field_widgets = {}
        jira_fds = [fd for fd in self.config_data["fields"]
                    if fd.get("jira_field") and fd.get("enabled", True)]

        for fd in jira_fds:
            key     = fd["key"]
            label   = fd["label"]
            ftype   = fd.get("type", "text")
            choices = fd.get("choices", [])
            req     = fd.get("required", False)

            if key == "project":
                default = self.config_data.get("jira_project_key", fd.get("default", ""))
            else:
                default = fd.get("default", "")

            row = tk.Frame(card, bg=C["card"]); row.pack(fill="x", pady=(0, 5))
            lf  = tk.Frame(row, bg=C["card"], width=130); lf.pack(side="left", fill="y")
            lf.pack_propagate(False)
            req_badge = " *" if req else ""
            tk.Label(lf, text=f"{label}{req_badge}",
                     bg=C["card"], fg=C["text"],
                     font=("Segoe UI", 9, "bold" if req else "normal"),
                     anchor="w").pack(side="left", padx=(0, 4), pady=4)

            var = tk.StringVar(value=default)
            if ftype == "dropdown" and choices:
                w = ttk.Combobox(row, textvariable=var, values=choices, state="readonly")
            else:
                w = ttk.Entry(row, textvariable=var)
            w.pack(side="left", fill="x", expand=True)

            if ftype == "dropdown":
                def _edit_jf(fd=fd, cb=w, v=var):
                    self._edit_dropdown_choices(fd, cb, v)
                tk.Button(row, text="✎", bg=C["surface"], fg=C["accent"],
                          relief="flat", font=("Segoe UI", 10), padx=6, pady=2,
                          activebackground=C["border"], cursor="hand2",
                          command=_edit_jf
                          ).pack(side="left", padx=(4, 0))

            self._jira_field_widgets[key] = var

        # auto-sync: Jira Summary → Doc File Name + Doc Title
        if "summary" in self._jira_field_widgets:
            def _sanitize_fname(s):
                s = re.sub(r'[\[\](){}<>:"/\\|?*\n\r\t]', '_', s)
                s = re.sub(r'\s+', '_', s.strip())
                s = re.sub(r'_+', '_', s).strip('_')
                return s[:80] if s else ""

            _cur_sum = self._jira_field_widgets["summary"].get()
            self._last_auto_doc_name  = _sanitize_fname(_cur_sum)
            self._last_auto_doc_title = _cur_sum
            self._syncing = False

            def _on_summary_change(*_):
                if getattr(self, "_syncing", False):
                    return
                self._syncing = True
                try:
                    sv = self._jira_field_widgets.get("summary")
                    if sv is None:
                        return
                    val = sv.get()
                    fw = getattr(self, "_field_widgets", {})
                    if "doc_name" in fw:
                        cur = fw["doc_name"].get()
                        if cur == getattr(self, "_last_auto_doc_name", None) or cur == "":
                            new = _sanitize_fname(val)
                            fw["doc_name"].set(new)
                            self._last_auto_doc_name = new
                    if "doc_title" in fw:
                        cur = fw["doc_title"].get()
                        if cur == getattr(self, "_last_auto_doc_title", None) or cur == "":
                            fw["doc_title"].set(val)
                            self._last_auto_doc_title = val
                finally:
                    self._syncing = False

            self._jira_field_widgets["summary"].trace_add("write", _on_summary_change)

    # Dynamic fields card — rebuilt whenever field definitions change
    def _build_dynamic_fields_card(self, parent):
        self._fields_card_parent = parent
        self._rebuild_fields_card()

    def _rebuild_fields_card(self):
        if hasattr(self, "_fields_card") and self._fields_card:
            self._fields_card.destroy()

        flat = tk.Frame(self._fields_card_parent, bg=C["bg"])
        flat.pack(fill="x")
        self._fields_card = flat

        self._field_widgets = {}
        enabled = [fd for fd in self.config_data["fields"]
                   if fd.get("enabled", True) and not fd.get("jira_field")]

        saved_vals = load_field_values()

        for idx, fd in enumerate(enabled):
            key = fd["key"]; label = fd["label"]
            ftype = fd.get("type","text"); choices = fd.get("choices",[])
            default = saved_vals.get(key, fd.get("default", ""))
            req = fd.get("required", False)
            row_bg = C["card"] if idx % 2 == 0 else C["bg"]

            outer = tk.Frame(flat, bg=row_bg)
            outer.pack(fill="x", padx=10, pady=0)
            row = tk.Frame(outer, bg=row_bg)
            row.pack(fill="x", padx=4, pady=(2, 2))
            lf  = tk.Frame(row, bg=row_bg, width=132); lf.pack(side="left", fill="y")
            lf.pack_propagate(False)

            lbl_w = tk.Label(lf, text=label,
                             bg=row_bg, fg=C["text"] if req else C["muted"],
                             font=("Segoe UI", 9, "bold" if req else "normal"),
                             anchor="w")
            lbl_w.pack(side="left", anchor="w", pady=4)
            if req:
                tk.Label(lf, text=" ●", bg=row_bg, fg=C["purple"],
                         font=("Segoe UI", 7)).pack(side="left", anchor="w")
            hint = FIELD_HINTS.get(key, "")
            if hint:
                Tooltip(lbl_w, hint)

            var = tk.StringVar(value=default)
            if ftype == "dropdown" and choices:
                w = ttk.Combobox(row, textvariable=var, values=choices, state="readonly")
            else:
                w = ttk.Entry(row, textvariable=var)
            w.pack(side="left", fill="x", expand=True)

            if key == "doc_name":
                tk.Button(row, text="⊞", bg=C["purple"], fg="#ffffff",
                          relief="flat", font=("Segoe UI", 9, "bold"), padx=7, pady=4,
                          activebackground="#a060cc", cursor="hand2",
                          command=self._choose_doc_dir).pack(side="left", padx=(4, 0))

            if ftype == "dropdown":
                def _edit_choices(fd=fd, cb=w, v=var):
                    self._edit_dropdown_choices(fd, cb, v)
                tk.Button(row, text="✎", bg=row_bg, fg=C["purple"],
                          relief="flat", font=("Segoe UI", 9), padx=5, pady=4,
                          activebackground=C["border"], cursor="hand2",
                          command=_edit_choices).pack(side="left", padx=(3, 0))

            self._field_widgets[key] = var

        # ── footer: description checkbox + defaults buttons ───────────────────
        sep = tk.Frame(flat, bg=C["border"], height=1)
        sep.pack(fill="x", padx=10, pady=(8, 4))
        cb_row = tk.Frame(flat, bg=C["bg"])
        cb_row.pack(fill="x", padx=10, pady=(0, 4))
        if not hasattr(self, "_include_in_desc_var"):
            self._include_in_desc_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(cb_row, text="Add Issue Details to Jira Description",
                        variable=self._include_in_desc_var).pack(side="left")

        def _save_defaults():
            vals = {k: v.get() for k, v in self._field_widgets.items()}
            save_field_values(vals)
            self._set_status("Issue Details defaults saved")

        def _clear_fields():
            for v in self._field_widgets.values():
                v.set("")

        btn_row = tk.Frame(flat, bg=C["bg"])
        btn_row.pack(fill="x", padx=10, pady=(0, 8))
        tk.Button(btn_row, text="Save as Default", bg=C["accent_dim"], fg=C["accent"],
                  relief="flat", font=("Segoe UI", 8, "bold"), padx=8, pady=3,
                  activebackground=C["border"], cursor="hand2",
                  command=_save_defaults).pack(side="left")
        tk.Button(btn_row, text="Clear All", bg=C["card2"], fg=C["muted"],
                  relief="flat", font=("Segoe UI", 8), padx=8, pady=3,
                  activebackground=C["border"], cursor="hand2",
                  command=_clear_fields).pack(side="left", padx=(4, 0))

    def _rebuild_jira_fields_section(self):
        if hasattr(self, "_jira_fields_card") and self._jira_fields_card:
            self._jira_fields_card.destroy()

        flat = tk.Frame(self._jira_fields_parent, bg=C["bg"])
        flat.pack(fill="x")
        self._jira_fields_card = flat

        self._jira_field_widgets = {}
        enabled = [fd for fd in self.config_data["fields"]
                   if fd.get("jira_field") and fd.get("enabled", True)]

        for idx, fd in enumerate(enabled):
            key     = fd["key"]; label = fd["label"]
            ftype   = fd.get("type", "text"); choices = fd.get("choices", [])
            default = (self.config_data.get("jira_project_key", fd.get("default", ""))
                       if key == "project" else fd.get("default", ""))
            req     = fd.get("required", False)
            row_bg  = C["card"] if idx % 2 == 0 else C["bg"]

            fr = tk.Frame(flat, bg=row_bg)
            fr.pack(fill="x", padx=10, pady=0)
            pad_f = tk.Frame(fr, bg=row_bg)
            pad_f.pack(fill="x", padx=4, pady=(2, 2))
            lf = tk.Frame(pad_f, bg=row_bg, width=132)
            lf.pack(side="left", fill="y")
            lf.pack_propagate(False)

            lbl_w = tk.Label(lf, text=label,
                             bg=row_bg, fg=C["text"] if req else C["muted"],
                             font=("Segoe UI", 9, "bold" if req else "normal"),
                             anchor="w")
            lbl_w.pack(side="left", anchor="w", pady=4)
            if req:
                tk.Label(lf, text=" ●", bg=row_bg, fg=C["accent"],
                         font=("Segoe UI", 7)).pack(side="left", anchor="w")

            var = tk.StringVar(value=default)
            if ftype == "dropdown" and choices:
                w = ttk.Combobox(pad_f, textvariable=var, values=choices, state="readonly")
            else:
                w = ttk.Entry(pad_f, textvariable=var)
            w.pack(side="left", fill="x", expand=True)

            if ftype == "dropdown":
                def _edit(fd=fd, cb=w, v=var):
                    self._edit_dropdown_choices(fd, cb, v)
                tk.Button(pad_f, text="✎", bg=row_bg, fg=C["accent"], relief="flat",
                          font=("Segoe UI", 9), padx=5, pady=4,
                          activebackground=C["border"],
                          cursor="hand2", command=_edit).pack(side="left", padx=(3, 0))

            self._jira_field_widgets[key] = var

        if "summary" in self._jira_field_widgets:
            def _sanitize(s):
                s = re.sub(r'[\[\](){}<>:"/\\|?*\n\r\t]', '_', s)
                s = re.sub(r'\s+', '_', s.strip())
                return re.sub(r'_+', '_', s).strip('_')[:80] if s else ""
            self._last_auto_doc_name  = _sanitize(self._jira_field_widgets["summary"].get())
            self._last_auto_doc_title = self._jira_field_widgets["summary"].get()
            self._syncing = False
            def _on_sum(*_):
                if getattr(self, "_syncing", False): return
                self._syncing = True
                try:
                    val = self._jira_field_widgets["summary"].get()
                    fw  = getattr(self, "_field_widgets", {})
                    if "doc_name" in fw:
                        cur = fw["doc_name"].get()
                        if cur in (getattr(self, "_last_auto_doc_name", None), ""):
                            new = _sanitize(val)
                            fw["doc_name"].set(new); self._last_auto_doc_name = new
                    if "doc_title" in fw:
                        cur = fw["doc_title"].get()
                        if cur in (getattr(self, "_last_auto_doc_title", None), ""):
                            fw["doc_title"].set(val); self._last_auto_doc_title = val
                finally:
                    self._syncing = False
            self._jira_field_widgets["summary"].trace_add("write", _on_sum)

    def _edit_dropdown_choices(self, fd, combobox, var):
        dlg = tk.Toplevel(self)
        dlg.title(f"Edit choices — {fd['label']}")
        dlg.geometry("400x360")
        dlg.configure(bg=C["bg"])
        dlg.grab_set()
        dlg.transient(self)

        tk.Frame(dlg, bg=C["accent"], height=3).pack(fill="x")
        hdr = tk.Frame(dlg, bg=C["surface"])
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=C["accent"], width=4).pack(side="left", fill="y")
        hdr_inner = tk.Frame(hdr, bg=C["surface"])
        hdr_inner.pack(side="left", padx=12, pady=10)
        tk.Label(hdr_inner, text=f"Edit Choices  —  {fd['label']}",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 11, "bold")).pack(anchor="w")
        tk.Label(hdr_inner, text="One value per line  •  Empty lines ignored",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(2, 0))
        tk.Frame(dlg, bg=C["border"], height=1).pack(fill="x")

        txt = scrolledtext.ScrolledText(dlg, height=10,
                                        bg=C["inp"], fg=C["text"],
                                        insertbackground=C["accent"],
                                        relief="flat", font=("Segoe UI", 10),
                                        highlightthickness=1,
                                        highlightbackground=C["border"],
                                        highlightcolor=C["accent"],
                                        padx=10, pady=8)
        txt.pack(fill="both", expand=True, padx=14, pady=(10, 6))
        txt.insert("1.0", "\n".join(fd.get("choices", [])))

        def _save():
            lines = [l.strip() for l in txt.get("1.0", "end-1c").splitlines() if l.strip()]
            fd["choices"] = lines
            if combobox:
                combobox.config(values=lines)
            # keep current value if still valid, else reset to first
            if var.get() not in lines and lines:
                var.set(lines[0])
            save_config(self.config_data)
            dlg.destroy()

        br = tk.Frame(dlg, bg=C["bg"]); br.pack(fill="x", padx=16, pady=(0, 14))
        tk.Button(br, text="✦  Save", bg=C["green"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=14, pady=6,
                  cursor="hand2", command=_save).pack(side="right")
        tk.Button(br, text="✕  Cancel", bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground="#c0392b", cursor="hand2",
                  command=dlg.destroy).pack(side="right", padx=(0, 8))

    def _build_desc_card(self, parent):
        card = self._card(parent, "Description", "", accent_color=C["muted"])
        self.desc_text = scrolledtext.ScrolledText(
            card, height=7, bg=C["inp"], fg=C["muted"],
            insertbackground=C["text"], relief="flat",
            font=("Segoe UI", 10), wrap="word", selectbackground=C["hover"],
            highlightthickness=1, highlightbackground=C["border"],
            highlightcolor=C["accent"])
        self.desc_text.pack(fill="x")
        self.desc_text.insert("1.0", "Describe the issue in detail...")

        def _clr(e):
            if self.desc_text.get("1.0", "end-1c") == "Describe the issue in detail...":
                self.desc_text.delete("1.0", "end")
                self.desc_text.config(fg=C["text"])
        self.desc_text.bind("<FocusIn>", _clr)

    def _build_actions_card(self, parent):
        card = self._card(parent, "Actions", "", accent_color=C["green"])

        # status label + progress bar
        self.result_var = tk.StringVar()
        rl = tk.Label(card, textvariable=self.result_var,
                      bg=C["card"], fg=C["green"],
                      font=("Segoe UI", 9, "bold"), cursor="hand2",
                      wraplength=400, justify="left")
        rl.pack(anchor="w", pady=(0, 6))
        rl.bind("<Button-1>", lambda e: self._jira_url and webbrowser.open(self._jira_url))

        self.progress = ttk.Progressbar(card, mode="indeterminate")
        self.progress.pack(fill="x", pady=(0, 16))

        # buttons — stacked full-width, right-aligned in the half-width column
        prev_btn = ttk.Button(card, text="  Preview Doc  ", style="Orange.TButton",
                              command=self._show_preview)
        prev_btn.pack(fill="x", pady=(0, 6))
        Tooltip(prev_btn, "Add change comments then generate the Word document.")

        self.sql_btn = ttk.Button(card, text="  SQL PR File  ", style="Accent.TButton",
                                  state="disabled", command=self._open_sql_popup)
        self.sql_btn.pack(fill="x", pady=(0, 6))
        Tooltip(self.sql_btn, "Combine all PR SQL files into a single runnable script.")

        att_row2 = tk.Frame(card, bg=C["card"]); att_row2.pack(fill="x", pady=(0, 4))
        tk.Label(att_row2, text="Attach:", bg=C["card"], fg=C["muted"],
                 font=("Segoe UI", 7)).pack(side="left", padx=(0, 4))
        ttk.Checkbutton(att_row2, text="Word Doc",
                        variable=self._attach_word_var).pack(side="left")
        ttk.Checkbutton(att_row2, text="SQL File",
                        variable=self._attach_sql_var).pack(side="left", padx=(8, 0))

        jira_btn = ttk.Button(card, text="  Create Jira Ticket  ", style="Success.TButton",
                              command=self._create_jira_thread)
        jira_btn.pack(fill="x")
        Tooltip(jira_btn, "Create a Jira ticket and attach Word doc + SQL file.")

        tk.Frame(card, bg=C["border"], height=1).pack(fill="x", pady=(8, 6))

        prev_jira_btn = ttk.Button(card, text="  Jira Preview  ", style="Accent.TButton",
                                   command=self._show_jira_preview)
        prev_jira_btn.pack(fill="x")
        Tooltip(prev_jira_btn, "Preview the Jira ticket that will be created — validate before submitting.")

    # ── Jira Preview window ───────────────────────────────────────────────────

    def _show_jira_preview(self):
        fv = self._collect_fields()

        # ── Jira-style light palette ──────────────────────────────────────────
        JC = {
            "bg":        "#F4F5F7",
            "surface":   "#FFFFFF",
            "header":    "#0052CC",
            "header_fg": "#FFFFFF",
            "border":    "#DFE1E6",
            "text":      "#172B4D",
            "muted":     "#5E6C84",
            "accent":    "#0052CC",
            "loz_open":  ("#0052CC", "#DEEBFF"),   # fg, bg
            "loz_done":  ("#006644", "#E3FCEF"),
            "loz_prog":  ("#0747A6", "#EAE6FF"),
            "bug":       "#FF5630",
            "story":     "#36B37E",
            "task":      "#0052CC",
            "epic":      "#6554C0",
            "section":   "#F4F5F7",
            "section_border": "#DFE1E6",
            "tag_bg":    "#EBECF0",
        }

        issue_type = fv.get("issue_type", "Bug")
        type_map   = {"Bug":"Bug","Enhancement":"Story","Task":"Task",
                      "Story":"Story","Epic":"Epic","Sub-task":"Sub-task","Incident":"Bug"}
        jira_type  = type_map.get(issue_type, "Bug")
        type_color = {"Bug": JC["bug"], "Story": JC["story"],
                      "Task": JC["task"], "Epic": JC["epic"]}.get(jira_type, JC["task"])
        type_icons = {"Bug": "⬤", "Story": "◆", "Task": "✔", "Epic": "⚡", "Sub-task": "↳"}
        type_icon  = type_icons.get(jira_type, "●")

        priority   = fv.get("priority", "Medium")
        pri_colors = {"Highest": "#FF5630", "High": "#FF7452",
                      "Medium":  "#FF991F", "Low":  "#2684FF", "Lowest": "#00B8D9"}
        pri_color  = pri_colors.get(priority, "#FF991F")
        pri_icons  = {"Highest": "▲▲", "High": "▲", "Medium": "●",
                      "Low": "▼", "Lowest": "▼▼"}
        pri_icon   = pri_icons.get(priority, "●")

        proj_key   = fv.get("project") or self.config_data.get("jira_project_key", "PROJ")
        summary    = fv.get("summary") or "(no summary)"
        reporter   = fv.get("reporter") or "—"
        description = fv.get("description") or "No description provided."

        # non-jira detail fields (only when checkbox enabled)
        _show_details = getattr(self, "_include_in_desc_var", None)
        _show_details = _show_details.get() if _show_details else True
        detail_rows = []
        if _show_details:
            for fd in self.config_data["fields"]:
                if not fd.get("enabled", True): continue
                if fd.get("jira_field"):        continue
                val = fv.get(fd["key"], "")
                if val and fd["key"] not in ("doc_name", "doc_title"):
                    detail_rows.append((fd["label"], val))

        file_rows = []
        for fn, status, jira_cmt, sql_cmt in self._build_changed_files():
            em = STATUS_EMOJI.get(status, "~")
            file_rows.append((fn, status, em, jira_cmt, sql_cmt))

        # ── Window ────────────────────────────────────────────────────────────
        win = tk.Toplevel(self)
        win.title("Jira Ticket Preview")
        win.geometry("980x720")
        win.configure(bg=JC["bg"])
        win.grab_set()

        # ── Jira top-bar ──────────────────────────────────────────────────────
        topbar = tk.Frame(win, bg=JC["header"], height=48)
        topbar.pack(fill="x")
        topbar.pack_propagate(False)
        inner_top = tk.Frame(topbar, bg=JC["header"])
        inner_top.pack(side="left", padx=16, fill="y")
        tk.Label(inner_top, text="◈  Jira", bg=JC["header"], fg="#FFFFFF",
                 font=("Segoe UI", 12, "bold")).pack(side="left", pady=12)
        tk.Label(inner_top, text=f"  /  {proj_key}  /  New Issue  (Preview)",
                 bg=JC["header"], fg="#B8D4FF",
                 font=("Segoe UI", 9)).pack(side="left", pady=12)
        tk.Label(topbar,
                 text="⚠  Not yet created  —  this is a preview only",
                 bg="#FF991F", fg="#172B4D",
                 font=("Segoe UI", 8, "bold"),
                 padx=10).pack(side="right", fill="y")

        # ── Scrollable body ───────────────────────────────────────────────────
        body_frame = tk.Frame(win, bg=JC["bg"])
        body_frame.pack(fill="both", expand=True)

        canvas = tk.Canvas(body_frame, bg=JC["bg"], highlightthickness=0)
        vsb    = ttk.Scrollbar(body_frame, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        scroll_frame = tk.Frame(canvas, bg=JC["bg"])
        wid = canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(wid, width=e.width))
        scroll_frame.bind("<Configure>",
                          lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<MouseWheel>",
                    lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        # ── Issue header area ─────────────────────────────────────────────────
        hdr_area = tk.Frame(scroll_frame, bg=JC["surface"],
                            relief="flat", bd=0)
        hdr_area.pack(fill="x", padx=0, pady=0)
        tk.Frame(hdr_area, bg=JC["border"], height=1).pack(fill="x")

        hdr_inner = tk.Frame(hdr_area, bg=JC["surface"])
        hdr_inner.pack(fill="x", padx=32, pady=18)

        # breadcrumb
        bc = tk.Frame(hdr_inner, bg=JC["surface"])
        bc.pack(anchor="w")
        tk.Label(bc, text=f"{proj_key}  ›", bg=JC["surface"], fg=JC["accent"],
                 font=("Segoe UI", 8), cursor="hand2").pack(side="left")
        tk.Label(bc, text="  New Issue", bg=JC["surface"], fg=JC["muted"],
                 font=("Segoe UI", 8)).pack(side="left")

        # issue type badge + key
        type_row = tk.Frame(hdr_inner, bg=JC["surface"])
        type_row.pack(anchor="w", pady=(6, 0))
        tk.Label(type_row, text=type_icon, bg=JC["surface"], fg=type_color,
                 font=("Segoe UI", 11, "bold")).pack(side="left")
        tk.Label(type_row, text=f"  {jira_type.upper()}",
                 bg=JC["surface"], fg=JC["muted"],
                 font=("Segoe UI", 9, "bold")).pack(side="left")
        tk.Label(type_row, text=f"  ·  {proj_key}-???",
                 bg=JC["surface"], fg=JC["muted"],
                 font=("Segoe UI", 9)).pack(side="left")

        # summary
        sum_lbl = tk.Label(hdr_inner, text=summary,
                           bg=JC["surface"], fg=JC["text"],
                           font=("Segoe UI", 17, "bold"),
                           wraplength=860, justify="left", anchor="w")
        sum_lbl.pack(anchor="w", pady=(8, 0))

        # status + priority row
        sp_row = tk.Frame(hdr_inner, bg=JC["surface"])
        sp_row.pack(anchor="w", pady=(10, 4))

        # status lozenge
        loz_bg = JC["loz_open"][1]; loz_fg = JC["loz_open"][0]
        loz = tk.Label(sp_row, text="  OPEN  ",
                       bg=loz_bg, fg=loz_fg,
                       font=("Segoe UI", 8, "bold"),
                       relief="flat", padx=4, pady=2)
        loz.pack(side="left")

        tk.Label(sp_row, text="  ", bg=JC["surface"]).pack(side="left")

        # priority lozenge
        tk.Label(sp_row, text=f"  {pri_icon}  {priority}  ",
                 bg=JC["tag_bg"], fg=pri_color,
                 font=("Segoe UI", 8, "bold"),
                 relief="flat", padx=4, pady=2).pack(side="left")

        tk.Frame(hdr_area, bg=JC["border"], height=1).pack(fill="x")

        # ── Two-column body ───────────────────────────────────────────────────
        cols = tk.Frame(scroll_frame, bg=JC["bg"])
        cols.pack(fill="both", expand=True, padx=20, pady=16)

        # ── LEFT column: description content ─────────────────────────────────
        left = tk.Frame(cols, bg=JC["bg"])
        left.pack(side="left", fill="both", expand=True, padx=(12, 16))

        def _section(parent, title):
            f = tk.Frame(parent, bg=JC["surface"],
                         relief="flat", bd=1,
                         highlightbackground=JC["border"],
                         highlightthickness=1)
            f.pack(fill="x", pady=(0, 12))
            hf = tk.Frame(f, bg=JC["section"])
            hf.pack(fill="x")
            tk.Frame(hf, bg=JC["accent"], width=3).pack(side="left", fill="y")
            tk.Label(hf, text=f"  {title}",
                     bg=JC["section"], fg=JC["text"],
                     font=("Segoe UI", 9, "bold"),
                     pady=7).pack(side="left")
            tk.Frame(f, bg=JC["border"], height=1).pack(fill="x")
            body = tk.Frame(f, bg=JC["surface"])
            body.pack(fill="x", padx=14, pady=10)
            return body

        # Description first
        desc_body = _section(left, "Description")
        tk.Label(desc_body, text=description,
                 bg=JC["surface"], fg=JC["text"],
                 font=("Segoe UI", 9),
                 wraplength=460, justify="left", anchor="nw").pack(anchor="w")

        # Changed Files
        if file_rows:
            fc_body = _section(left, "Changed Files")
            st_colors = {"added": "#1a7f37", "modified": "#0969da",
                         "removed": "#cf222e", "renamed": "#8250df"}
            for fname, status, emoji, jcmt, sqlcmt in file_rows:
                frow = tk.Frame(fc_body, bg="#F8F9FB",
                                highlightbackground=JC["border"],
                                highlightthickness=1)
                frow.pack(fill="x", pady=2)
                lborder = tk.Frame(frow, bg=st_colors.get(status, "#0969da"), width=3)
                lborder.pack(side="left", fill="y")
                finner = tk.Frame(frow, bg="#F8F9FB")
                finner.pack(side="left", fill="x", expand=True, padx=8, pady=5)
                tk.Label(finner,
                         text=f"{emoji}  {fname}  [{status.upper()}]",
                         bg="#F8F9FB", fg=JC["text"],
                         font=("Consolas", 8)).pack(anchor="w")
                if jcmt:
                    tk.Label(finner, text=f"    ↳  {jcmt}",
                             bg="#F8F9FB", fg=JC["muted"],
                             font=("Segoe UI", 8),
                             wraplength=420, justify="left").pack(anchor="w")
                if sqlcmt:
                    tk.Label(finner, text=f"    ↳  SQL: {sqlcmt}",
                             bg="#F8F9FB", fg=JC["yellow"],
                             font=("Segoe UI", 8),
                             wraplength=420, justify="left").pack(anchor="w")

        # Issue Details table — last in left column
        if detail_rows:
            det_body = _section(left, "Issue Details")
            hdr_row = tk.Frame(det_body, bg=JC["section"])
            hdr_row.pack(fill="x", pady=(0, 2))
            tk.Label(hdr_row, text="Field",
                     bg=JC["section"], fg=JC["muted"],
                     font=("Segoe UI", 8, "bold"), width=20, anchor="w",
                     padx=4, pady=3).pack(side="left")
            tk.Frame(hdr_row, bg=JC["border"], width=1).pack(side="left", fill="y")
            tk.Label(hdr_row, text="Value",
                     bg=JC["section"], fg=JC["muted"],
                     font=("Segoe UI", 8, "bold"), anchor="w",
                     padx=8, pady=3).pack(side="left", fill="x", expand=True)
            tk.Frame(det_body, bg=JC["border"], height=1).pack(fill="x", pady=(0, 2))
            for i, (lbl, val) in enumerate(detail_rows):
                row_bg = JC["surface"] if i % 2 == 0 else "#F8F9FB"
                row = tk.Frame(det_body, bg=row_bg)
                row.pack(fill="x")
                tk.Label(row, text=lbl,
                         bg=row_bg, fg=JC["muted"],
                         font=("Segoe UI", 9), width=20, anchor="w",
                         padx=4, pady=3).pack(side="left")
                tk.Frame(row, bg=JC["border"], width=1).pack(side="left", fill="y")
                tk.Label(row, text=val,
                         bg=row_bg, fg=JC["text"],
                         font=("Segoe UI", 9), anchor="w",
                         padx=8, pady=3, wraplength=400,
                         justify="left").pack(side="left", fill="x", expand=True)
                tk.Frame(det_body, bg=JC["border"], height=1).pack(fill="x")

        # ── RIGHT column: details sidebar ─────────────────────────────────────
        right = tk.Frame(cols, bg=JC["surface"],
                         highlightbackground=JC["border"],
                         highlightthickness=1,
                         width=260)
        right.pack(side="right", fill="y", padx=(0, 12))
        right.pack_propagate(False)

        def _detail_row(parent, label, value, val_color=None):
            f = tk.Frame(parent, bg=JC["surface"])
            f.pack(fill="x", padx=14, pady=5)
            tk.Frame(parent, bg=JC["border"], height=1).pack(fill="x", padx=14)
            tk.Label(f, text=label,
                     bg=JC["surface"], fg=JC["muted"],
                     font=("Segoe UI", 8), anchor="w").pack(anchor="w")
            tk.Label(f, text=value,
                     bg=JC["surface"], fg=val_color or JC["text"],
                     font=("Segoe UI", 9, "bold"), anchor="w",
                     wraplength=220, justify="left").pack(anchor="w")

        tk.Label(right, text="Details",
                 bg=JC["section"], fg=JC["text"],
                 font=("Segoe UI", 9, "bold"), pady=8,
                 anchor="w", padx=14).pack(fill="x")
        tk.Frame(right, bg=JC["border"], height=1).pack(fill="x")

        _detail_row(right, "Project",    proj_key,         JC["accent"])
        _detail_row(right, "Issue Type", jira_type,        type_color)
        _detail_row(right, "Status",     "Open")
        _detail_row(right, "Priority",   priority,         pri_color)
        _detail_row(right, "Reporter",   reporter)

        comp = fv.get("component", "")
        if comp:
            _detail_row(right, "Component", comp)
        fxv = fv.get("fix_version", "")
        if fxv:
            _detail_row(right, "Fix Version", fxv)

        git = fv.get("git_link", "")
        if git:
            tk.Frame(right, bg=JC["border"], height=1).pack(fill="x", padx=14)
            f_git = tk.Frame(right, bg=JC["surface"])
            f_git.pack(fill="x", padx=14, pady=5)
            tk.Label(f_git, text="PR Link", bg=JC["surface"], fg=JC["muted"],
                     font=("Segoe UI", 8), anchor="w").pack(anchor="w")
            git_lbl = tk.Label(f_git, text=git, bg=JC["surface"], fg=JC["accent"],
                               font=("Segoe UI", 8), anchor="w", cursor="hand2",
                               wraplength=220, justify="left")
            git_lbl.pack(anchor="w")
            git_lbl.bind("<Button-1>", lambda e, u=git: webbrowser.open(u))

        # ── Footer buttons ────────────────────────────────────────────────────
        foot = tk.Frame(win, bg=JC["surface"])
        foot.pack(fill="x", side="bottom")
        tk.Frame(foot, bg=JC["border"], height=1).pack(fill="x")
        btn_row = tk.Frame(foot, bg=JC["surface"])
        btn_row.pack(fill="x", padx=24, pady=10)
        ttk.Button(btn_row, text="✕  Close Preview",
                   style="Ghost.TButton",
                   command=win.destroy).pack(side="right", padx=(8, 0))
        ttk.Button(btn_row, text="✦  Create Jira Ticket",
                   style="Success.TButton",
                   command=lambda: [win.destroy(),
                                    self._create_jira_thread()]).pack(side="right")

        def _copy_api_json():
            proj_key_ = self.config_data.get("jira_project_key", "PROJ")
            _type_map = {"Bug":"Bug","Enhancement":"Story","Task":"Task",
                         "Story":"Story","Epic":"Epic","Sub-task":"Sub-task","Incident":"Bug"}
            _issue_type = _type_map.get(fv.get("issue_type","Bug"), "Bug")
            _inc = getattr(self, "_include_in_desc_var", None)
            _inc = _inc.get() if _inc else True
            _desc_lines = [fv.get("description") or ""]
            _changed = self._build_changed_files()
            if _changed:
                _desc_lines.append("")
                _desc_lines.append("h3. Changed Files")
                for _fn, _st, _jc, _sc in _changed:
                    _t = f"{STATUS_EMOJI.get(_st,'~')} {_fn}  [{_st.upper()}]"
                    if _jc: _t += f"\n  {_jc}"
                    if _sc: _t += f"\n  SQL: {_sc}"
                    _desc_lines.append(f"* {_t}")
            if _inc:
                _det_rows = [(fd["label"], fv.get(fd["key"], ""))
                             for fd in self.config_data["fields"]
                             if fd.get("enabled", True) and not fd.get("jira_field")
                             and fd["key"] not in ("doc_name","doc_title")
                             and fv.get(fd["key"], "")]
                if _det_rows:
                    _desc_lines.append("")
                    _desc_lines.append("h3. Issue Details")
                    _desc_lines.append("||Field||Value||")
                    for _lbl, _val in _det_rows:
                        _desc_lines.append(f"|{_lbl}|{_val}|")
            _payload = {"fields":{
                "project":     {"key": fv.get("project") or proj_key_},
                "summary":     fv.get("summary") or "(no summary)",
                "issuetype":   {"name": _issue_type},
                "description": "\n".join(_desc_lines),
            }}
            if fv.get("reporter"):
                _payload["fields"]["reporter"] = {"name": fv["reporter"]}
            self.clipboard_clear()
            self.clipboard_append(json.dumps(_payload, indent=2))
            messagebox.showinfo("Copied", "Jira API payload copied to clipboard.", parent=win)

        tk.Button(btn_row, text="⧉  Copy API JSON",
                  bg=JC["surface"], fg=JC["accent"], relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground=JC["border"], cursor="hand2",
                  command=_copy_api_json).pack(side="left")

    # ── Settings ─────────────────────────────────────────────────────────────


    # ── settings window ───────────────────────────────────────────────────────

    def _open_settings(self):
        win = tk.Toplevel(self)
        win.title(f"{APP_NAME}  —  Settings")
        win.geometry("660x560")
        win.configure(bg=C["bg"])
        win.grab_set()

        # top accent stripe
        tk.Frame(win, bg=C["orange"], height=3).pack(fill="x")
        hdr = tk.Frame(win, bg=C["surface"])
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=C["orange"], width=4).pack(side="left", fill="y")
        hdr_inner = tk.Frame(hdr, bg=C["surface"])
        hdr_inner.pack(side="left", padx=14, pady=12)
        hdr_row = tk.Frame(hdr_inner, bg=C["surface"])
        hdr_row.pack(anchor="w")
        tk.Label(hdr_row, text="⚙", bg=C["surface"], fg=C["orange"],
                 font=("Segoe UI", 13)).pack(side="left", padx=(0, 8))
        tk.Label(hdr_row, text="Connection Settings",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 13, "bold")).pack(side="left")
        tk.Label(hdr_inner, text="GitHub and Jira API credentials  •  Output directory",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(2, 0))
        tk.Frame(win, bg=C["border"], height=1).pack(fill="x")

        c = tk.Canvas(win, bg=C["bg"], highlightthickness=0)
        v = ttk.Scrollbar(win, orient="vertical", command=c.yview)
        c.configure(yscrollcommand=v.set)
        v.pack(side="right", fill="y")
        c.pack(side="left", fill="both", expand=True)
        sf  = tk.Frame(c, bg=C["bg"])
        wid = c.create_window((0, 0), window=sf, anchor="nw")
        c.bind("<Configure>", lambda e: c.itemconfig(wid, width=e.width))
        sf.bind("<Configure>", lambda e: c.configure(scrollregion=c.bbox("all")))

        defs = [
            ("GitHub Token File",         "github_token_file",   "⎇"),
            ("GitHub API URL",            "github_api_url",      "⎇"),
            ("GitHub Owner",              "github_owner",        "⎇"),
            ("GitHub Repo",               "github_repo",         "⎇"),
            ("GitHub SSL Certificate",    "ssl_cert_file",       "⎇"),
            ("Jira Token File",           "jira_token_file",     "⊡"),
            ("Jira API URL",              "jira_base_url",       "≋"),
            ("Jira Project Key",          "jira_project_key",    "⊡"),
            ("Jira Email",                "jira_email",          "≋"),
            ("Jira SSL Certificate",      "jira_ssl_cert_file",  "⊡"),
            ("Output Directory",          "word_doc_output_dir", "⊞"),
        ]
        vars_ = {}
        for idx, (label, key, icon) in enumerate(defs):
            row_bg = C["card"] if idx % 2 == 0 else C["bg"]
            row = tk.Frame(sf, bg=row_bg)
            row.pack(fill="x", padx=16, pady=0)
            inner = tk.Frame(row, bg=row_bg)
            inner.pack(fill="x", padx=6, pady=5)
            lbl_f = tk.Frame(inner, bg=row_bg)
            lbl_f.pack(side="left")
            tk.Label(lbl_f, text=icon, bg=row_bg, fg=C["accent"],
                     font=("Segoe UI", 9), width=2).pack(side="left")
            tk.Label(lbl_f, text=label, bg=row_bg, fg=C["muted"],
                     width=22, anchor="w", font=("Segoe UI", 9)).pack(side="left")
            v2 = tk.StringVar(value=str(self.config_data.get(key, "")))
            vars_[key] = v2
            ttk.Entry(inner, textvariable=v2).pack(side="left", fill="x", expand=True)
            if "file" in key or "dir" in key.lower():
                is_dir = "dir" in key.lower()
                is_pem = key in ("ssl_cert_file", "jira_ssl_cert_file")
                def _browse(var=v2, d=is_dir, pem=is_pem):
                    if d:
                        p = filedialog.askdirectory()
                    elif pem:
                        p = filedialog.askopenfilename(
                            filetypes=[("PEM Certificate","*.pem *.crt *.cer"),("All","*.*")])
                    else:
                        p = filedialog.askopenfilename(
                            filetypes=[("Text","*.txt"),("All","*.*")])
                    if p: var.set(p)
                ttk.Button(inner, text="⊞", style="Subtle.TButton",
                           command=_browse).pack(side="left", padx=(6, 0))

        # ── Auto-open Word doc checkbox ───────────────────────────────────────
        chk_row_bg = C["card"] if len(defs) % 2 == 0 else C["bg"]
        chk_row = tk.Frame(sf, bg=chk_row_bg)
        chk_row.pack(fill="x", padx=16, pady=0)
        chk_inner = tk.Frame(chk_row, bg=chk_row_bg)
        chk_inner.pack(fill="x", padx=6, pady=5)
        chk_lbl_f = tk.Frame(chk_inner, bg=chk_row_bg)
        chk_lbl_f.pack(side="left")
        tk.Label(chk_lbl_f, text="⊞", bg=chk_row_bg, fg=C["accent"],
                 font=("Segoe UI", 9), width=2).pack(side="left")
        tk.Label(chk_lbl_f, text="Auto-open Word Doc", bg=chk_row_bg, fg=C["muted"],
                 width=22, anchor="w", font=("Segoe UI", 9)).pack(side="left")
        auto_open_var = tk.BooleanVar(value=bool(self.config_data.get("auto_open_word_doc", False)))
        ttk.Checkbutton(chk_inner, variable=auto_open_var).pack(side="left")
        tk.Label(chk_inner, text="Open generated .docx automatically after saving",
                 bg=chk_row_bg, fg=C["muted"], font=("Segoe UI", 8)).pack(side="left", padx=(6, 0))

        # ── SSL verify toggles (GitHub + Jira) ───────────────────────────────
        def _make_ssl_toggle(parent_frame, label, icon, cfg_key, default=True):
            n = len([w for w in parent_frame.winfo_children()])
            row_bg = C["card"] if n % 2 == 0 else C["bg"]
            row = tk.Frame(parent_frame, bg=row_bg)
            row.pack(fill="x", padx=16, pady=0)
            inner = tk.Frame(row, bg=row_bg)
            inner.pack(fill="x", padx=6, pady=5)
            lf = tk.Frame(inner, bg=row_bg)
            lf.pack(side="left")
            tk.Label(lf, text=icon, bg=row_bg, fg=C["accent"],
                     font=("Segoe UI", 9), width=2).pack(side="left")
            tk.Label(lf, text=label, bg=row_bg, fg=C["muted"],
                     width=22, anchor="w", font=("Segoe UI", 9)).pack(side="left")
            var = tk.BooleanVar(value=bool(self.config_data.get(cfg_key, default)))

            def _refresh(btn, v):
                if v.get():
                    btn.configure(text="✓  Enabled", bg=C["green"],  fg="#000000")
                else:
                    btn.configure(text="✗  Disabled", bg=C["red"], fg="#ffffff")

            btn = tk.Button(inner, font=("Segoe UI", 9, "bold"),
                            relief="flat", padx=10, pady=3, cursor="hand2")
            btn.configure(command=lambda: [var.set(not var.get()), _refresh(btn, var)])
            _refresh(btn, var)
            btn.pack(side="left")
            tk.Label(inner, text="SSL certificate verification",
                     bg=row_bg, fg=C["muted"],
                     font=("Segoe UI", 8)).pack(side="left", padx=(10, 0))
            return var

        gh_ssl_var   = _make_ssl_toggle(sf, "GitHub SSL Verify", "⎇",
                                        "github_ssl_verify", default=True)
        jira_ssl_var = _make_ssl_toggle(sf, "Jira SSL Verify",   "⊡",
                                        "jira_ssl_verify",   default=True)

        # ── Jira Auth Method selector ─────────────────────────────────────────
        _auth_methods = [("Basic Auth (email + API token)", "basic"),
                         ("Bearer Token (PAT / Server)",    "bearer"),
                         ("OAuth 2.0 (email + access token)", "oauth")]
        _auth_row_n = len([w for w in sf.winfo_children()])
        auth_row_bg = C["card"] if _auth_row_n % 2 == 0 else C["bg"]
        auth_row = tk.Frame(sf, bg=auth_row_bg)
        auth_row.pack(fill="x", padx=16, pady=0)
        auth_inner = tk.Frame(auth_row, bg=auth_row_bg)
        auth_inner.pack(fill="x", padx=6, pady=6)
        auth_lbl_f = tk.Frame(auth_inner, bg=auth_row_bg)
        auth_lbl_f.pack(side="left")
        tk.Label(auth_lbl_f, text="⊡", bg=auth_row_bg, fg=C["accent"],
                 font=("Segoe UI", 9), width=2).pack(side="left")
        tk.Label(auth_lbl_f, text="Jira Auth Method", bg=auth_row_bg, fg=C["muted"],
                 width=22, anchor="w", font=("Segoe UI", 9)).pack(side="left")
        jira_auth_method_var = tk.StringVar(
            value=self.config_data.get("jira_auth_method", "basic"))
        btn_f = tk.Frame(auth_inner, bg=auth_row_bg)
        btn_f.pack(side="left")
        _auth_btns = {}
        def _refresh_auth_btns(selected):
            for _v, _b in _auth_btns.items():
                if _v == selected:
                    _b.configure(bg=C["accent"], fg="#ffffff", relief="flat")
                else:
                    _b.configure(bg=C["card2"], fg=C["muted"], relief="flat")
        for _lbl, _val in _auth_methods:
            def _pick(v=_val):
                jira_auth_method_var.set(v)
                _refresh_auth_btns(v)
            b = tk.Button(btn_f, text=_lbl, font=("Segoe UI", 8),
                          relief="flat", padx=8, pady=3, cursor="hand2",
                          command=_pick)
            b.pack(side="left", padx=(0, 4))
            _auth_btns[_val] = b
        _refresh_auth_btns(jira_auth_method_var.get())

        def _save():
            for k, v2 in vars_.items():
                self.config_data[k] = v2.get()
            self.config_data["auto_open_word_doc"]  = auto_open_var.get()
            self.config_data["github_ssl_verify"]   = gh_ssl_var.get()
            self.config_data["jira_ssl_verify"]     = jira_ssl_var.get()
            self.config_data["jira_auth_method"]    = jira_auth_method_var.get()
            save_config(self.config_data)
            self._set_status("Settings saved")
            win.destroy()

        def _export_config():
            path = filedialog.asksaveasfilename(
                parent=win, defaultextension=".json",
                filetypes=[("JSON", "*.json"), ("All", "*.*")],
                title="Export Config")
            if not path: return
            try:
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(self.config_data, f, indent=2, default=str)
                messagebox.showinfo("Export", f"Config exported to:\n{path}", parent=win)
            except Exception as e:
                messagebox.showerror("Export Error", str(e), parent=win)

        def _import_config():
            path = filedialog.askopenfilename(
                parent=win, filetypes=[("JSON", "*.json"), ("All", "*.*")],
                title="Import Config")
            if not path: return
            try:
                with open(path, "r", encoding="utf-8") as f:
                    loaded = json.load(f)
                required = {"jira_base_url", "github_owner", "fields"}
                missing = required - set(loaded.keys())
                if missing:
                    messagebox.showerror("Import Error",
                        f"Missing required keys: {', '.join(missing)}", parent=win)
                    return
                self.config_data.update(loaded)
                save_config(self.config_data)
                for k, v2 in vars_.items():
                    v2.set(str(self.config_data.get(k, "")))
                auto_open_var.set(bool(self.config_data.get("auto_open_word_doc", False)))
                messagebox.showinfo("Import", "Config imported and applied.", parent=win)
            except Exception as e:
                messagebox.showerror("Import Error", str(e), parent=win)

        br = tk.Frame(sf, bg=C["bg"]); br.pack(fill="x", padx=20, pady=20)
        ttk.Button(br, text="✦  Save Settings", style="Success.TButton",
                   command=_save).pack(side="right")
        tk.Button(br, text="✕  Cancel",
                  bg=C["red"], fg="#ffffff", relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground="#c0392b", cursor="hand2",
                  command=win.destroy).pack(side="right", padx=(0, 8))
        tk.Button(br, text="⬆  Import",
                  bg=C["card2"], fg=C["text"], relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground=C["border"], cursor="hand2",
                  command=_import_config).pack(side="left", padx=(0, 6))
        tk.Button(br, text="⬇  Export",
                  bg=C["card2"], fg=C["text"], relief="flat",
                  font=("Segoe UI", 10), padx=12, pady=6,
                  activebackground=C["border"], cursor="hand2",
                  command=_export_config).pack(side="left")

        # ── [TEST] Jira connection test button ────────────────────────────────
        def _test_jira():
            snap = dict(self.config_data)
            for k, v2 in vars_.items():
                snap[k] = v2.get()
            snap["jira_ssl_verify"]   = jira_ssl_var.get()
            snap["github_ssl_verify"] = gh_ssl_var.get()
            snap["jira_auth_method"]  = jira_auth_method_var.get()
            self._test_jira_connection(parent=win, cfg=snap)

        tk.Button(br, text="⚡  Test Jira",
                  bg=C["yellow"], fg="#000000", relief="flat",
                  font=("Segoe UI", 10, "bold"), padx=12, pady=6,
                  activebackground=C["orange"], cursor="hand2",
                  command=_test_jira).pack(side="left", padx=(12, 0))

    # ── [TEST] Jira connection tester ─────────────────────────────────────────

    def _test_jira_connection(self, parent=None, cfg=None):
        """Fetch up to 20 tickets reported-by current Jira user and show popup."""
        cfg = cfg or self.config_data
        parent = parent or self

        tok_file = cfg.get("jira_token_file", "")
        email    = cfg.get("jira_email", "").strip()
        base_url = cfg.get("jira_base_url", "").rstrip("/")

        try:
            jira_tok = read_token(tok_file)
        except Exception as e:
            messagebox.showerror("Jira Test — Token Error", str(e), parent=parent)
            return

        method = cfg.get("jira_auth_method", "basic")
        hdrs   = jira_auth_headers(email, jira_tok, method)
        verify = _jira_ssl_verify(cfg)

        def _fetch():
            jql    = f'reporter = currentUser() ORDER BY created DESC'
            params = {"jql": jql, "maxResults": 20,
                      "fields": "summary,status,issuetype,created,key"}
            try:
                r = requests.get(f"{base_url}/search",
                                 headers=hdrs, params=params,
                                 timeout=15, verify=verify)
                r.raise_for_status()
                return r.json(), None
            except requests.HTTPError as e:
                return None, f"HTTP {e.response.status_code}: {e.response.text[:300]}"
            except Exception as e:
                return None, str(e)

        # run in thread — show "connecting…" popup first
        popup = tk.Toplevel(parent)
        popup.title(f"{APP_NAME}  —  Jira Connection Test")
        popup.geometry("700x520")
        popup.configure(bg=C["bg"])
        popup.grab_set()

        tk.Frame(popup, bg=C["accent"], height=3).pack(fill="x")
        hdr = tk.Frame(popup, bg=C["surface"])
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=C["accent"], width=4).pack(side="left", fill="y")
        hi = tk.Frame(hdr, bg=C["surface"])
        hi.pack(side="left", padx=14, pady=10)
        tk.Label(hi, text="⚡  Jira Connection Test",
                 bg=C["surface"], fg=C["text"],
                 font=("Segoe UI", 12, "bold")).pack(anchor="w")
        tk.Label(hi, text=f"{base_url}  •  {email}",
                 bg=C["surface"], fg=C["muted"],
                 font=("Segoe UI", 8)).pack(anchor="w")
        tk.Frame(popup, bg=C["border"], height=1).pack(fill="x")

        status_var = tk.StringVar(value="Connecting…")
        tk.Label(popup, textvariable=status_var,
                 bg=C["bg"], fg=C["yellow"],
                 font=("Segoe UI", 9, "italic")).pack(anchor="w", padx=16, pady=(8, 0))

        # scrollable results frame
        c2    = tk.Canvas(popup, bg=C["bg"], highlightthickness=0)
        vsb   = ttk.Scrollbar(popup, orient="vertical", command=c2.yview)
        c2.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        c2.pack(side="left", fill="both", expand=True)
        sf2   = tk.Frame(c2, bg=C["bg"])
        wid2  = c2.create_window((0, 0), window=sf2, anchor="nw")
        c2.bind("<Configure>", lambda e: c2.itemconfig(wid2, width=e.width))
        sf2.bind("<Configure>", lambda e: c2.configure(scrollregion=c2.bbox("all")))

        def _populate(data, err):
            if err:
                status_var.set(f"Connection FAILED")
                tk.Label(sf2, text=err, bg=C["bg"], fg=C["red"],
                         font=("Segoe UI", 9), wraplength=620,
                         justify="left").pack(anchor="w", padx=16, pady=8)
                return

            issues = data.get("issues", [])
            total  = data.get("total", 0)
            status_var.set(f"Connected  ✓   Showing {len(issues)} of {total} tickets reported by you")

            if not issues:
                tk.Label(sf2, text="No tickets found.", bg=C["bg"], fg=C["muted"],
                         font=("Segoe UI", 9)).pack(anchor="w", padx=16, pady=8)
                return

            # header row
            hrow = tk.Frame(sf2, bg=C["card2"])
            hrow.pack(fill="x", padx=12, pady=(8, 2))
            for txt, w in [("Key", 12), ("Type", 10), ("Status", 14), ("Created", 12), ("Summary", 42)]:
                tk.Label(hrow, text=txt, bg=C["card2"], fg=C["muted"],
                         font=("Segoe UI", 8, "bold"), width=w, anchor="w").pack(side="left", padx=4)

            for i, iss in enumerate(issues):
                f  = iss.get("fields", {})
                row_bg = C["card"] if i % 2 == 0 else C["bg"]
                row = tk.Frame(sf2, bg=row_bg, cursor="hand2")
                row.pack(fill="x", padx=12, pady=1)
                key     = iss.get("key", "")
                itype   = (f.get("issuetype") or {}).get("name", "")
                status  = (f.get("status")    or {}).get("name", "")
                created = (f.get("created")   or "")[:10]
                summary = f.get("summary", "")

                url = f"{base_url}/browse/{key}"
                for txt, w, col in [
                    (key,     12, C["accent"]),
                    (itype,   10, C["purple"]),
                    (status,  14, C["green"]),
                    (created, 12, C["muted"]),
                    (summary, 42, C["text"]),
                ]:
                    lbl = tk.Label(row, text=txt, bg=row_bg, fg=col,
                                   font=("Segoe UI", 9), width=w, anchor="w")
                    lbl.pack(side="left", padx=4, pady=3)
                    lbl.bind("<Button-1>", lambda e, u=url: webbrowser.open(u))
                    lbl.bind("<Enter>",    lambda e, r2=row, b=row_bg: r2.configure(bg=C["accent_dim"]))
                    lbl.bind("<Leave>",    lambda e, r2=row, b=row_bg: r2.configure(bg=b))

        def _thread():
            data, err = _fetch()
            popup.after(0, lambda: _populate(data, err))

        threading.Thread(target=_thread, daemon=True).start()

    # ── Field managers ────────────────────────────────────────────────────────

    def _open_jira_fields_mgr(self):
        """Popup: enable/disable Jira API fields shown on the left panel."""
        dlg = FieldToggleDialog(
            self, self.config_data["fields"],
            "Jira Fields",
            "Enable / disable fields sent directly to the Jira API  •  left panel",
            lambda fd: fd.get("jira_field", False),
            accent_color=C["accent"],
        )
        self.wait_window(dlg)
        if dlg.result is not None:
            self.config_data["fields"] = dlg.result
            save_config(self.config_data)
            self._rebuild_jira_fields_section()
            self._set_status("Jira fields updated")

    def _open_desc_info_mgr(self):
        """Popup: manage Description Info fields — toggle, set values, add custom."""
        dlg = DescInfoDialog(self, self.config_data["fields"])
        self.wait_window(dlg)
        if dlg.result is not None:
            self.config_data["fields"] = dlg.result
            self.config_data["field_defaults"] = {
                fd["key"]: fd["default"] for fd in dlg.result}
            save_config(self.config_data)
            self._rebuild_fields_card()
            self._set_status("Description fields updated")

    def _open_field_mgr(self):
        """Legacy full field manager (advanced)."""
        dlg = FieldManagerDialog(self, self.config_data["fields"])
        self.wait_window(dlg)
        if dlg.result is not None:
            self.config_data["fields"] = dlg.result
            self.config_data["field_defaults"] = {
                fd["key"]: fd["default"] for fd in dlg.result}
            save_config(self.config_data)
            self._rebuild_jira_fields_section()
            self._rebuild_fields_card()
            self._set_status("Fields updated — defaults saved")

    # ── Fetch PR ──────────────────────────────────────────────────────────────

    def _fetch_pr_thread(self):
        threading.Thread(target=self._fetch_pr, daemon=True).start()

    def _fetch_pr(self):
        self.after(0, self.progress.start)
        self.after(0, lambda: self._set_status("Fetching PR from GitHub..."))
        try:
            token = read_token(self.config_data["github_token_file"])
        except FileNotFoundError as e:
            self.after(0, lambda m=str(e): messagebox.showerror("Token Error", m))
            self.after(0, self.progress.stop); return

        raw   = self.pr_url_var.get().strip()
        owner = self.config_data.get("github_owner", "")
        repo  = self.config_data.get("github_repo", "")
        num   = None

        if raw.isdigit():
            num = raw
        elif "github.com" in raw:
            parts = raw.rstrip("/").split("/")
            try:
                idx = parts.index("pull")
                num = parts[idx+1]; owner = parts[idx-2]; repo = parts[idx-1]
            except (ValueError, IndexError):
                pass

        if not num:
            self.after(0, lambda: messagebox.showerror(
                "Error", "Enter a PR number or full GitHub PR URL"))
            self.after(0, self.progress.stop); return
        if not owner or not repo:
            self.after(0, lambda: messagebox.showerror(
                "Error", "Set GitHub Owner and Repo in Settings"))
            self.after(0, self.progress.stop); return

        hdrs       = gh_headers(token)
        api_base   = self.config_data.get("github_api_url", "https://api.github.com").rstrip("/")
        verify     = _ssl_verify(self.config_data)
        try:
            pr_r = requests.get(
                f"{api_base}/repos/{owner}/{repo}/pulls/{num}",
                headers=hdrs, timeout=15, verify=verify)
            pr_r.raise_for_status()
            pr   = pr_r.json()

            files_r = requests.get(
                f"{api_base}/repos/{owner}/{repo}/pulls/{num}/files?per_page=100",
                headers=hdrs, timeout=15, verify=verify)
            files_r.raise_for_status()
            files = files_r.json()

            # fetch per-file review comments
            rev_r = requests.get(
                f"{api_base}/repos/{owner}/{repo}/pulls/{num}/comments?per_page=100",
                headers=hdrs, timeout=15, verify=verify)
            review_comments = rev_r.json() if rev_r.ok else []

            # fetch PR-level issue comments
            iss_r = requests.get(
                f"{api_base}/repos/{owner}/{repo}/issues/{num}/comments?per_page=100",
                headers=hdrs, timeout=15, verify=verify)
            issue_comments = iss_r.json() if iss_r.ok else []

            # fetch PR reviews (to get reviewer names)
            reviews_r = requests.get(
                f"{api_base}/repos/{owner}/{repo}/pulls/{num}/reviews?per_page=100",
                headers=hdrs, timeout=15, verify=verify)
            pr_reviews = reviews_r.json() if reviews_r.ok else []

        except requests.HTTPError as e:
            msg = f"GitHub API error {e.response.status_code}:\n{e.response.text[:300]}"
            self.after(0, lambda m=msg: messagebox.showerror("GitHub Error", m))
            self.after(0, self.progress.stop); return
        except Exception as e:
            self.after(0, lambda m=str(e): messagebox.showerror("Network Error", m))
            self.after(0, self.progress.stop); return

        # group review comments by file — sort newest-first so index 0 = latest comment
        from collections import defaultdict
        rev_by_file = defaultdict(list)
        for c in sorted(review_comments,
                        key=lambda x: x.get("updated_at") or x.get("created_at", ""),
                        reverse=True):
            path = c.get("path", "")
            body = c.get("body", "").strip()
            user = (c.get("user") or {}).get("login", "")
            if path and body:
                rev_by_file[path].append(f"{user}: {body}" if user else body)

        # build file_comments — use latest comment only per file
        file_comments = []
        for f in files:
            fname    = f.get("filename", "")
            comments = rev_by_file.get(fname, [])
            jira_cmt = comments[0] if comments else ""
            word_cmt = jira_cmt
            file_comments.append(dict(
                filename=fname, status=f.get("status", "modified"),
                additions=f.get("additions", 0), deletions=f.get("deletions", 0),
                word_comment=word_cmt, jira_comment=jira_cmt,
                comment=word_cmt, screenshots=[]))

        # pick reviewer: first APPROVED reviewer, else first reviewer in any review
        reviewer_login = ""
        approved = [r for r in pr_reviews if r.get("state") == "APPROVED"]
        if approved:
            reviewer_login = (approved[0].get("user") or {}).get("login", "")
        elif pr_reviews:
            reviewer_login = (pr_reviews[0].get("user") or {}).get("login", "")
        # fall back to first requested_reviewer if no reviews yet
        if not reviewer_login:
            req_reviewers = pr.get("requested_reviewers") or []
            if req_reviewers:
                reviewer_login = (req_reviewers[0] or {}).get("login", "")

        self.pr_cache          = pr
        self.pr_files          = files
        self.file_comments     = file_comments
        self.pr_review_comments= review_comments
        self.pr_issue_comments = issue_comments
        self.pr_reviews        = pr_reviews

        def _ui():
            self.pr_info_var.set(
                f"  #{pr['number']}  {pr['title']}   "
                f"{pr['state'].upper()}   by {pr['user']['login']}   "
                f"{(pr.get('created_at') or '')[:10]}")

            jfw = getattr(self, "_jira_field_widgets", {})

            # PR title → Jira Summary
            if "summary" in jfw and not jfw["summary"].get().strip():
                jfw["summary"].set(pr["title"])

            # PR creator → Reporter
            author = (pr.get("user") or {}).get("login", "")
            if "reporter" in jfw and author:
                jfw["reporter"].set(author)

            # PR reviewer → Assignee
            if "assignee" in jfw and reviewer_login:
                jfw["assignee"].set(reviewer_login)

            # PR body → Description text area
            pr_body = (pr.get("body") or "").strip()
            if pr_body:
                cur = self.desc_text.get("1.0", "end-1c").strip()
                if not cur or cur == "Describe the issue in detail...":
                    self.desc_text.delete("1.0", "end")
                    self.desc_text.config(fg=C["text"])
                    self.desc_text.insert("1.0", pr_body)

            self.pr_url_var.set(pr.get("html_url", ""))

            total = len(files)
            adds  = sum(1 for f in files if f.get("status")=="added")
            mods  = sum(1 for f in files if f.get("status")=="modified")
            rems  = sum(1 for f in files if f.get("status")=="removed")
            other = total - adds - mods - rems
            parts = [f"{total} file{'s' if total!=1 else ''} changed"]
            if adds:  parts.append(f"A {adds} added")
            if mods:  parts.append(f"M {mods} modified")
            if rems:  parts.append(f"D {rems} removed")
            if other: parts.append(f"~ {other} other")
            self.files_sum_var.set("   ".join(parts))
            if hasattr(self, "sql_btn"):
                self.sql_btn.config(state="normal")
            if hasattr(self, "_gh_preview_btn"):
                self._gh_preview_btn.config(state="normal")
            self.progress.stop()
            self._set_status(f"PR #{pr['number']} loaded — {total} files"
                             + (f"  |  {len(review_comments)} review comments" if review_comments else ""))

        self.after(0, _ui)

    # ── File changes ──────────────────────────────────────────────────────────

    def _open_sql_popup(self):
        if not self.pr_files:
            messagebox.showinfo("No PR loaded", "Fetch a PR first.", parent=self)
            return
        dlg = SqlFilePopup(self, self.pr_files, self.pr_cache or {}, self.config_data, self.file_comments)
        self.wait_window(dlg)
        result = dlg.get_result()
        if result and os.path.exists(result):
            self._last_sql_path = result
        sql_cmts = dlg.get_sql_comments()
        if sql_cmts:
            self._sql_file_comments = sql_cmts
        self._sql_keyword_findings = dlg.get_keyword_findings()

    # ── Mock / test data ──────────────────────────────────────────────────────

    def _load_mock_data(self):
        self.pr_cache = {
            "number":     42,
            "title":      "feat: inventory & order management SQL overhaul",
            "state":      "open",
            "html_url":   "https://github.com/example/db-repo/pull/42",
            "created_at": "2026-05-10T09:00:00Z",
            "updated_at": "2026-05-19T14:00:00Z",
            "user":       {"login": "anantharuban"},
            "head":       {"ref": "feature/sql-overhaul"},
            "base":       {"ref": "main"},
            "body": (
                "## Summary\n"
                "- Refactored `v_customer_summary` view to include loyalty tier\n"
                "- Updated `sp_update_order_status` procedure for new workflow states\n"
                "- Added new `pkg_inventory_mgmt` package for stock tracking\n"
                "- Removed deprecated `v_order_summary_old` view (replaced by pkg_inventory_mgmt)\n\n"
                "## Test plan\n"
                "- [ ] Run regression on order status transitions\n"
                "- [ ] Verify inventory deduction on new orders\n"
                "- [ ] Check view returns correct loyalty tier labels\n"
                "- [ ] Confirm no dependencies on removed view"
            ),
        }

        # ── File 1: SMALL change — view, 3 additions / 2 deletions ─────────────
        patch_small = (
            "@@ -1,14 +1,15 @@\n"
            " CREATE OR REPLACE VIEW v_customer_summary AS\n"
            " SELECT\n"
            "     c.customer_id,\n"
            "     c.full_name,\n"
            "     c.email,\n"
            "-    c.created_date,\n"
            "-    COUNT(o.order_id)  AS total_orders\n"
            "+    c.created_date,\n"
            "+    COUNT(o.order_id)     AS total_orders,\n"
            "+    NVL(lp.tier, 'BRONZE') AS loyalty_tier\n"
            " FROM customers c\n"
            " LEFT JOIN orders o  ON o.customer_id = c.customer_id\n"
            "+LEFT JOIN loyalty_points lp ON lp.customer_id = c.customer_id\n"
            " WHERE c.status = 'ACTIVE'\n"
            " GROUP BY\n"
            "     c.customer_id, c.full_name, c.email,\n"
            "-    c.created_date;\n"
            "+    c.created_date, lp.tier;\n"
        )

        # ── File 2: MEDIUM change — stored procedure, 18 add / 8 del ──────────
        patch_medium = (
            "@@ -1,6 +1,8 @@\n"
            " CREATE OR REPLACE PROCEDURE sp_update_order_status (\n"
            "     p_order_id   IN  orders.order_id%TYPE,\n"
            "     p_new_status IN  VARCHAR2,\n"
            "+    p_updated_by IN  VARCHAR2 DEFAULT USER,\n"
            "+    p_notes      IN  VARCHAR2 DEFAULT NULL,\n"
            "     p_result     OUT VARCHAR2\n"
            " ) AS\n"
            "@@ -12,18 +14,26 @@\n"
            "     v_current_status  VARCHAR2(30);\n"
            "     v_allowed_next    VARCHAR2(200);\n"
            "+    v_audit_id        audit_log.audit_id%TYPE;\n"
            " BEGIN\n"
            "     SELECT status\n"
            "     INTO   v_current_status\n"
            "     FROM   orders\n"
            "     WHERE  order_id = p_order_id;\n"
            " \n"
            "-    IF p_new_status NOT IN ('PENDING','PROCESSING','SHIPPED','DELIVERED','CANCELLED') THEN\n"
            "+    IF p_new_status NOT IN ('PENDING','PROCESSING','PICKED','SHIPPED',\n"
            "+                            'OUT_FOR_DELIVERY','DELIVERED','CANCELLED','RETURNED') THEN\n"
            "         p_result := 'ERROR: invalid status';\n"
            "         RETURN;\n"
            "     END IF;\n"
            " \n"
            "-    UPDATE orders SET status = p_new_status WHERE order_id = p_order_id;\n"
            "+    UPDATE orders\n"
            "+    SET    status      = p_new_status,\n"
            "+           updated_by  = p_updated_by,\n"
            "+           updated_date = SYSDATE\n"
            "+    WHERE  order_id = p_order_id;\n"
            " \n"
            "+    INSERT INTO order_status_history (order_id, old_status, new_status,\n"
            "+                                      changed_by, changed_date, notes)\n"
            "+    VALUES (p_order_id, v_current_status, p_new_status,\n"
            "+            p_updated_by, SYSDATE, p_notes);\n"
            "+\n"
            "-    COMMIT;\n"
            "+    COMMIT;\n"
            "     p_result := 'OK';\n"
            " EXCEPTION\n"
            "     WHEN OTHERS THEN\n"
            "-        ROLLBACK;\n"
            "+        ROLLBACK;\n"
            "         p_result := 'ERROR: ' || SQLERRM;\n"
            " END sp_update_order_status;\n"
            " /\n"
        )

        # ── File 3: BIG change — new package, 72 additions ─────────────────────
        patch_big = (
            "@@ -0,0 +1,72 @@\n"
            "+CREATE OR REPLACE PACKAGE pkg_inventory_mgmt AS\n"
            "+\n"
            "+    -- Reserve stock for an order line\n"
            "+    PROCEDURE reserve_stock (\n"
            "+        p_product_id  IN  NUMBER,\n"
            "+        p_qty         IN  NUMBER,\n"
            "+        p_order_id    IN  NUMBER,\n"
            "+        p_result      OUT VARCHAR2\n"
            "+    );\n"
            "+\n"
            "+    -- Release reserved stock (cancellation / rejection)\n"
            "+    PROCEDURE release_stock (\n"
            "+        p_product_id  IN  NUMBER,\n"
            "+        p_qty         IN  NUMBER,\n"
            "+        p_order_id    IN  NUMBER,\n"
            "+        p_result      OUT VARCHAR2\n"
            "+    );\n"
            "+\n"
            "+    -- Confirm stock deduction on dispatch\n"
            "+    PROCEDURE confirm_dispatch (\n"
            "+        p_order_id    IN  NUMBER,\n"
            "+        p_result      OUT VARCHAR2\n"
            "+    );\n"
            "+\n"
            "+    -- Returns current available qty (on_hand - reserved)\n"
            "+    FUNCTION get_available_qty (\n"
            "+        p_product_id  IN  NUMBER\n"
            "+    ) RETURN NUMBER;\n"
            "+\n"
            "+END pkg_inventory_mgmt;\n"
            "+/\n"
            "+\n"
            "+CREATE OR REPLACE PACKAGE BODY pkg_inventory_mgmt AS\n"
            "+\n"
            "+    PROCEDURE reserve_stock (p_product_id IN NUMBER, p_qty IN NUMBER,\n"
            "+                             p_order_id IN NUMBER, p_result OUT VARCHAR2) AS\n"
            "+        v_avail NUMBER;\n"
            "+    BEGIN\n"
            "+        SELECT on_hand_qty - reserved_qty\n"
            "+        INTO   v_avail\n"
            "+        FROM   inventory\n"
            "+        WHERE  product_id = p_product_id\n"
            "+        FOR UPDATE;\n"
            "+\n"
            "+        IF v_avail < p_qty THEN\n"
            "+            p_result := 'ERROR: insufficient stock (' || v_avail || ' available)';\n"
            "+            RETURN;\n"
            "+        END IF;\n"
            "+\n"
            "+        UPDATE inventory\n"
            "+        SET    reserved_qty = reserved_qty + p_qty\n"
            "+        WHERE  product_id = p_product_id;\n"
            "+\n"
            "+        INSERT INTO inventory_reservations (product_id, order_id, qty, reserved_date)\n"
            "+        VALUES (p_product_id, p_order_id, p_qty, SYSDATE);\n"
            "+\n"
            "+        COMMIT;\n"
            "+        p_result := 'OK';\n"
            "+    EXCEPTION WHEN OTHERS THEN ROLLBACK; p_result := 'ERROR: ' || SQLERRM;\n"
            "+    END reserve_stock;\n"
            "+\n"
            "+    PROCEDURE release_stock (p_product_id IN NUMBER, p_qty IN NUMBER,\n"
            "+                             p_order_id IN NUMBER, p_result OUT VARCHAR2) AS\n"
            "+    BEGIN\n"
            "+        UPDATE inventory\n"
            "+        SET    reserved_qty = GREATEST(0, reserved_qty - p_qty)\n"
            "+        WHERE  product_id = p_product_id;\n"
            "+\n"
            "+        DELETE FROM inventory_reservations\n"
            "+        WHERE  product_id = p_product_id AND order_id = p_order_id;\n"
            "+\n"
            "+        COMMIT; p_result := 'OK';\n"
            "+    EXCEPTION WHEN OTHERS THEN ROLLBACK; p_result := 'ERROR: ' || SQLERRM;\n"
            "+    END release_stock;\n"
            "+\n"
            "+    PROCEDURE confirm_dispatch (p_order_id IN NUMBER, p_result OUT VARCHAR2) AS\n"
            "+    BEGIN\n"
            "+        UPDATE inventory i\n"
            "+        SET    on_hand_qty  = on_hand_qty  - ir.qty,\n"
            "+               reserved_qty = reserved_qty - ir.qty\n"
            "+        FROM   inventory_reservations ir\n"
            "+        WHERE  ir.product_id = i.product_id AND ir.order_id = p_order_id;\n"
            "+\n"
            "+        DELETE FROM inventory_reservations WHERE order_id = p_order_id;\n"
            "+        COMMIT; p_result := 'OK';\n"
            "+    EXCEPTION WHEN OTHERS THEN ROLLBACK; p_result := 'ERROR: ' || SQLERRM;\n"
            "+    END confirm_dispatch;\n"
            "+\n"
            "+    FUNCTION get_available_qty (p_product_id IN NUMBER) RETURN NUMBER AS\n"
            "+        v_qty NUMBER := 0;\n"
            "+    BEGIN\n"
            "+        SELECT on_hand_qty - NVL(reserved_qty, 0)\n"
            "+        INTO   v_qty FROM inventory WHERE product_id = p_product_id;\n"
            "+        RETURN v_qty;\n"
            "+    EXCEPTION WHEN NO_DATA_FOUND THEN RETURN 0;\n"
            "+    END get_available_qty;\n"
            "+\n"
            "+END pkg_inventory_mgmt;\n"
            "+/\n"
        )

        # ── File 4: DELETED file — deprecated view, 14 deletions ──────────────
        patch_deleted = (
            "@@ -1,14 +0,0 @@\n"
            "-CREATE OR REPLACE VIEW v_order_summary_old AS\n"
            "-SELECT\n"
            "-    o.order_id,\n"
            "-    o.customer_id,\n"
            "-    o.status,\n"
            "-    o.order_total,\n"
            "-    o.created_date,\n"
            "-    o.updated_date,\n"
            "-    c.full_name        AS customer_name,\n"
            "-    c.email            AS customer_email\n"
            "-FROM orders o\n"
            "-JOIN customers c ON c.customer_id = o.customer_id\n"
            "-WHERE o.created_date < ADD_MONTHS(SYSDATE, -24)\n"
            "-  AND o.status IN ('CANCELLED', 'RETURNED');\n"
        )

        # Full post-change file content for modified files (simulates raw_url fetch)
        full_small = (
            "CREATE OR REPLACE VIEW v_customer_summary AS\n"
            "SELECT\n"
            "    c.customer_id,\n"
            "    c.full_name,\n"
            "    c.email,\n"
            "    c.created_date,\n"
            "    COUNT(o.order_id)       AS total_orders,\n"
            "    NVL(lp.tier, 'BRONZE')  AS loyalty_tier\n"
            "FROM customers c\n"
            "LEFT JOIN orders o         ON o.customer_id  = c.customer_id\n"
            "LEFT JOIN loyalty_points lp ON lp.customer_id = c.customer_id\n"
            "WHERE c.status = 'ACTIVE'\n"
            "GROUP BY\n"
            "    c.customer_id, c.full_name, c.email,\n"
            "    c.created_date, lp.tier;\n"
        )

        full_medium = (
            "CREATE OR REPLACE PROCEDURE sp_update_order_status (\n"
            "    p_order_id   IN  orders.order_id%TYPE,\n"
            "    p_new_status IN  VARCHAR2,\n"
            "    p_updated_by IN  VARCHAR2 DEFAULT USER,\n"
            "    p_notes      IN  VARCHAR2 DEFAULT NULL,\n"
            "    p_result     OUT VARCHAR2\n"
            ") AS\n"
            "    v_current_status  VARCHAR2(30);\n"
            "    v_allowed_next    VARCHAR2(200);\n"
            "    v_audit_id        audit_log.audit_id%TYPE;\n"
            "BEGIN\n"
            "    SELECT status\n"
            "    INTO   v_current_status\n"
            "    FROM   orders\n"
            "    WHERE  order_id = p_order_id;\n"
            "\n"
            "    IF p_new_status NOT IN ('PENDING','PROCESSING','PICKED','SHIPPED',\n"
            "                            'OUT_FOR_DELIVERY','DELIVERED','CANCELLED','RETURNED') THEN\n"
            "        p_result := 'ERROR: invalid status';\n"
            "        RETURN;\n"
            "    END IF;\n"
            "\n"
            "    UPDATE orders\n"
            "    SET    status       = p_new_status,\n"
            "           updated_by   = p_updated_by,\n"
            "           updated_date = SYSDATE\n"
            "    WHERE  order_id = p_order_id;\n"
            "\n"
            "    INSERT INTO order_status_history\n"
            "           (order_id, old_status, new_status, changed_by, changed_date, notes)\n"
            "    VALUES (p_order_id, v_current_status, p_new_status,\n"
            "            p_updated_by, SYSDATE, p_notes);\n"
            "\n"
            "    COMMIT;\n"
            "    p_result := 'OK';\n"
            "EXCEPTION\n"
            "    WHEN OTHERS THEN\n"
            "        ROLLBACK;\n"
            "        p_result := 'ERROR: ' || SQLERRM;\n"
            "END sp_update_order_status;\n"
            "/\n"
        )

        self.pr_files = [
            {
                "filename":     "db/views/v_customer_summary.sql",
                "status":       "modified",
                "additions":    3,
                "deletions":    2,
                "changes":      5,
                "patch":        patch_small,
                "full_content": full_small,
                "blob_url":     "https://github.com/example/db-repo/blob/feature/sql-overhaul/db/views/v_customer_summary.sql",
                "raw_url":      "",
                "contents_url": "",
            },
            {
                "filename":     "db/procedures/sp_update_order_status.sql",
                "status":       "modified",
                "additions":    18,
                "deletions":    8,
                "changes":      26,
                "patch":        patch_medium,
                "full_content": full_medium,
                "blob_url":     "https://github.com/example/db-repo/blob/feature/sql-overhaul/db/procedures/sp_update_order_status.sql",
                "raw_url":      "",
                "contents_url": "",
            },
            {
                "filename":     "db/packages/pkg_inventory_mgmt.sql",
                "status":       "added",
                "additions":    72,
                "deletions":    0,
                "changes":      72,
                "patch":        patch_big,
                "full_content": None,
                "blob_url":     "https://github.com/example/db-repo/blob/feature/sql-overhaul/db/packages/pkg_inventory_mgmt.sql",
                "raw_url":      "",
                "contents_url": "",
            },
            {
                "filename":     "db/views/v_order_summary_old.sql",
                "status":       "removed",
                "additions":    0,
                "deletions":    14,
                "changes":      14,
                "patch":        patch_deleted,
                "full_content": None,
                "blob_url":     "",
                "raw_url":      "",
                "contents_url": "",
            },
        ]

        # Mock per-file review comments (simulate GitHub PR review comments)
        self.pr_review_comments = [
            {"path": "db/views/v_customer_summary.sql",
             "body": "NVL fallback to BRONZE looks good. Should we index loyalty_points.customer_id?",
             "user": {"login": "reviewer_alice"}, "created_at": "2026-05-11T10:22:00Z",
             "line": 8, "commit_id": "abc123"},
            {"path": "db/views/v_customer_summary.sql",
             "body": "GROUP BY clause updated correctly — approved.",
             "user": {"login": "reviewer_bob"}, "created_at": "2026-05-11T11:05:00Z",
             "line": 14, "commit_id": "abc123"},
            {"path": "db/procedures/sp_update_order_status.sql",
             "body": "New status values PICKED and OUT_FOR_DELIVERY match the logistics workflow spec.",
             "user": {"login": "reviewer_alice"}, "created_at": "2026-05-12T09:10:00Z",
             "line": 17, "commit_id": "def456"},
            {"path": "db/procedures/sp_update_order_status.sql",
             "body": "order_status_history insert is missing an index on order_id — raise a follow-up ticket.",
             "user": {"login": "reviewer_carol"}, "created_at": "2026-05-12T14:30:00Z",
             "line": 22, "commit_id": "def456"},
            {"path": "db/packages/pkg_inventory_mgmt.sql",
             "body": "Package spec looks clean. Confirm reserve_stock handles concurrent transactions via SELECT FOR UPDATE.",
             "user": {"login": "reviewer_bob"}, "created_at": "2026-05-13T08:55:00Z",
             "line": 5, "commit_id": "ghi789"},
        ]

        # Mock PR-level issue comments
        self.pr_issue_comments = [
            {"body": "LGTM overall. The inventory package is a great addition — much cleaner than the old inline logic.",
             "user": {"login": "reviewer_alice"}, "created_at": "2026-05-14T09:00:00Z"},
            {"body": "Merge after confirming loyalty_points index. Adding that as a follow-up Jira ticket.",
             "user": {"login": "reviewer_carol"}, "created_at": "2026-05-14T15:45:00Z"},
        ]

        # Build file_comments from mock review comments — sort newest-first, use latest per file
        from collections import defaultdict
        rev_by_file = defaultdict(list)
        for c in sorted(self.pr_review_comments,
                        key=lambda x: x.get("updated_at") or x.get("created_at", ""),
                        reverse=True):
            path = c.get("path", "")
            body = c.get("body", "").strip()
            user = (c.get("user") or {}).get("login", "")
            if path and body:
                rev_by_file[path].append(f"{user}: {body}" if user else body)

        self.file_comments = []
        for f in self.pr_files:
            fname    = f["filename"]
            comments = rev_by_file.get(fname, [])
            jira_cmt = comments[0] if comments else ""
            self.file_comments.append(dict(
                filename=fname, status=f["status"],
                additions=f["additions"], deletions=f["deletions"],
                word_comment=jira_cmt, jira_comment=jira_cmt,
                comment=jira_cmt, screenshots=[]))

        pr = self.pr_cache
        self.pr_info_var.set(
            f"  #{pr['number']}  {pr['title']}   "
            f"{pr['state'].upper()}   by {pr['user']['login']}   "
            f"2026-05-10  [MOCK DATA]"
        )
        self.files_sum_var.set(
            "4 files changed   A 1 added   M 2 modified   D 1 removed"
        )
        if hasattr(self, "sql_btn"):
            self.sql_btn.config(state="normal")
        if hasattr(self, "_gh_preview_btn"):
            self._gh_preview_btn.config(state="normal")
        # Auto-fill Jira Summary from PR title
        jfw = getattr(self, "_jira_field_widgets", {})
        if "summary" in jfw:
            jfw["summary"].set(pr["title"])

        # Auto-fill Description from PR body
        pr_body = (pr.get("body") or "").strip()
        if pr_body:
            self.desc_text.delete("1.0", "end")
            self.desc_text.config(fg=C["text"])
            self.desc_text.insert("1.0", pr_body)

        # Auto-fill Reporter from PR creator
        if "reporter" in jfw:
            jfw["reporter"].set(pr["user"]["login"])

        # Auto-fill Assignee from first mock reviewer
        if "assignee" in jfw:
            jfw["assignee"].set("reviewer_alice")

        self._set_status("Mock data loaded — 4 SQL files · 5 review comments · auto-filled summary, description, reporter & assignee")

    # ── Preview ───────────────────────────────────────────────────────────────

    def _show_gh_preview(self):
        if not self.pr_cache:
            messagebox.showinfo("No PR", "Fetch a PR first.", parent=self)
            return
        GitHubPRPreviewPopup(
            self,
            pr=self.pr_cache,
            files=self.pr_files,
            review_comments=getattr(self, "pr_review_comments", []),
            issue_comments=getattr(self, "pr_issue_comments", []),
        )

    def _show_preview(self):
        fv = self._collect_fields()
        html_content = build_preview_html(
            fv, self.config_data["fields"], self.pr_cache, self.file_comments)

        popup_ref = [None]

        def _on_generate(updated_comments):
            self.file_comments = updated_comments

            def _after_save(path):
                p = popup_ref[0]
                if p and p.winfo_exists():
                    p._enable_open_word(path)

            self._gen_word_thread(after_save=_after_save)

        def _build_html(fc):
            fv2 = self._collect_fields()
            return build_preview_html(fv2, self.config_data["fields"], self.pr_cache, fc)

        popup = DocPreviewWindow(self, html_content, self._last_doc_path,
                                 pr_files=self.pr_files,
                                 file_comments=self.file_comments,
                                 on_comments_saved=_on_generate,
                                 build_html_func=_build_html)
        popup_ref[0] = popup

    # ── Word doc ─────────────────────────────────────────────────────────────

    def _gen_word_thread(self, after_save=None):
        threading.Thread(target=self._gen_word, args=(after_save,), daemon=True).start()

    def _gen_word(self, after_save=None):
        self.after(0, self.progress.start)
        self.after(0, lambda: self._set_status("Generating Word document..."))
        fv      = self._collect_fields()
        out_dir = self.config_data.get("word_doc_output_dir", BASE_DIR)
        try:
            path = generate_word_doc(fv, self.config_data["fields"],
                                     self.pr_cache, self.file_comments, out_dir,
                                     keyword_findings=getattr(self, "_sql_keyword_findings", {}))
        except Exception as e:
            self.after(0, lambda m=str(e): messagebox.showerror("Word Error", m))
            self.after(0, self.progress.stop); return
        self._last_doc_path = path
        self.after(0, self.progress.stop)
        self.after(0, lambda: self._set_status("Word document saved"))
        if self.config_data.get("auto_open_word_doc") and path:
            self.after(0, lambda p=path: os.startfile(p))
        if after_save:
            self.after(0, lambda p=path: after_save(p))

    # ── Jira ticket ───────────────────────────────────────────────────────────

    def _create_jira_thread(self):
        threading.Thread(target=self._create_jira, daemon=True).start()

    def _create_jira(self):
        self.after(0, self.progress.start)
        self.after(0, lambda: self._set_status("Creating Jira ticket..."))

        fv = self._collect_fields()

        # validate required fields
        for fd in self.config_data["fields"]:
            if fd.get("required") and fd.get("enabled", True):
                if not fv.get(fd["key"]):
                    msg = f"{fd['label']} is required"
                    self.after(0, lambda m=msg: messagebox.showerror("Validation", m))
                    self.after(0, self.progress.stop); return

        try:
            jira_tok = read_token(self.config_data["jira_token_file"])
        except FileNotFoundError as e:
            self.after(0, lambda m=str(e): messagebox.showerror("Token Error", m))
            self.after(0, self.progress.stop); return

        email    = self.config_data.get("jira_email", "")
        base_url = self.config_data.get("jira_base_url", "").rstrip("/")
        proj_key = self.config_data.get("jira_project_key", "")

        # build plain-text description (Jira REST API v2 — string, not ADF)
        # order: description → changed files → issue details table
        desc_lines = []
        desc_lines.append(fv.get("description") or "")

        changed_files = self._build_changed_files()
        if changed_files:
            desc_lines.append("")
            desc_lines.append("h3. Changed Files")
            for fn, status, jira_cmt, sql_cmt in changed_files:
                line = f"{STATUS_EMOJI.get(status,'~')} {fn}  [{status.upper()}]"
                if jira_cmt:
                    line += f"\n  {jira_cmt}"
                if sql_cmt:
                    line += f"\n  SQL: {sql_cmt}"
                desc_lines.append(f"* {line}")

        include_details = getattr(self, "_include_in_desc_var", None)
        include_details = include_details.get() if include_details else True
        if include_details:
            detail_rows = []
            for fd in self.config_data["fields"]:
                if not fd.get("enabled", True): continue
                if fd.get("jira_field"): continue
                if fd["key"] in ("doc_name", "doc_title"): continue
                val = fv.get(fd["key"], "")
                if val:
                    detail_rows.append((fd["label"], val))
            if detail_rows:
                desc_lines.append("")
                desc_lines.append("h3. Issue Details")
                desc_lines.append("||Field||Value||")
                for lbl, val in detail_rows:
                    desc_lines.append(f"|{lbl}|{val}|")

        # keyword findings from SQL generation
        kw_findings = getattr(self, "_sql_keyword_findings", {})
        if kw_findings:
            desc_lines.append("")
            desc_lines.append("h3. ⚠ Keyword Findings (Test Cases)")
            desc_lines.append("||File||Keywords Found||")
            for fn, kws in kw_findings.items():
                desc_lines.append(f"|{fn}|{', '.join(kws)}|")

        description = "\n".join(desc_lines)

        type_map = {"Bug":"Bug","Enhancement":"Story","Task":"Task",
                    "Story":"Story","Epic":"Epic","Sub-task":"Sub-task","Incident":"Bug"}
        issue_type = type_map.get(fv.get("issue_type","Bug"), "Bug")

        payload = {"fields": {
            "project":     {"key": fv.get("project") or proj_key},
            "summary":     fv.get("summary") or "(no summary)",
            "issuetype":   {"name": issue_type},
            "description": description,
        }}

        if fv.get("reporter"):
            payload["fields"]["reporter"] = {"name": fv["reporter"]}

        # additional enabled Jira API fields (beyond the 4 core hardcoded above)
        _CORE_KEYS  = {"project", "summary", "issue_type", "reporter"}
        _MULTI_OBJ  = {"components", "fixVersions", "versions"}
        _MULTI_STR  = {"labels"}
        _NAME_FIELDS= {"priority", "issuetype", "assignee"}
        for fd in self.config_data["fields"]:
            if not fd.get("jira_field"):         continue
            if not fd.get("enabled", True):      continue
            if fd["key"] in _CORE_KEYS:          continue
            jk  = fd.get("jira_key")
            if not jk:                           continue
            val = fv.get(fd["key"], "").strip()
            if not val:                          continue
            if jk in _MULTI_OBJ:
                payload["fields"][jk] = [{"name": v.strip()} for v in val.split(",") if v.strip()]
            elif jk in _MULTI_STR:
                payload["fields"][jk] = [v.strip() for v in val.split(",") if v.strip()]
            elif jk in _NAME_FIELDS:
                payload["fields"][jk] = {"name": val}
            elif fd.get("type") == "number":
                try:
                    payload["fields"][jk] = float(val) if "." in val else int(val)
                except ValueError:
                    payload["fields"][jk] = val
            else:
                payload["fields"][jk] = val

        method = self.config_data.get("jira_auth_method", "basic")
        hdrs   = jira_auth_headers(email, jira_tok, method)
        verify = _jira_ssl_verify(self.config_data)
        try:
            resp = requests.post(f"{base_url}/issue",
                                 json=payload, headers=hdrs, timeout=20, verify=verify)
            resp.raise_for_status()
            data = resp.json()
        except requests.HTTPError as e:
            msg = f"Jira API error {e.response.status_code}:\n{e.response.text[:400]}"
            self.after(0, lambda m=msg: messagebox.showerror("Jira Error", m))
            self.after(0, self.progress.stop); return
        except Exception as e:
            self.after(0, lambda m=str(e): messagebox.showerror("Network Error", m))
            self.after(0, self.progress.stop); return

        issue_key = data.get("key", "")
        issue_url = f"{base_url}/browse/{issue_key}"
        self._jira_url       = issue_url
        self._last_issue_key = issue_key

        # auto-save issue details field values as last-used defaults
        save_field_values({k: v.get() for k, v in getattr(self, "_field_widgets", {}).items()})

        # generate Word doc now (so it's ready when user clicks Attach)
        try:
            doc_path = generate_word_doc(fv, self.config_data["fields"],
                                         self.pr_cache, self.file_comments,
                                         self.config_data.get("word_doc_output_dir", BASE_DIR),
                                         keyword_findings=getattr(self, "_sql_keyword_findings", {}))
            self._last_doc_path = doc_path
        except Exception:
            pass

        def _done():
            self.progress.stop()
            self.result_var.set(f"  Ticket created: {issue_key}   Click to open in Jira")
            self._set_status(f"Ticket {issue_key} created — click Attach Files to upload docs")
            pr_note = f"\n\nLinked PR: #{pr['number']} — {pr.get('title','')[:40]}" \
                      if pr.get("number") else ""
            messagebox.showinfo("Ticket Created",
                f"Jira ticket created!\n\n{issue_key}\n{issue_url}{pr_note}"
                f"\n\nUse 'Attach Files to Ticket' button to upload Word doc / SQL file.")
            if hasattr(self, "_attach_btn"):
                self._attach_btn.configure(
                    state="normal",
                    text=f"⊞  Attach Files to {issue_key}")
        self.after(0, _done)

    # ── Attach files to existing ticket ──────────────────────────────────────

    def _attach_files_thread(self):
        threading.Thread(target=self._attach_files_to_jira, daemon=True).start()

    def _attach_files_to_jira(self):
        issue_key = self._last_issue_key
        if not issue_key:
            self.after(0, lambda: messagebox.showerror(
                "No Ticket", "No ticket created yet in this session."))
            return

        cfg      = self.config_data
        tok_file = cfg.get("jira_token_file", "")
        email    = cfg.get("jira_email", "").strip()
        base_url = cfg.get("jira_base_url", "").rstrip("/")
        method   = cfg.get("jira_auth_method", "basic")
        verify   = _jira_ssl_verify(cfg)

        try:
            jira_tok = read_token(tok_file)
        except Exception as e:
            self.after(0, lambda m=str(e): messagebox.showerror("Token Error", m))
            return

        do_word = getattr(self._attach_word_var, "get", lambda: True)()
        do_sql  = getattr(self._attach_sql_var,  "get", lambda: True)()

        if not do_word and not do_sql:
            self.after(0, lambda: messagebox.showwarning(
                "Nothing Selected",
                "Select at least one file type to attach\n(Word Doc or SQL File)."))
            return

        self.after(0, self.progress.start)
        self.after(0, lambda: self._set_status(f"Attaching files to {issue_key}..."))

        def _post_file(path, mime):
            ah = jira_auth_headers(email, jira_tok, method)
            ah["X-Atlassian-Token"] = "no-check"
            ah.pop("Content-Type", None)
            with open(path, "rb") as fh:
                r = requests.post(
                    f"{base_url}/issue/{issue_key}/attachments",
                    headers=ah,
                    files={"file": (os.path.basename(path), fh, mime)},
                    timeout=30, verify=verify)
            r.raise_for_status()

        attached = []
        errors   = []

        if do_word:
            doc_path = self._last_doc_path
            if doc_path and os.path.exists(doc_path):
                try:
                    _post_file(doc_path,
                               "application/vnd.openxmlformats-officedocument"
                               ".wordprocessingml.document")
                    attached.append(os.path.basename(doc_path))
                except Exception as e:
                    errors.append(f"Word doc: {e}")
            else:
                errors.append("Word doc not found — generate it first via 'Generate Word Doc'.")

        if do_sql:
            sql_path = self._last_sql_path
            if sql_path and os.path.exists(sql_path):
                try:
                    _post_file(sql_path, "text/plain")
                    attached.append(os.path.basename(sql_path))
                except Exception as e:
                    errors.append(f"SQL file: {e}")
            else:
                errors.append("SQL file not found — generate it first via 'SQL PR File'.")

        def _done():
            self.progress.stop()
            if attached:
                self._set_status(f"Attached {len(attached)} file(s) to {issue_key}")
                self._attach_btn.configure(
                    text=f"✓  Attached to {issue_key}", state="disabled")
            msg = ""
            if attached:
                msg += f"Attached to {issue_key}:\n" + "\n".join(f"  • {a}" for a in attached)
            if errors:
                msg += ("\n\n" if msg else "") + "Errors:\n" + "\n".join(f"  • {e}" for e in errors)
            if msg:
                (messagebox.showinfo if attached else messagebox.showerror)(
                    "Attach Files", msg)
        self.after(0, _done)

    # ── helpers ───────────────────────────────────────────────────────────────

    def _build_changed_files(self):
        """Return list of (filename, status, jira_comment, sql_comment) for all changed files.
        Merges file_comments (from word doc flow) and _sql_file_comments (from SQL popup)."""
        sql_cmts  = getattr(self, "_sql_file_comments", {})
        fc_by_name = {fc["filename"]: fc for fc in self.file_comments}

        # ordered unique filenames: file_comments order first, then any SQL-only files
        all_fnames = list(dict.fromkeys(
            [fc["filename"] for fc in self.file_comments] + list(sql_cmts.keys())
        ))

        rows = []
        pr_status = {pf["filename"]: pf.get("status", "modified")
                     for pf in getattr(self, "pr_files", [])}
        for fn in all_fnames:
            fc        = fc_by_name.get(fn, {})
            status    = fc.get("status") or pr_status.get(fn, "modified")
            jira_cmt  = (fc.get("jira_comment") or fc.get("comment", "")).strip()
            sql_cmt   = sql_cmts.get(fn, "")
            rows.append((fn, status, jira_cmt, sql_cmt))
        return rows

    def _collect_fields(self):
        fv = {}
        for key, var in self._field_widgets.items():
            fv[key] = var.get().strip()
        for key, var in getattr(self, "_jira_field_widgets", {}).items():
            fv[key] = var.get().strip()
        desc = self.desc_text.get("1.0", "end-1c").strip()
        if desc == "Describe the issue in detail...":
            desc = ""
        fv["description"] = desc
        fv["pr_url"]      = self.pr_url_var.get().strip()
        return fv

    def _choose_doc_dir(self):
        p = filedialog.askdirectory()
        if p:
            self.config_data["word_doc_output_dir"] = p

    def _open_word(self):
        import subprocess, sys
        if not self._last_doc_path or not os.path.exists(self._last_doc_path):
            messagebox.showinfo("No document", "Generate a Word doc first."); return
        if sys.platform == "win32":
            os.startfile(self._last_doc_path)
        elif sys.platform == "darwin":
            subprocess.run(["open", self._last_doc_path])
        else:
            subprocess.run(["xdg-open", self._last_doc_path])

    def _set_status(self, msg, color=None):
        self.status_var.set(msg)
        if hasattr(self, "_status_dot"):
            if color:
                self._status_dot.config(fg=color)
            elif any(w in msg.lower() for w in ("error", "failed", "not found")):
                self._status_dot.config(fg=C["red"])
            elif any(w in msg.lower() for w in ("fetching", "generating", "creating", "loading")):
                self._status_dot.config(fg=C["yellow"])
            else:
                self._status_dot.config(fg=C["green"])


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    app = App()
    app.mainloop()
